
# Standard Library
import logging
import uuid
from random import randint
from io import BytesIO

# Third-Party
import pydf
from docx import Document
from django_fsm import FSMIntegerField
from django_fsm import transition
from django_fsm_log.decorators import fsm_log_by
from django_fsm_log.models import StateLog
from dry_rest_permissions.generics import allow_staff_or_superuser
from dry_rest_permissions.generics import authenticated_users
from model_utils import Choices
from model_utils.models import TimeStampedModel
from django.db.models import Sum, Max, Avg, StdDev, Count, Q, F, Func
from django.contrib.postgres.fields import ArrayField, JSONField
from django_fsm import RETURN_VALUE
from django.db.models.functions import DenseRank, RowNumber, Rank
from django.core.validators import MaxValueValidator
from django.core.validators import MinValueValidator
from django.db import models
from django.db.models import Min, Max, Count, Avg
from django.core.exceptions import ValidationError
# Django
from django.apps import apps
from django.contrib.contenttypes.fields import GenericRelation
from django.core.files.base import ContentFile
from django.db import models
from django.db.models import F, Window
from django.template.loader import render_to_string
from django.utils.functional import cached_property
from django.conf import settings
from django.utils.timezone import now

from .tasks import build_email
from .tasks import send_publish_email_from_round
from .tasks import send_publish_report_email_from_round
from .tasks import send_psa_email_from_panelist
from .tasks import save_psa_from_panelist
from .tasks import send_complete_email_from_appearance
from .tasks import save_reports_from_round

from .fields import FileUploadPath


from .managers import AppearanceManager
from .managers import PanelistManager
from .managers import SongManager
from .managers import ScoreManager

log = logging.getLogger(__name__)


class Appearance(TimeStampedModel):
    """
    An appearance of a group on stage.

    The Appearance is meant to be a private resource.
    """

    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-30, 'disqualified', 'Disqualified',),
        (-20, 'scratched', 'Scratched',),
        (-10, 'completed', 'Completed',),
        (0, 'new', 'New',),
        (7, 'built', 'Built',),
        (10, 'started', 'Started',),
        (20, 'finished', 'Finished',),
        (25, 'variance', 'Variance',),
        (30, 'verified', 'Verified',),
        (40, 'advanced', 'Advanced',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
        help_text="""The order of appearance for this round.""",
    )

    draw = models.IntegerField(
        help_text="""The draw for the next round.""",
        null=True,
        blank=True,
    )

    is_private = models.BooleanField(
        help_text="""Copied from entry.""",
        default=False,
    )

    is_single = models.BooleanField(
        help_text="""Single-round group""",
        default=False,
    )

    participants = models.CharField(
        help_text='Director(s) or Members (listed TLBB)',
        max_length=255,
        blank=True,
        default='',
    )

    representing = models.CharField(
        help_text='Representing entity',
        max_length=255,
        blank=True,
        default='',
    )

    onstage = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    actual_start = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    actual_finish = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    pos = models.IntegerField(
        help_text='Actual Participants-on-Stage',
        null=True,
        blank=True,
    )

    legacy_group = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    stats = JSONField(
        null=True,
        blank=True,
    )

    base = models.FloatField(
        help_text="""
            The incoming base score used to determine most-improved winners.""",
        null=True,
        blank=True,
    )

    variance_report = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )

    csa = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )

    # Appearance FKs
    owners = models.ManyToManyField(
        settings.AUTH_USER_MODEL,
        related_name='appearances',
    )

    round = models.ForeignKey(
        'Round',
        related_name='appearances',
        on_delete=models.CASCADE,
    )

    group_id = models.UUIDField(
        null=True,
        blank=True,
    )

    entry = models.ForeignKey(
        'smanager.entry',
        related_name='appearances',
        null=True,
        blank=True,
        on_delete=models.SET_NULL,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='appearances',
    )

    @cached_property
    def round__kind(self):
        return self.round.kind

    @cached_property
    def run_total(self):
        Score = apps.get_model('rmanager.score')
        Panelist = apps.get_model('rmanager.panelist')
        run_total = Score.objects.filter(
            song__appearance__round__session=self.round.session,
            song__appearance__round__num__lte=self.round.num,
            song__appearance__group=self.group,
            panelist__kind=Panelist.KIND.official,
        ).aggregate(
            sum=Sum('points'),
            avg=Avg('points'),
            mus=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            sng=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
            per=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
        )
        return run_total


    # Appearance Internals
    objects = AppearanceManager()

    def clean(self):
        if self.group.kind != self.group.KIND.vlq:
            if self.group.kind != self.round.session.kind:
                raise ValidationError(
                    {'group': 'Group kind must match session'}
                )

    class Meta:
        ordering = [
            '-round__num',
            'num',
        ]
        # unique_together = (
        #     ('round', 'num'),
        # )

    class JSONAPIMeta:
        resource_name = "appearance"

    def __str__(self):
        return str(self.id)
        # return "{0} {1}".format(
        #     self.round,
        #     self.group,
        # )

    # Methods
    def get_variance(self):
        Chart = apps.get_model('bhs.chart')
        Person = apps.get_model('bhs.person')
        Score = apps.get_model('rmanager.score')
        Panelist = apps.get_model('rmanager.panelist')

        # Songs Block
        songs = self.songs.annotate(
            tot_score=Avg(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
        ).order_by('num')
        scores = Score.objects.filter(
            panelist__kind=Panelist.KIND.official,
            song__appearance=self,
        ).order_by(
            'category',
            # 'panelist__person__last_name',
            'song__num',
        )
        panelists = self.round.panelists.filter(
            kind=Panelist.KIND.official,
            category__gt=Panelist.CATEGORY.ca,
        ).order_by(
            'category',
            # 'person__last_name',
        )
        for panelist in panelists:
            person = Person.objects.get(id=panelist.person_id)
            panelist.person = person
        variances = []
        for song in songs:
            chart = Chart.objects.get(song.chart_id)
            song.chart = chart
            variances.extend(song.dixons)
            variances.extend(song.asterisks)
        variances = list(set(variances))
        tot_points = scores.aggregate(sum=Sum('points'))['sum']
        context = {
            'appearance': self,
            'songs': songs,
            'scores': scores,
            'panelists': panelists,
            'variances': variances,
            'tot_points': tot_points,
        }
        rendered = render_to_string('reports/variance.html', context)
        pdf = pydf.generate_pdf(
            rendered,
            enable_smart_shrinking=False,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
        )
        content = ContentFile(pdf)
        return content

    def save_variance(self):
        content = self.get_variance()
        self.variance_report.save("variance_report", content)

    def mock(self):
        # Mock Appearance
        Chart = apps.get_model('bhs.chart')
        Panelist = apps.get_model('rmanager.panelist')
        if self.group.kind == self.group.KIND.chorus:
            pos = self.group.members.filter(
                status=self.group.members.model.STATUS.active,
            ).count()
            self.pos = pos
        # Try to approximate reality
        prelim = getattr(getattr(self, "entry"), "prelim", None)
        if not prelim:
            prelim = self.group.appearances.annotate(
                avg=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    )
                )
            ).latest('round__date').avg or randint(60, 70)
        songs = self.songs.all()
        for song in songs:
            # song.chart = Chart.objects.filter(
            #     status=Chart.STATUS.active
            # ).order_by("?").first()
            song.save()
            scores = song.scores.all()
            for score in scores:
                d = randint(-3, 3)
                score.points = prelim + d
                score.save()
        if self.status == self.STATUS.new:
            raise RuntimeError("Out of state")
        if self.status == self.STATUS.built:
            self.start()
            self.finish()
            self.verify()
            return
        if self.status == self.STATUS.started:
            self.finish()
            self.verify()
            return
        if self.status == self.STATUS.finished:
            self.verify()
            return

    def check_variance(self):
        # Set flag
        variance = False
        # Run checks for all songs and save.
        for song in self.songs.all():
            asterisks = song.get_asterisks()
            if asterisks:
                song.asterisks = asterisks
                variance = True
            dixons = song.get_dixons()
            if dixons:
                song.dixons = dixons
                variance = True
            if variance:
                song.save()
        return variance

    def get_csa(self):
        Chart = apps.get_model('bhs.chart')
        Person = apps.get_model('bhs.person')
        Panelist = apps.get_model('rmanager.panelist')
        Song = apps.get_model('rmanager.song')
        Score = apps.get_model('rmanager.score')

        # Appearancers Block
        group = self.group
        stats = Score.objects.select_related(
            'song__appearance__group',
            'song__appearance__round__session',
            'song__appearance__round',
            'panelist',
        ).filter(
            song__appearance__group=self.group,
            song__appearance__round__session=self.round.session,
            panelist__kind=Panelist.KIND.official,
        ).aggregate(
            max=Max(
                'song__appearance__round__num',
            ),
            tot_points=Sum(
                'points',
            ),
            mus_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            tot_score=Avg(
                'points',
            ),
            mus_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
        )
        appearances = Appearance.objects.select_related(
            'group',
            'round',
            'round__session',
        ).prefetch_related(
            'songs__scores',
            'songs__scores__panelist',
        ).filter(
            group=self.group,
            round__session=self.round.session,
        ).annotate(
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
            tot_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
        )

        # Monkeypatch
        for key, value in stats.items():
            setattr(group, key, value)
        for appearance in appearances:
            songs = appearance.songs.prefetch_related(
                'scores',
                'scores__panelist',
            ).order_by(
                'num',
            ).annotate(
                tot_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
                tot_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )
            for song in songs:
                penalties_map = {
                    10: "†",
                    30: "‡",
                    40: "✠",
                    50: "✶",
                }
                items = " ".join([penalties_map[x] for x in song.penalties])
                song.penalties_patched = items
            appearance.songs_patched = songs
        group.appearances_patched = appearances

        # Panelists
        panelists = Panelist.objects.select_related(
            # 'person',
        ).filter(
            kind=Panelist.KIND.official,
            round__session=self.round.session,
            round__num=1,
            category__gt=10,
        ).order_by('num')

        # Score Block
        initials = []
        for panelist in panelists:
            person = Person.objects.get(id=panelist.person_id)
            initials.append(person.initials)


        # Hackalicious
        category_count = {
            'Music': 0,
            'Performance': 0,
            'Singing': 0,
        }
        for panelist in panelists:
            category_count[panelist.get_category_display()] += 1
        songs = Song.objects.filter(
            appearance__round__session=self.round.session,
            appearance__group=self.group,
        ).order_by(
            'appearance__round__kind',
            'num',
        )
        for song in songs:
            chart = Chart.objects.get(id=song.chart_id)
            song.chart = chart
            scores = song.scores.filter(
                panelist__kind=Panelist.KIND.official,
            ).order_by('panelist__num')
            class_map = {
                Panelist.CATEGORY.music: 'warning',
                Panelist.CATEGORY.performance: 'success',
                Panelist.CATEGORY.singing: 'info',
            }
            items = []
            for score in scores:
                items.append((score.points, class_map[score.panelist.category]))
            song.scores_patched = items

        # Category Block
        categories = {
            'Music': [],
            'Performance': [],
            'Singing': [],
        }
        # panelists from above
        for panelist in panelists:
            item = categories[panelist.get_category_display()]
            person = Person.objects.get(id=panelist.person_id)
            item.append(person.common_name)

        # Penalties Block
        array = Song.objects.filter(
            appearance__round__session=self.round.session,
            appearance__group=self.group,
            penalties__len__gt=0,
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            10: "† Score(s) penalized due to violation of Article IX.A.1 of the BHS Contest Rules.",
            30: "‡ Score(s) penalized due to violation of Article IX.A.2 of the BHS Contest Rules.",
            40: "✠ Score(s) penalized due to violation of Article IX.A.3 of the BHS Contest Rules.",
            50: "✶ Score(s) penalized due to violation of Article X.B of the BHS Contest Rules.",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        context = {
            'appearance': self,
            'round': self.round,
            'group': group,
            'initials': initials,
            'songs': songs,
            'categories': categories,
            'penalties': penalties,
            'category_count': category_count,
        }
        rendered = render_to_string(
            'reports/csa.html',
            context,
        )
        statelog = self.round.statelogs.latest('timestamp')
        footer = 'Published by {0} at {1}'.format(
            statelog.by,
            statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
        )
        file = pydf.generate_pdf(
            rendered,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def save_csa(self):
        content = self.get_csa()
        return self.csa.save('csa', content)


    def get_complete_email(self):
        Panelist = apps.get_model('rmanager.panelist')
        Score = apps.get_model('rmanager.score')
        Chart = apps.get_model('bhs.chart')

        # Context
        group = self.group
        stats = Score.objects.select_related(
            'song__appearance__group',
            'song__appearance__round__session',
            'song__appearance__round',
            'panelist',
        ).filter(
            song__appearance__group=self.group,
            song__appearance__round__session=self.round.session,
            panelist__kind=Panelist.KIND.official,
        ).aggregate(
            max=Max(
                'song__appearance__round__num',
            ),
            tot_points=Sum(
                'points',
            ),
            mus_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            tot_score=Avg(
                'points',
            ),
            mus_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
        )
        appearances = Appearance.objects.select_related(
            'group',
            'round',
            'round__session',
        ).prefetch_related(
            'songs__scores',
            'songs__scores__panelist',
        ).filter(
            group=self.group,
            round__session=self.round.session,
        ).annotate(
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
            tot_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
        )

        # Monkeypatch
        for key, value in stats.items():
            setattr(group, key, value)
        for a in appearances:
            songs = a.songs.prefetch_related(
                'scores',
                'scores__panelist',
            ).order_by(
                'num',
            ).annotate(
                tot_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
                tot_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )
            for song in songs:
                chart = Chart.objects.get(id=song.chart_id)
                song.chart = chart
                penalties_map = {
                    10: "†",
                    30: "‡",
                    40: "✠",
                    50: "✶",
                }
                items = " ".join([penalties_map[x] for x in song.penalties])
                song.penalties_patched = items
            a.songs_patched = songs
        group.appearances_patched = appearances
        context = {'group': group}

        template = 'emails/appearance_complete.txt'
        subject = "[Barberscore] CSA for {0}".format(
            self.group.name,
        )
        to = self.group.get_officer_emails()
        cc = self.round.session.convention.get_drcj_emails()
        cc.extend(self.round.session.convention.get_ca_emails())

        if self.csa:
            pdf = self.csa.file
        else:
            pdf = self.get_csa()
        file_name = '{0} CSA.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]
        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email

    def send_complete_email(self):
        if self.status != self.STATUS.completed:
            raise ValueError("Do not send CSAs unless completed")
        email = self.get_complete_email()
        return email.send()

    # Appearance Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return any([
            self.round.status == self.round.STATUS.published,
            self.round.owners.filter(id__contains=request.user.id),
        ])

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.round.owners.filter(id__contains=request.user.id),
                self.round.status != self.round.STATUS.published,
            ]),
        ])

    # Appearance Conditions
    def can_verify(self):
        try:
            if self.group.kind == self.group.KIND.chorus and not self.pos:
                is_pos = False
            else:
                is_pos = True
        except AttributeError:
            is_pos = False
        return all([
            is_pos,
            not self.songs.filter(scores__points__isnull=True),
        ])

    # Appearance Transitions
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.new],
        target=STATUS.built,
    )
    def build(self, *args, **kwargs):
        """Sets up the Appearance."""
        Panelist = apps.get_model('rmanager.panelist')
        panelists = self.round.panelists.filter(
            category__gt=Panelist.CATEGORY.ca,
        )
        i = 1
        while i <= 2:  # Number songs constant
            song = self.songs.create(
                num=i
            )
            for panelist in panelists:
                song.scores.create(
                    panelist=panelist,
                )
            i += 1
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.built],
        target=STATUS.started,
    )
    def start(self, *args, **kwargs):
        """Indicates when they start singing."""
        self.actual_start = now()
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.started],
        target=STATUS.finished,
    )
    def finish(self, *args, **kwargs):
        """Indicates when they finish singing."""
        self.actual_finish = now()
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.finished, STATUS.variance],
        target=RETURN_VALUE(STATUS.variance, STATUS.verified,),
        conditions=[can_verify],
    )
    def verify(self, *args, **kwargs):
        # Checks for variance.  Returns Verified or Variance accordingly.
        if self.status == self.STATUS.finished:
            variance = self.check_variance()
            if variance:
                # Run variance report and save file.
                self.save_variance()
        # Variance is only checked once.
        else:
            variance = False
        return self.STATUS.variance if variance else self.STATUS.verified

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.verified, STATUS.advanced],
        target=STATUS.advanced,
    )
    def advance(self, *args, **kwargs):
        # Advances the Group.
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.verified, STATUS.advanced, STATUS.completed],
        target=STATUS.completed,
    )
    def complete(self, *args, **kwargs):
        # Completes the Group.
        # Saves CSA via post-transition signal to avoid race condition
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[
            STATUS.new,
            STATUS.built,
        ],
        target=STATUS.scratched,
    )
    def scratch(self, *args, **kwargs):
        # Scratches the group.
        # Remove songs
        self.songs.delete()
        return

    @fsm_log_by
    @transition(
        field=status,
        source='*',
        target=STATUS.disqualified,
    )
    def disqualify(self, *args, **kwargs):
        # Disqualify the group.
        self.songs.delete()
        # Notify the group?
        return


class Contender(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-10, 'excluded', 'Excluded',),
        (0, 'new', 'New',),
        (10, 'included', 'Included',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    # FKs
    appearance = models.ForeignKey(
        'Appearance',
        related_name='contenders',
        on_delete=models.CASCADE,
    )

    outcome = models.ForeignKey(
        'Outcome',
        related_name='contenders',
        on_delete=models.CASCADE,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='contenders',
    )

    # Internals
    class Meta:
        ordering = (
            'outcome__num',
        )
        # unique_together = (
        #     ('appearance', 'outcome',),
        # )

    class JSONAPIMeta:
        resource_name = "contender"

    def __str__(self):
        return str(self.id)

    # Methods

    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return True

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            'SCJC' in request.user.roles,
            all([
                self.outcome.round.owners.filter(id__contains=request.user.id),
                self.outcome.round.status != self.outcome.round.STATUS.started,
            ]),
        ])

    # contender Transitions
    @fsm_log_by
    @transition(field=status, source='*', target=STATUS.included)
    def include(self, *args, **kwargs):
        return

    @fsm_log_by
    @transition(field=status, source='*', target=STATUS.excluded)
    def exclude(self, *args, **kwargs):
        return


class Outcome(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-10, 'inactive', 'Inactive',),
        (0, 'new', 'New',),
        (10, 'active', 'Active',),
    )

    status = models.IntegerField(
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
        blank=True,
        null=True,
    )

    name = models.CharField(
        max_length=1024,
        null=True,
        blank=True,
    )

    legacy_num = models.IntegerField(
        blank=True,
        null=True,
    )

    legacy_name = models.CharField(
        max_length=1024,
        null=True,
        blank=True,
    )

    # FKs
    round = models.ForeignKey(
        'Round',
        related_name='outcomes',
        on_delete=models.CASCADE,
    )

    award = models.ForeignKey(
        'cmanager.award',
        related_name='outcomes',
        null=True,
        blank=True,
        on_delete=models.SET_NULL,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='outcomes',
    )

    # Methods
    def get_name(self):
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('rmanager.panelist')
        if self.round.kind != self.round.KIND.finals and not self.award.is_single:
            return "(Result determined in Finals)"
        if self.award.level == self.award.LEVEL.deferred:
            return "(Result determined post-contest)"
        if self.award.level in [self.award.LEVEL.manual, self.award.LEVEL.raw, self.award.LEVEL.standard]:
            return "MUST ENTER WINNER MANUALLY"
        # if self.award.level == self.award.LEVEL.raw:
        #     group_ids = Group.objects.filter(
        #         appearances__contenders__outcome=self,
        #     ).values_list('id', flat=True)
        #     group = Group.objects.filter(
        #         id__in=group_ids,
        #     ).annotate(
        #         avg=Avg(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         base=Avg(
        #             'appearances__base',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #             ),
        #         ),
        #         diff=F('avg') - F('base'),
        #         sng=Sum(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #                 appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing
        #             ),
        #         ),
        #         per=Sum(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #                 appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance
        #             ),
        #         ),
        #     ).filter(
        #         diff__gt=0,
        #     ).order_by(
        #        '-diff',
        #     ).first()
        #     try:
        #         winner = group.name
        #     except AttributeError:
        #         winner = "(No Award Winner)"
        #     return winner
        # if self.award.level == self.award.LEVEL.standard:
        #     group_ids = Group.objects.filter(
        #         appearances__contenders__outcome=self,
        #     ).values_list('id', flat=True)
        #     groups = Group.objects.filter(
        #         id__in=group_ids,
        #     ).annotate(
        #         score=Avg(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         base=Avg(
        #             'appearances__base',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #             ),
        #         ),
        #     )
        #     totals = groups.aggregate(
        #         Avg('score'),
        #         Avg('base'),
        #     )
        #     winner = groups.annotate(
        #         score_ratio=F('score') / totals['score__avg'],
        #         base_ratio=F('base') / totals['base__avg'],
        #         diff=F('score_ratio') - F('base_ratio'),
        #     ).order_by(
        #         '-diff',
        #     ).first().name
        if self.award.level == self.award.LEVEL.qualifier:
            threshold = self.award.threshold
            group_ids = Group.objects.filter(
                appearances__contenders__outcome=self,
            ).values_list('id', flat=True)
            qualifiers = Group.objects.filter(
                id__in=group_ids,
            ).annotate(
                avg=Avg(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__round__session=self.round.session,
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
            ).filter(
                avg__gte=threshold,
            ).order_by(
                'name',
            ).values_list('name', flat=True)
            if qualifiers:
                return ", ".join(qualifiers)
            return "(No Qualifiers)"
        if self.award.level in [self.award.LEVEL.championship, self.award.LEVEL.representative]:
            winner = Group.objects.filter(
                appearances__contenders__outcome=self,
            ).annotate(
                tot=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                sng=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing
                    ),
                ),
                per=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance
                    ),
                ),
            ).earliest(
                '-tot',
                '-sng',
                '-per',
            )
            if winner:
                return str(winner.name)
            return "(No Recipient)"
        raise RuntimeError("Level mismatch")

    # Internals
    # class Meta:
    #     unique_together = (
    #         ('round', 'award',),
    #         ('round', 'num',),
    #     )

    class JSONAPIMeta:
        resource_name = "outcome"

    def __str__(self):
        return str(self.id)

    def clean(self):
        pass

    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return any([
            self.round.status == self.round.STATUS.published,
            self.round.owners.filter(id__contains=request.user.id),
        ])


    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.round.owners.filter(id__contains=request.user.id),
                self.round.status < self.round.STATUS.verified,
            ]),
        ])


class Panelist(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-10, 'inactive', 'Inactive',),
        (-5, 'released', 'Released',),
        (0, 'new', 'New',),
        (10, 'active', 'Active',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.active,
    )

    num = models.IntegerField(
        blank=True,
        null=True,
    )

    KIND = Choices(
        (10, 'official', 'Official'),
        (20, 'practice', 'Practice'),
        (30, 'observer', 'Observer'),
    )

    kind = models.IntegerField(
        choices=KIND,
    )

    CATEGORY = Choices(
        (5, 'drcj', 'DRCJ'),
        (10, 'ca', 'CA'),
        (30, 'music', 'Music'),
        (40, 'performance', 'Performance'),
        (50, 'singing', 'Singing'),
    )

    category = models.IntegerField(
        choices=CATEGORY,
        null=True,
        blank=True,
    )

    psa = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )

    legacy_person = models.CharField(
        max_length=255,
        null=True,
        blank=True,
    )

    representing = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    # FKs
    round = models.ForeignKey(
        'Round',
        related_name='panelists',
        on_delete=models.CASCADE,
    )

    user = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        related_name='panelists',
        null=True,
        blank=True,
        on_delete=models.SET_NULL,
    )


    person_id = models.UUIDField(
        null=True,
        blank=True,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='panelists',
    )

    @cached_property
    def row_class(self):
        if self.category == self.CATEGORY.music:
            row_class = 'warning'
        elif self.category == self.CATEGORY.performance:
            row_class = 'success'
        elif self.category == self.CATEGORY.singing:
            row_class = 'info'
        else:
            row_class = None
        return row_class


    # Internals
    objects = PanelistManager()

    # class Meta:
        # unique_together = (
        #     ('round', 'num',),
        #     ('round', 'person', 'category',),
        # )

    class JSONAPIMeta:
        resource_name = "panelist"

    def __str__(self):
        return str(self.id)
        # return "{0} {1}".format(
        #     self.round,
        #     self.person,
        # )

    def clean(self):
        if self.kind > self.KIND.practice:
            raise ValidationError(
                {'category': 'Panel may only contain Official and Practice'}
            )
        if self.num and self.num > 50 and self.kind == self.KIND.official:
            raise ValidationError(
                {'num': 'Official Num must be less than 50'}
            )
        if self.num and self.num <= 50 and self.kind == self.KIND.practice:
            raise ValidationError(
                {'num': 'Practice Num must be greater than 50'}
            )
        if self.num and self.num and self.category == self.CATEGORY.ca:
            raise ValidationError(
                {'num': 'CAs must not have a num.'}
            )

    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return True

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.round.owners.filter(id__contains=request.user.id),
                self.round.status < self.round.STATUS.started,
            ]),
        ])

    def get_psa(self):
        Group = apps.get_model('bhs.group')
        Chart = apps.get_model('bhs.chart')
        Appearance = apps.get_model('rmanager.appearance')
        Score = apps.get_model('rmanager.score')
        # Score block
        group_ids = self.round.appearances.exclude(
            # Don't include advancers or MTs on PSA
            draw__gt=0,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications
            status=Appearance.STATUS.disqualified,
        ).values_list('group__id', flat=True)
        groups = Group.objects.filter(
            id__in=group_ids,
        ).prefetch_related(
            'appearances__songs__scores',
            'appearances__songs__scores__panelist',
            'appearances__round__session',
        ).annotate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.round.session,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.round.session,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.round.session,
                ),
            ),
        ).order_by(
            '-tot_points',
            '-sng_points',
            '-per_points',
        )
        # Monkeypatching
        class_map = {
            Panelist.CATEGORY.music: 'badge badge-warning mono-font',
            Panelist.CATEGORY.performance: 'badge badge-success mono-font',
            Panelist.CATEGORY.singing: 'badge badge-info mono-font',
        }
        for group in groups:
            appearances = group.appearances.filter(
                round__session=self.round.session,
            ).order_by('round__kind')
            for appearance in appearances:
                songs = appearance.songs.order_by(
                    'num',
                ).prefetch_related(
                    'scores',
                    'scores__panelist',
                ).annotate(
                    avg=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    Music=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    Performance=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    Singing=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart = song
                    scores2 = song.scores.select_related(
                        'panelist',
                    ).filter(
                        panelist__kind=Panelist.KIND.official,
                    ).order_by('points')
                    out = []
                    for score in scores2:
                        if score.points == 0:
                            score.points = "00"
                        span_class = class_map[score.panelist.category]
                        if score.panelist == self:
                            span_class = "{0} black-font".format(span_class)
                        out.append((score.points, span_class))
                    song.scores_patched = out
                    panelist_score = song.scores.get(
                        panelist=self,
                    )
                    category = self.get_category_display()
                    diff = panelist_score.points - getattr(song, category)
                    song.diff_patched = diff
                    pp = song.scores.get(panelist=self).points
                    song.pp = pp
                appearance.songs_patched = songs
            group.appearances_patched = appearances
        panelist = self
        person = Person.objects.get(id=panelist.person_id)
        panelist.person = person


        context = {
            'panelist': panelist,
            'groups': groups,
        }
        rendered = render_to_string(
            'reports/psa.html',
            context,
        )
        statelog = self.round.statelogs.latest('timestamp')
        footer = 'Published by {0} at {1}'.format(
            statelog.by,
            statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
        )
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def save_psa(self):
        content = self.get_psa()
        return self.psa.save('psa', content)

    def get_psa_email(self):
        Person = apps.get_model('bhs.person')
        context = {'panelist': self}

        template = 'emails/panelist_released.txt'
        person = Person.objects.get(id=self.person_id)
        subject = "[Barberscore] PSA for {0}".format(
            person.common_name,
        )
        to = ["{0} <{1}>".format(person.common_name, person.email)]
        cc = self.round.session.convention.get_ca_emails()

        if self.psa:
            pdf = self.psa.file
        else:
            pdf = self.get_psa()
        file_name = '{0} PSA.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]

        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email

    def send_psa_email(self):
        email = self.get_psa_email()
        return email.send()

    # Transitions
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.new, STATUS.active, STATUS.released],
        target=STATUS.released
    )
    def release(self, *args, **kwargs):
        # Saves PSA through post-transition signal to avoid race condition
        return


class Round(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
        (10, 'built', 'Built',),
        (20, 'started', 'Started',),
        (25, 'completed', 'Completed',),
        (27, 'verified', 'Verified',),
        (30, 'published', 'Published',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    KIND = Choices(
        (1, 'finals', 'Finals'),
        (2, 'semis', 'Semi-Finals'),
        (3, 'quarters', 'Quarter-Finals'),
    )

    kind = models.IntegerField(
        choices=KIND,
    )

    num = models.IntegerField(
        default=0,
    )

    spots = models.IntegerField(
        default=0,
    )

    date = models.DateField(
        null=True,
        blank=True,
    )

    footnotes = models.TextField(
        help_text="""
            Freeform text field; will print on OSS.""",
        blank=True,
    )

    oss = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )
    legacy_oss = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )
    sa = models.FileField(
        upload_to=FileUploadPath(),
        blank=True,
        default='',
    )

    is_reviewed = models.BooleanField(
        help_text="""Reviewed for history app""",
        default=False,
    )

    # FKs
    owners = models.ManyToManyField(
        settings.AUTH_USER_MODEL,
        related_name='rounds',
    )

    session = models.ForeignKey(
        'smanager.session',
        related_name='rounds',
        on_delete=models.CASCADE,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='rounds',
    )

    # Internals
    class Meta:
        # unique_together = (
        #     ('session', 'kind',),
        # )
        get_latest_by = [
            'num',
        ]

    class JSONAPIMeta:
        resource_name = "round"

    def __str__(self):
        return "{0} {1} {2}".format(
            self.session.convention,
            self.session.get_kind_display(),
            self.get_kind_display(),
        )

    # Methods
    def get_oss(self, zoom=1):
        Group = apps.get_model('bhs.group')
        Chart = apps.get_model('bhs.chart')
        Person = apps.get_model('bhs.person')
        Panelist = apps.get_model('rmanager.panelist')
        Appearance = apps.get_model('rmanager.appearance')
        Song = apps.get_model('rmanager.song')

        # Score Block - get completeds
        group_ids = self.appearances.filter(
            is_private=False,
        ).exclude(
            # Don't include advancers on OSS
            draw__gt=0,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            # Don't include mic testers on OSS
            num__lte=0,
        ).values_list('group__id', flat=True)
        groups = Group.objects.filter(
            id__in=group_ids,
        ).prefetch_related(
            'appearances',
            'appearances__songs__scores',
            'appearances__songs__scores__panelist',
            'appearances__round__session',
        ).annotate(
            max_round=Max(
                'appearances__round__num',
                filter=Q(
                    appearances__num__gt=0,
                    appearances__round__session=self.session,
                ),
            ),
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            mus_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            tot_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            mus_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            per_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            sng_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            tot_rank=Window(
                expression=RowNumber(),
                order_by=(
                    F('tot_points').desc(),
                    F('sng_points').desc(),
                    F('per_points').desc(),
                )
            ),
        ).order_by(
            '-tot_points',
            '-sng_points',
            '-per_points',
        )

        # Monkeypatching
        for group in groups:
            group.tot_rank = group.tot_rank + self.spots
            appearances = group.appearances.filter(
                num__gt=0,
                round__session=self.session,
                round__num__lte=self.num,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).order_by(
                'round__kind',
            ).annotate(
                tot_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                tot_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )
            for appearance in appearances:
                songs = appearance.songs.prefetch_related(
                    'scores',
                    'scores__panelist',
                ).order_by(
                    'num',
                ).annotate(
                    tot_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    mus_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    per_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    sng_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart = chart
                    penalties_map = {
                        10: "†",
                        30: "‡",
                        40: "✠",
                        50: "✶",
                    }
                    items = " ".join([penalties_map[x] for x in song.penalties])
                    song.penalties_patched = items
                appearance.songs_patched = songs
            recent = appearances.get(round__num=self.num)
            contesting = recent.contenders.order_by(
                'outcome__num',
            ).values_list(
                'outcome__num',
                flat=True
            )
            group.contesting_patched = ", ".join([str(x) for x in contesting])
            group.pos_patched = recent.pos
            group.representing_patched = recent.representing
            group.participants_patched = recent.participants
            group.appearances_patched = appearances


        # Penalties Block
        array = Song.objects.select_related(
            'appearance__round',
            'appearance__group',
        ).filter(
            appearance__round__session=self.session, # Using any session appearance
            penalties__len__gt=0, # Where there are penalties
            appearance__is_private=False, # If a public appearance
            appearance__group__id__in=group_ids,  # Only completeds
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            10: "† Score(s) penalized due to violation of Article IX.A.1 of the BHS Contest Rules.",
            30: "‡ Score(s) penalized due to violation of Article IX.A.2 of the BHS Contest Rules.",
            40: "✠ Score(s) penalized due to violation of Article IX.A.3 of the BHS Contest Rules.",
            50: "✶ Score(s) penalized due to violation of Article X.B of the BHS Contest Rules.",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        # Missing flag
        # use group_ids - these are the completeds
        is_missing = bool(Song.objects.filter(
            appearance__round__session=self.session,
            chart_id__isnull=True,
            appearance__group__id__in=group_ids,
        ))
        # Eval Only Block
        privates = self.appearances.prefetch_related(
            'group',
        ).filter(
            is_private=True,
        ).order_by(
            'group__name',
        ).values_list('group__name', flat=True)
        privates = list(privates)

        # Disqualification Block
        disqualifications = self.appearances.prefetch_related(
            'group',
        ).filter(
            status=Appearance.STATUS.disqualified,
        ).order_by(
            'group__name',
        ).values_list('group__name', flat=True)
        disqualifications = list(disqualifications)

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancers = self.appearances.filter(
                draw__gt=0,
            ).select_related(
                'group',
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'group__name',
            )
            advancers = list(advancers)
            try:
                mt = self.appearances.get(
                    draw=0,
                ).group.name
                advancers.append(('MT', mt))
            except self.appearances.model.DoesNotExist:
                pass
        else:
            advancers = None

        # Panelist Block
        panelists_raw = self.panelists.select_related(
            # 'person',
        ).filter(
            kind=Panelist.KIND.official,
            category__gte=Panelist.CATEGORY.ca,
        ).order_by(
            'num',
        )
        categories_map = {
            10: 'CA',
            30: 'MUS',
            40: 'PER',
            50: 'SNG',
        }
        panelists = []
        for key, value in categories_map.items():
            sections = panelists_raw.filter(
                category=key,
            ).select_related(
                # 'person',
            ).order_by(
                'num',
            )
            persons = []
            for x in sections:
                try:
                    person = Person.objects.get(id=x.person_id)
                    persons.append(
                        "{0} {1}".format(
                            person.common_name,
                            person.district,
                        )
                    )
                except AttributeError:
                    persons.append("(Unknown)")
            names = ", ".join(persons)
            panelists.append((value, names))


        # Outcome Block
        items = self.outcomes.select_related(
            'award',
        ).order_by(
            'num',
        ).values_list(
            'num',
            'award__name',
            'name',
        )
        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} {1}".format(item[0], item[1]),
                    item[2],
                )
            )

        context = {
            'round': self,
            'groups': groups,
            'penalties': penalties,
            'privates': privates,
            'disqualifications': disqualifications,
            'advancers': advancers,
            'panelists': panelists,
            'outcomes': outcomes,
            'is_missing': is_missing,
        }
        rendered = render_to_string('reports/oss.html', context)

        if self.session.convention.district == 'BHS':
            if self.session.convention.name == 'International Youth Convention':
                page_size = 'Legal'
            elif self.session.kind == self.session.KIND.quartet and self.kind == self.KIND.semis:
                page_size = 'Legal'
            else:
                page_size = 'Letter'
        else:
            if self.session.rounds.count() == 1:
                if groups.count() <= 15:
                    page_size = 'Letter'
                else:
                    page_size = 'Legal'
            else:
                if self.kind == self.KIND.finals:
                    if groups.count() >= 8:
                        page_size = 'Legal'
                    else:
                        page_size = 'Letter'
                elif self.kind == self.KIND.semis:
                    if groups.count() >= 8:
                        page_size = 'Legal'
                    else:
                        page_size = 'Letter'
                else:
                    page_size = 'Legal'
        try:
            statelog = self.statelogs.latest('timestamp')
            footer = 'Published by {0} at {1}'.format(
                statelog.by,
                statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
            )
        except StateLog.DoesNotExist:
            footer = '(Unknown)'
        file = pydf.generate_pdf(
            rendered,
            page_size=page_size,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
            zoom=zoom,
        )
        content = ContentFile(file)
        return content


    # Placeholder for single-page recursion tool.
    # from builtins import round as rnd
    # from PyPDF2 import PdfFileReader

    # r = Round.objects.get(id='0c02d601-357b-4315-9b73-063f563aa410')

    # f = r.get_oss()
    # obj = PdfFileReader(f)
    # num_pages = obj.getNumPages()

    # zoom = 1
    # while num_pages > 1:
    #     zoom = rnd((zoom - .05),2)
    #     print(zoom)
    #     f = r.get_oss(zoom=zoom)
    #     obj = PdfFileReader(f)
    #     num_pages = obj.getNumPages()
    # print(type(f))

    # r.oss.save('oss', f)


    def save_oss(self):
        oss = self.get_oss()
        return self.oss.save('oss', oss)

    def get_sa(self):
        Group = apps.get_model('bhs.group')
        Person = apps.get_model('bhs.person')
        Chart = apps.get_model('bhs.chart')
        Appearance = apps.get_model('rmanager.appearance')
        Panelist = apps.get_model('rmanager.panelist')
        Song = apps.get_model('rmanager.song')

        # Score Block
        group_ids = self.appearances.exclude(
            # DOn't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            # Don't include advancers on SA
            draw__gt=0,
        ).values_list('group__id', flat=True)
        groups = Group.objects.prefetch_related(
            'appearances',
            'appearances__songs__scores',
            'appearances__songs__scores__panelist',
            'appearances__round__session',
        ).filter(
            id__in=group_ids,
        ).annotate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            mus_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            tot_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            mus_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            per_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            sng_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            tot_rank=Window(
                expression=RowNumber(),
                order_by=(
                    F('tot_points').desc(),
                    F('sng_points').desc(),
                    F('per_points').desc(),
                )
            ),
            mus_rank=Window(
                expression=Rank(),
                order_by=F('mus_points').desc(),
            ),
            per_rank=Window(
                expression=Rank(),
                order_by=F('per_points').desc(),
            ),
            sng_rank=Window(
                expression=Rank(),
                order_by=F('sng_points').desc(),
            ),
        ).order_by(
            '-tot_points',
            '-sng_points',
            '-per_points',
        )

        persons = Person.objects.filter(
            panelists__round__session=self.session,
            panelists__category__gt=10,
            panelists__kind__in=[
                Panelist.KIND.official,
                Panelist.KIND.practice,
            ],
        ).order_by(
            'panelists__num',
        ).distinct()

        mus_persons_qs = persons.filter(
            panelists__category=Panelist.CATEGORY.music,
            panelists__round__session=self.session,
        ).order_by(
            'panelists__num',
        ).distinct()
        mus_persons = []
        for p in mus_persons_qs:
            practice = bool(p.panelists.get(round=self).kind == Panelist.KIND.practice)
            mus_persons.append((p.common_name, practice, p.initials))

        per_persons_qs = persons.filter(
            panelists__category=Panelist.CATEGORY.performance,
            panelists__round__session=self.session,
        ).order_by(
            'panelists__num',
        ).distinct()
        per_persons = []
        for p in per_persons_qs:
            practice = bool(p.panelists.get(round=self).kind == Panelist.KIND.practice)
            per_persons.append((p.common_name, practice, p.initials))

        sng_persons_qs = persons.filter(
            panelists__category=Panelist.CATEGORY.singing,
            panelists__round__session=self.session,
        ).order_by(
            'panelists__num',
        ).distinct()
        sng_persons = []
        for p in sng_persons_qs:
            practice = bool(p.panelists.get(round=self).kind == Panelist.KIND.practice)
            sng_persons.append((p.common_name, practice, p.initials))

        # Monkeypatching
        for group in groups:
            # Populate the group block
            appearances = group.appearances.filter(
                round__session=self.session,
                round__num__lte=self.num,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).annotate(
                tot_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                sng_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
                per_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                tot_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            ).order_by(
                'round__kind',
            )
            for appearance in appearances:
                songs = appearance.songs.prefetch_related(
                    'scores',
                    'scores__panelist',
                ).order_by(
                    'num',
                ).annotate(
                    tot_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    mus_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    per_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    sng_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                    tot_dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    mus_dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    per_dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    sng_dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    penalties_map = {
                        10: "†",
                        30: "‡",
                        40: "✠",
                        50: "✶",
                    }
                    items = " ".join([penalties_map[x] for x in song.penalties])
                    song.penalties_patched = items


                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart = chart

                    mus_scores = []
                    for person in mus_persons_qs:
                        raw_music_scores = song.scores.filter(
                            panelist__person=person,
                        ).order_by(
                            'panelist__num',
                        )
                        if raw_music_scores:
                            for m in raw_music_scores:
                                diff = abs(m.points - m.song.mus_score) > 5
                                practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                mus_scores.append((m.points, diff, practice))
                        else:
                            mus_scores.append((None, False, False))
                    song.mus_scores = mus_scores

                    per_scores = []
                    for person in per_persons_qs:
                        raw_performance_scores = song.scores.filter(
                            panelist__person=person,
                        ).order_by(
                            'panelist__num',
                        )
                        if raw_performance_scores:
                            for m in raw_performance_scores:
                                diff = abs(m.points - m.song.per_score) > 5
                                practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                per_scores.append((m.points, diff, practice))
                        else:
                            per_scores.append((None, False, False))
                    song.per_scores = per_scores

                    sng_scores = []
                    for person in sng_persons_qs:
                        raw_singing_scores = song.scores.filter(
                            panelist__person=person,
                        ).order_by(
                            'panelist__num',
                        )
                        if raw_singing_scores:
                            for m in raw_singing_scores:
                                diff = abs(m.points - m.song.sng_score) > 5
                                practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                sng_scores.append((m.points, diff, practice))
                        else:
                            sng_scores.append((None, False, False))
                    song.sng_scores = sng_scores


                appearance.songs_patched = songs
            group.appearances_patched = appearances

        # Build stats
        stats = Song.objects.select_related(
            'appearance__round',
        ).prefetch_related(
            'scores__panelist__kind',
        ).filter(
            appearance__round__session=self.session,
            appearance__round__num__lte=self.num,
        ).annotate(
            tot_dev=StdDev(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_dev=StdDev(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                    scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_dev=StdDev(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                    scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_dev=StdDev(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                    scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
        ).aggregate(
            tot=Avg('tot_dev'),
            mus=Avg('mus_dev'),
            per=Avg('per_dev'),
            sng=Avg('sng_dev'),
        )

        # Penalties Block
        array = Song.objects.select_related(
            'appearance__round',
            'appearance__group',
        ).filter(
            appearance__round__session=self.session, # Using any session appearance
            penalties__len__gt=0, # Where there are penalties
            appearance__is_private=False, # If a public appearance
            appearance__group__id__in=group_ids,  # Only completeds
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            10: "† Score(s) penalized due to violation of Article IX.A.1 of the BHS Contest Rules.",
            30: "‡ Score(s) penalized due to violation of Article IX.A.2 of the BHS Contest Rules.",
            40: "✠ Score(s) penalized due to violation of Article IX.A.3 of the BHS Contest Rules.",
            50: "✶ Score(s) penalized due to violation of Article X.B of the BHS Contest Rules.",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        context = {
            'round': self,
            'groups': groups,
            'mus_persons': mus_persons,
            'per_persons': per_persons,
            'sng_persons': sng_persons,
            'stats': stats,
            'penalties': penalties,
        }
        rendered = render_to_string('reports/sa.html', context)
        statelog = self.statelogs.latest('timestamp')
        footer = 'Published by {0} at {1}'.format(
            statelog.by,
            statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
        )
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Landscape',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def save_sa(self):
        sa = self.get_sa()
        return self.sa.save('sa', sa)

    def save_reports(self):
        oss = self.get_oss()
        self.oss.save('oss', oss, save=False)
        sa = self.get_sa()
        self.sa.save('sa', sa, save=False)
        return self.save()


    def get_legacy_oss(self):
        Chart = apps.get_model('bhs.chart')
        Person = apps.get_model('bhs.person')
        Panelist = apps.get_model('rmanager.panelist')
        Appearance = apps.get_model('rmanager.appearance')
        Song = apps.get_model('rmanager.song')
        Score = apps.get_model('rmanager.score')

        # Get the Groups
        group_ids = self.appearances.filter(
            is_private=False,
        ).exclude(
            # Don't include advancers on OSS
            draw__gt=0,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            # Don't include mic testers on OSS
            num__lte=0,
        ).values_list('group__id', flat=True)

        session = self.session

        rounds = session.rounds.filter(
            num__lte=self.num
        ).order_by('kind')

        # Monkeypatching
        for round in rounds:
            appearances = round.appearances.filter(
                group__id__in=group_ids,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).annotate(
                sub_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
            )
            for appearance in appearances:
                if appearance.round.num == 1:
                    appearance.prev_points = None
                    appearance.sub_points = None
                songs = appearance.songs.prefetch_related(
                    'scores',
                    'scores__panelist',
                ).order_by(
                    'num',
                ).annotate(
                    tot_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    mus_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    per_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    sng_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart = chart
                    penalties_map = {
                        10: "†",
                        30: "‡",
                        40: "✠",
                        50: "✶",
                    }
                    items = " ".join([penalties_map[x] for x in song.penalties])
                    song.penalties_patched = items
                contesting = appearance.contenders.order_by(
                    'outcome__num',
                ).values_list('outcome__num', flat=True)
                appearance.contesting = ", ".join([str(x) for x in contesting])
                tot_score = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num__lte=round.num,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(tot_score=Avg('points'))['tot_score']
                appearance.tot_score = tot_score
                tot_points = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num__lte=round.num,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(tot_points=Sum('points'))['tot_points']
                appearance.tot_points = tot_points if tot_points else 0
                prev_points = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num=round.num - 1,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(prev_points=Sum('points'))['prev_points']
                appearance.prev_points = prev_points
                appearance.songs_patched = songs
            appearances = sorted(appearances, reverse=True, key=lambda t: t.tot_points)
            round.appearances_patched = appearances

        # Penalties Block
        array = Song.objects.select_related(
            'appearance__round',
            'appearance__group',
        ).filter(
            appearance__round__session=self.session, # Using any session appearance
            penalties__len__gt=0, # Where there are penalties
            appearance__is_private=False, # If a public appearance
            appearance__group__id__in=group_ids,  # Only completeds
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            10: "† Score(s) penalized due to violation of Article IX.A.1 of the BHS Contest Rules.",
            30: "‡ Score(s) penalized due to violation of Article IX.A.2 of the BHS Contest Rules.",
            40: "✠ Score(s) penalized due to violation of Article IX.A.3 of the BHS Contest Rules.",
            50: "✶ Score(s) penalized due to violation of Article X.B of the BHS Contest Rules.",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        # Missing flag
        # use group_ids - these are the completeds
        is_missing = bool(Song.objects.filter(
            appearance__round__session=self.session,
            chart_id__isnull=True,
            appearance__group__id__in=group_ids,
        ))

        # Eval Only Block
        privates = self.appearances.prefetch_related(
            'group',
        ).filter(
            is_private=True,
        ).order_by(
            'group__name',
        ).values_list('group__name', flat=True)
        privates = list(privates)

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancers = self.appearances.filter(
                draw__gt=0,
            ).select_related(
                'group',
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'group__name',
            )
            advancers = list(advancers)
            try:
                mt = self.appearances.get(
                    draw=0,
                ).group.name
                advancers.append(('MT', mt))
            except self.appearances.model.DoesNotExist:
                pass
        else:
            advancers = None

        # Panelist Block
        panelists_raw = self.panelists.select_related(
            # 'person',
        ).filter(
            kind=Panelist.KIND.official,
            category__gte=Panelist.CATEGORY.ca,
        ).order_by(
            'num',
        )
        categories_map = {
            10: 'CA',
            30: 'Music',
            40: 'Performance',
            50: 'Singing',
        }
        panelists = {}
        for key, value in categories_map.items():
            sections = panelists_raw.filter(
                category=key,
            ).select_related(
                # 'person',
            ).order_by(
                'num',
            )
            persons = []
            for x in sections:
                person = Person.objects.get(id=x.person_id)
                persons.append("{0} - {1}".format(person.common_name, person.district))
            panelists[value] = persons


        # Outcome Block
        items = self.outcomes.select_related(
            'award',
        ).order_by(
            'num',
        ).values_list(
            'num',
            'award__name',
            'name',
        )
        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} is the winner of the {1}".format(item[2], item[1])
                )
            )
        contests = []
        for item in items:
            contests.append(
                "{0}={1}".format(item[0], item[1])
            )
        contest_legend = ", ".join(contests)

        context = {
            'round': self,
            'contest_legend': contest_legend,
            'rounds': rounds,
            'penalties': penalties,
            'privates': privates,
            'advancers': advancers,
            'panelists': panelists,
            'outcomes': outcomes,
            'is_missing': is_missing,
        }
        rendered = render_to_string('reports/legacy_oss.html', context)

        # if self.session.rounds.count() == 1:
        #     if groups.count() < 16:
        #         page_size = 'Letter'
        #     else:
        #         page_size = 'Legal'
        # else:
        #     if groups.count() > 8 and self.kind == self.KIND.finals:
        #         page_size = 'Legal'
        #     else:
        #         page_size = 'Letter'
        page_size = 'Letter'
        try:
            statelog = self.statelogs.latest('timestamp')
            footer = 'Published by {0} at {1}'.format(
                statelog.by,
                statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
            )
        except StateLog.DoesNotExist:
            footer = ""
        file = pydf.generate_pdf(
            rendered,
            page_size=page_size,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Courier',
            footer_font_size=6,
            # zoom=zoom,
        )
        content = ContentFile(file)
        return content

    def save_legacy_oss(self):
        content = self.get_legacy_oss()
        self.legacy_oss.save('legacy_oss', content)


    def get_titles(self):
        Song = apps.get_model('rmanager.song')
        Chart = apps.get_model('bhs.chart')
        appearances = self.appearances.filter(
            draw__gt=0,
        ).order_by(
            'draw',
        )
        for appearance in appearances:
            songs = Song.objects.filter(
                appearance__group=appearance.group,
                appearance__round__session=self.session,
            ).distinct().order_by(
                'appearance__round__num',
                'num',
            )
            titles = []
            for song in songs:
                chart = Chart.objects.get(id=song.chart_id)
                try:
                    title = chart.title
                except AttributeError:
                    title = "Unknown (Not in Repertory)"
                row = "{0} Song {1}: {2}".format(
                    song.appearance.round.get_kind_display(),
                    song.num,
                    title,
                )
                titles.append(row)
            appearance.titles_patched = titles

        context = {
            'appearances': appearances,
            'round': self,
        }
        rendered = render_to_string('reports/titles.html', context)
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
        )
        content = ContentFile(file)
        return content

    def get_announcements(self):
        Panelist = apps.get_model('rmanager.panelist')
        Group = apps.get_model('bhs.group')
        Appearance = apps.get_model('rmanager.appearance')
        appearances = self.appearances.filter(
            draw__gt=0,
        ).order_by(
            'draw',
        )
        group_ids = appearances.values_list('group__id', flat=True)
        mt = self.appearances.filter(
            draw=0,
        ).first()
        outcomes = self.outcomes.order_by(
            '-num',
        )
        if self.kind == self.KIND.finals:
            groups = Group.objects.filter(
                appearances__round__session=self.session,
            ).exclude(
                status__in=[
                    Appearance.STATUS.disqualified,
                    Appearance.STATUS.scratched,
                ],
            ).prefetch_related(
                'appearances__songs__scores',
                'appearances__songs__scores__panelist',
                'appearances__round__session',
            ).annotate(
                tot_points=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__round__session=self.session,
                        appearances__round__num__lte=self.num,
                    ),
                ),
                sng_points=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                        appearances__round__session=self.session,
                        appearances__round__num__lte=self.num,
                    ),
                ),
                per_points=Sum(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                        appearances__round__session=self.session,
                        appearances__round__num__lte=self.num,
                    ),
                ),
                raw_score=Avg(
                    'appearances__songs__scores__points',
                    filter=Q(
                        appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                        appearances__round__session=self.session,
                        appearances__round__num__lte=self.num,
                    ),
                ),
                tot_score=Func(
                    F('raw_score'),
                    function='ROUND',
                    template='%(function)s(%(expressions)s, 1)'
                ),
            ).order_by(
                'tot_points',
                'sng_points',
                'per_points',
            )
            groups = list(groups)[-5:]
        else:
            groups = None
        pos = self.appearances.aggregate(sum=Sum('pos'))['sum']
        context = {
            'round': self,
            'appearances': appearances,
            'mt': mt,
            'outcomes': outcomes,
            'groups': groups,
            'pos': pos,
        }

        rendered = render_to_string('reports/announcements.html', context)
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
        )
        content = ContentFile(file)

        document = Document()

        document.add_heading(
            'Announcements {0}'.format(self)
        )
        document.add_heading('Evaluations')
        document.add_paragraph(
            "Remember to state Start time(s) for each round, Rotation time and general location. For choruses it is critical to provide location for each chorus. 'It's on your CSA or eval schedule' is only appropriate for quartets."
        )
        document.add_heading('Administration')
        document.add_paragraph(
            "CSA: sent to group admins"
        )
        if pos:
            document.add_paragraph(
                "Total participants on stage: {0}".format(pos)
            )
        document.add_heading('Awards')
        for outcome in outcomes:
            document.add_paragraph("{0}: {1}".format(outcome.award.name, outcome.name))
        if appearances:
            document.add_heading('Draw')
            for appearance in appearances:
                document.add_paragraph(
                    "{0}: {1}".format(appearance.draw, appearance.group.name),
                    # style='List Bullet',
                )
            document.add_paragraph(
                "MT: {0}".format(mt.group.name),
                # style='List Bullet',
            )
        if groups:
            document.add_heading('Results')
            for group in groups:
                document.add_paragraph(
                    "With a score of {0}, a {1} average: {2}".format(
                        group.tot_points,
                        group.tot_score,
                        group.name,
                    ),
                    style='List Bullet',
                )

        buff = BytesIO()
        document.save(buff)
        content = ContentFile(buff.getvalue())
        return content


    def get_judge_emails(self):
        Panelist = apps.get_model('rmanager.panelist')
        judges = self.panelists.filter(
            # status=Panelist.STATUS.active,
            category__gt=Panelist.CATEGORY.ca,
            # person__email__isnull=False,
        ).order_by(
            'kind',
            'category',
            # 'person__last_name',
            # 'person__first_name',
        )
        seen = set()
        result = [
            "{0} ({1} {2}) <{3}>".format(judge.user.common_name, judge.get_kind_display(), judge.get_category_display(), judge.user.email,)
            for judge in judges
            if not (
                "{0} ({1} {2}) <{3}>".format(judge.user.common_name, judge.get_kind_display(), judge.get_category_display(), judge.user.email,) in seen or seen.add(
                    "{0} ({1} {2}) <{3}>".format(judge.user.common_name, judge.get_kind_display(), judge.get_category_display(), judge.user.email,)
                )
            )
        ]
        return result

    def mock(self):
        Appearance = apps.get_model('rmanager.appearance')
        if self.status != self.STATUS.started:
            raise RuntimeError("Round not Started")
        appearances = self.appearances.exclude(
            status=Appearance.STATUS.verified,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            status=Appearance.STATUS.scratched,
        )
        for appearance in appearances:
            appearance.mock()
            appearance.save()
        return


    def get_publish_email(self):
        Appearance = apps.get_model('rmanager.appearance')
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('rmanager.panelist')
        group_ids = self.appearances.filter(
            is_private=False,
        ).exclude(
            # Don't include advancers on OSS
            draw__gt=0,
        ).exclude(
            # Don't include mic testers on OSS
            num__lte=0,
        ).values_list('group__id', flat=True)
        completes = Group.objects.filter(
            id__in=group_ids,
        ).prefetch_related(
            'appearances',
            'appearances__songs__scores',
            'appearances__songs__scores__panelist',
            'appearances__round__session',
        ).annotate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            tot_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session=self.session,
                    appearances__round__num__lte=self.num,
                ),
            ),
            raw_rank=Window(
                expression=RowNumber(),
                order_by=(
                    F('tot_points').desc(),
                    F('sng_points').desc(),
                    F('per_points').desc(),
                )
            ),
        ).order_by(
            '-tot_points',
            '-sng_points',
            '-per_points',
        )
        # Monkeypatch results
        for complete in completes:
            complete.tot_rank = complete.raw_rank + self.spots

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancers = self.appearances.filter(
                status=Appearance.STATUS.advanced,
            ).select_related(
                'group',
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'group__name',
            )
            advancers = list(advancers)
            try:
                mt = self.appearances.get(
                    draw=0,
                ).group.name
                advancers.append(('MT', mt))
            except Appearance.DoesNotExist:
                pass
        else:
            advancers = None

        # Outcome Block
        items = self.outcomes.select_related(
            'award',
        ).order_by(
            'num',
        ).values_list(
            'num',
            'award__name',
            'name',
        )
        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} {1}".format(item[0], item[1]),
                    item[2],
                )
            )

        context = {
            'round': self,
            'advancers': advancers,
            'completes': completes,
            'outcomes': outcomes,
        }
        template = 'emails/round_publish.txt'
        subject = "[Barberscore] {0} Results and OSS".format(
            self,
        )
        to = self.session.convention.get_ca_emails()
        cc = self.session.convention.get_drcj_emails()
        cc.extend(self.get_judge_emails())
        bcc = self.session.get_participant_emails()

        if self.oss:
            pdf = self.oss.file
        else:
            pdf = self.get_oss()
        file_name = '{0} OSS.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]
        context['bcc'] = [x.partition(" <")[0] for x in bcc]

        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            bcc=bcc,
            attachments=attachments,
        )
        return email


    def send_publish_email(self):
        email = self.get_publish_email()
        return email.send()


    def get_publish_report_email(self):
        template = 'emails/round_publish_report.txt'
        context = {
            'round': self,
        }
        subject = "[Barberscore] {0} Reports and SA".format(
            self,
        )
        to = self.session.convention.get_ca_emails()
        cc = self.session.convention.get_drcj_emails()
        cc.extend(self.get_judge_emails())
        attachments = []
        if self.sa:
            pdf = self.sa.file
        else:
            pdf = self.get_sa()
        file_name = '{0} SA.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]
        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email


    def send_publish_report_email(self):
        email = self.get_publish_report_email()
        return email.send()

    # Round Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return True

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.round.owners.filter(id__contains=request.user.id),
                self.round.status != self.round.STATUS.published,
            ]),
        ])

    # Round Conditions
    def can_build(self):
        return True

    def can_complete(self):
        Appearance = apps.get_model('rmanager.appearance')
        return all([
            self.appearances.filter(
                status__in=[
                    Appearance.STATUS.verified,
                    Appearance.STATUS.disqualified,
                    Appearance.STATUS.scratched,
                ]
            ),
        ])

    def can_verify(self):
        return all([
            not self.outcomes.filter(
                name='MUST ENTER WINNER MANUALLY',
            ),
        ])

    def can_publish(self):
        return True

    # Round Transitions
    @fsm_log_by
    @transition(
        field=status,
        source='*',
        target=STATUS.new,
    )
    def reset(self, *args, **kwargs):
        self.oss.delete()
        self.sa.delete()
        panelists = self.panelists.all()
        appearances = self.appearances.all()
        outcomes = self.outcomes.all()
        panelists.delete()
        outcomes.delete()
        appearances.delete()
        return

    @fsm_log_by
    @transition(
        field=status,
        source=STATUS.new,
        target=STATUS.built,
        conditions=[can_build],
    )
    def build(self, *args, **kwargs):
        # Reset for indempodence
        self.reset()

        # Instantiate prior round
        if self.num == 1:
            prior_round = None
        else:
            prior_round = self.session.rounds.get(num=self.num - 1)

        # Create Panelsists
        Assignment = apps.get_model('cmanager.assignment')
        cas = self.session.convention.assignments.filter(
            status=Assignment.STATUS.active,
            kind__in=[
                Assignment.KIND.official,
                Assignment.KIND.practice,
            ],
            category=Assignment.CATEGORY.ca,
        ).order_by(
            'kind',
            'category',
            # 'person__last_name',
            # 'person__nick_name',
            # 'person__first_name',
        )
        for ca in cas:
            self.panelists.create(
                kind=ca.kind,
                category=ca.category,
                # person=ca.person,
            )
        officials = self.session.convention.assignments.filter(
            status=Assignment.STATUS.active,
            kind=Assignment.KIND.official,
            category__gt=Assignment.CATEGORY.ca,
        ).order_by(
            'kind',
            'category',
            # 'person__last_name',
            # 'person__nick_name',
            # 'person__first_name',
        )
        i = 0
        for official in officials:
            i += 1
            self.panelists.create(
                num=i,
                kind=official.kind,
                category=official.category,
                # person=official.person,
            )

        practices = self.session.convention.assignments.filter(
            status=Assignment.STATUS.active,
            kind=Assignment.KIND.practice,
            category__gt=Assignment.CATEGORY.ca,
        ).order_by(
            'kind',
            'category',
            # 'person__last_name',
            # 'person__nick_name',
            # 'person__first_name',
        )
        p = 50
        for practice in practices:
            p += 1
            self.panelists.create(
                num=p,
                kind=practice.kind,
                category=practice.category,
                # person=practice.person,
            )

        # Create Outcomes
        # Create from contests if no prior round
        if not prior_round:
            contests = self.session.contests.select_related(
                'award',
            ).filter(
                status__gt=0,
            ).annotate(
                cnt=Count(
                    'contestants',
                    filter=Q(
                        contestants__status=10,
                    )
                ),
            ).exclude(
                cnt=0,
            ).order_by(
                'award__tree_sort',
            )
            i = 0
            for contest in contests:
                i += 1
                self.outcomes.create(
                    num=i,
                    award=contest.award,
                )
        else:
            prior_outcomes = prior_round.outcomes.exclude(
                award__is_single=True,
            )
            for prior_outcome in prior_outcomes:
                new_outcome = self.outcomes.create(
                    num=prior_outcome.num,
                    award=prior_outcome.award,
                )

        # Create Appearances
        Appearance = apps.get_model('rmanager.appearance')
        Entry = apps.get_model('smanager.entry')
        # If the first round, populate from entries
        if not prior_round:
            entries = self.session.entries.filter(
                status=Entry.STATUS.approved,
            )
            z = 0
            for entry in entries:
                # TODO - MT hack
                if entry.is_mt:
                    entry.draw = z
                    z -= 1
                # Pull active contestants
                contestants = entry.contestants.filter(
                    status__gt=0,
                )
                # Set is_single=True if they are only in single-round contests
                is_single = not bool(
                    contestants.filter(
                        contest__award__is_single=False,
                    )
                )
                # Create and start group
                appearance = self.appearances.create(
                    entry=entry,
                    group=entry.group,
                    num=entry.draw,
                    is_single=is_single,
                    is_private=entry.is_private,
                    participants=entry.participants,
                    representing=entry.representing,
                )
                # Create contenders
                awards = contestants.values_list('contest__award', flat=True)
                outcomes = self.outcomes.filter(award__in=awards)
                for outcome in outcomes:
                    outcome.contenders.create(
                        appearance=appearance,
                    )
        # Otherwise, populate from prior round
        else:
            new_outcomes = self.outcomes.all()
            prior_appearances = prior_round.appearances.filter(
                status=Appearance.STATUS.advanced,
            )
            for prior_appearance in prior_appearances:
                # Create and start group
                appearance = self.appearances.create(
                    entry=prior_appearance.entry,
                    group=prior_appearance.group,
                    num=prior_appearance.draw,
                    is_single=prior_appearance.is_single,
                    is_private=prior_appearance.is_private,
                    participants=prior_appearance.participants,
                    representing=prior_appearance.representing,
                )
                for new_outcome in new_outcomes:
                    curry = bool(
                        prior_appearance.contenders.filter(outcome__award=new_outcome.award)
                    )
                    if curry:
                        new_outcome.contenders.create(
                            appearance=appearance,
                        )

            mts = prior_round.appearances.filter(
                draw__lte=0,
            )
            for mt in mts:
                appearance = self.appearances.create(
                    entry=mt.entry,
                    group=mt.group,
                    num=mt.draw,
                    is_single=mt.is_single,
                    is_private=True,
                    participants=mt.participants,
                    representing=mt.representing,
                )


    @fsm_log_by
    @transition(field=status, source=[STATUS.built], target=STATUS.started)
    def start(self, *args, **kwargs):
        # Build the appearances
        appearances = self.appearances.all()
        for appearance in appearances:
            appearance.build()
            appearance.save()
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.started, STATUS.completed,],
        target=STATUS.completed,
        conditions=[can_complete],)
    def complete(self, *args, **kwargs):
        Appearance = apps.get_model('rmanager.appearance')
        Panelist = apps.get_model('rmanager.panelist')
        # Run outcomes
        outcomes = self.outcomes.all()
        for outcome in outcomes:
            outcome.name = outcome.get_name()
            outcome.save()

        # If there is no next round simply return
        if self.kind == self.KIND.finals:
            return

        # Otherwise, figure out the Draw.
        # First, get spots available
        spots = self.spots

        # Get all multi appearances and annotate average.
        multis = self.appearances.filter(
            status=Appearance.STATUS.verified,
            is_single=False,
        ).annotate(
            avg=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                )
            ),
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                )
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                )
            ),
        )
        # If spots are constricted, find those who advance
        if spots:
            # All those above 73.0 advance automatically, regardless of spots available
            if self.session.convention.district == 'BHS':
                ordered = multis.order_by(
                    '-tot_points',
                    '-sng_points',
                    '-per_points',
                )
                advancer__ids = [x.id for x in ordered[:spots]]
                mt = ordered[spots:spots+1][0]
            else:
                automatics = multis.filter(
                    avg__gte=73.0,
                )
                # generate list of the advancers, as appearance IDs
                advancer__ids = [x.id for x in automatics]
                # Generate remaining multi appearances
                remains = multis.exclude(
                    id__in=advancer__ids,
                ).order_by(
                    '-tot_points',
                    '-sng_points',
                    '-per_points',
                )
                # Figure out remaining spots
                diff = spots - automatics.count()
                # If there are additional remaining spots, add them up to available
                if diff > 0:
                    adds = remains[:diff]
                    for a in adds:
                        advancer__ids.append(a.id)
                    try:
                        mt = remains[diff:diff+1][0]
                    except IndexError:
                        mt = None
                else:
                    # If MT available add, otherwise none
                    try:
                        mt = remains.first()
                    except AttributeError:
                        mt = None
        # Otherwise, advance all
        else:
            advancer__ids = [a.id for a in multis]
            mt = None

        # Reset draw
        self.appearances.update(draw=None)

        # Randomize the advancers and set the initial draw
        appearances = self.appearances.filter(
            id__in=advancer__ids,
        ).order_by('?')
        i = 1
        for appearance in appearances:
            appearance.draw = i
            appearance.save()
            i += 1
        # create Mic Tester at draw 0
        if mt:
            mt.draw = 0
            mt.save()
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.completed],
        target=STATUS.verified,
        conditions=[can_verify],)
    def verify(self, *args, **kwargs):
        Appearance = apps.get_model('rmanager.appearance')
        Panelist = apps.get_model('rmanager.panelist')
        completed_appearances = self.appearances.filter(
            status=Appearance.STATUS.verified,
        ).exclude(
            draw__gt=0,
        )
        for appearance in completed_appearances:
            appearance.complete()
            appearance.save()
        advancing_appearances = self.appearances.filter(
            status=Appearance.STATUS.verified,
            draw__gt=0,
        )
        for appearance in advancing_appearances:
            appearance.advance()
            appearance.save()
        panelists = self.panelists.filter(
            category__gt=Panelist.CATEGORY.ca,
        )
        for panelist in panelists:
            panelist.release()
            panelist.save()
        # Saves reports through transition signal to avoid race condition
        return

    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.verified],
        target=STATUS.published,
        conditions=[can_publish],)
    def publish(self, *args, **kwargs):
        """Publishes the results and notifies all parties"""
        Appearance = apps.get_model('rmanager.appearance')
        Panelist = apps.get_model('rmanager.panelist')
        # Send the OSS
        send_publish_email_from_round.delay(self)
        # Send the CSAs
        completed_appearances = self.appearances.filter(
            status=Appearance.STATUS.completed,
        )
        for appearance in completed_appearances:
            send_complete_email_from_appearance.delay(appearance)
        # Send the SAs
        send_publish_report_email_from_round.delay(self)
        # Send the PSAs
        panelists = self.panelists.filter(
            category__gt=Panelist.CATEGORY.ca,
            status=Panelist.STATUS.released,
        )
        for panelist in panelists:
            send_psa_email_from_panelist.delay(panelist)
        return


class Score(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
        (10, 'verified', 'Verified',),
        # (20, 'entered', 'Entered',),
        (25, 'cleared', 'Cleared',),
        (30, 'flagged', 'Flagged',),
        (35, 'revised', 'Revised',),
        (40, 'confirmed', 'Confirmed',),
        # (50, 'final', 'Final',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    # CATEGORY = Choices(
    #     (30, 'music', 'Music'),
    #     (40, 'performance', 'Performance'),
    #     (50, 'singing', 'Singing'),
    # )

    # category = models.IntegerField(
    #     choices=CATEGORY,
    # )

    # KIND = Choices(
    #     (10, 'official', 'Official'),
    #     (20, 'practice', 'Practice'),
    #     (30, 'composite', 'Composite'),
    # )

    # kind = models.IntegerField(
    #     choices=KIND,
    # )

    # num = models.IntegerField(
    # )

    points = models.IntegerField(
        help_text="""
            The number of points (0-100)""",
        null=True,
        blank=True,
        validators=[
            MaxValueValidator(
                100,
                message='Points must be between 0 - 100',
            ),
            MinValueValidator(
                0,
                message='Points must be between 0 - 100',
            ),
        ]
    )

    # FKs
    song = models.ForeignKey(
        'Song',
        related_name='scores',
        on_delete=models.CASCADE,
    )

    panelist = models.ForeignKey(
        'Panelist',
        related_name='scores',
        on_delete=models.CASCADE,
    )

    objects = ScoreManager()

    # class Meta:
    #     unique_together = (
    #         ('song', 'panelist',),
    #     )

    class JSONAPIMeta:
        resource_name = "score"

    def __str__(self):
        return str(self.pk)


    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return any([
            # Assigned owners can always see
            self.round.owners.filter(id__contains=request.user.id),

            # Panelists can always see their own scores
            self.panelist.user == request.user,

            # Panelists can see others' scores if Appearance is complete.
            all([
                self.round.panelists.filter(user__contains=request.user),
                self.song.appearance.status <= self.song.appearance.STATUS.completed
            ]),

            # Group onwers can see their own scores if complete.
            all([
                self.song.appearance.owners.filter(id__contains=request.user.id),
                self.song.appearance.status <= self.song.appearance.STATUS.completed
            ]),
        ])

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.song.appearance.round.owners.filter(id__contains=request.user.id),
                self.song.appearance.round.status < self.song.appearance.round.STATUS.verified,
            ]),
        ])


class Song(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
        (10, 'verified', 'Verified',),
        # (20, 'entered', 'Entered',),
        # (30, 'flagged', 'Flagged',),
        # (35, 'verified', 'Verified',),
        (38, 'finished', 'Finished',),
        (40, 'confirmed', 'Confirmed',),
        (50, 'final', 'Final',),
        (90, 'announced', 'Announced',),
        (95, 'archived', 'Archived',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
    )

    legacy_chart = models.CharField(
        max_length=255,
        null=True,
        blank=True,
    )

    asterisks = ArrayField(
        base_field=models.IntegerField(
        ),
        default=list,
        blank=True,
    )

    dixons = ArrayField(
        base_field=models.IntegerField(
        ),
        default=list,
        blank=True,
    )

    PENALTY = Choices(
        (10, 'patreg', 'Primarily Patriotic/Religious Intent',),
        (30, 'accompaniment', 'Instrumental Accompaniment',),
        (40, 'texture', 'Chorus Exceeding 4-Part Texture',),
        (50, 'enhancement', 'Sound Equipment or Electronic Enhancement',),
    )

    penalties = ArrayField(
        base_field=models.IntegerField(
            choices=PENALTY,
        ),
        default=list,
        blank=True,
    )

    stats = JSONField(
        null=True,
        blank=True,
    )

    # FKs
    appearance = models.ForeignKey(
        'Appearance',
        related_name='songs',
        on_delete=models.CASCADE,
    )

    chart_id = models.UUIDField(
        null=True,
        blank=True,
    )

    # Internals
    objects = SongManager()

    class Meta:
        # unique_together = (
        #     ('appearance', 'num',),
        # )
        get_latest_by = ['num']

    class JSONAPIMeta:
        resource_name = "song"

    def __str__(self):
        return str(self.id)

    # Methods
    def get_asterisks(self):
        """
        Check to see if the song produces a category variance (asterisk)

        Returns a list of categories that produced an asterisk.
        """
        # Set Flag
        asterisks = []
        # Get Averages by category
        categories = self.scores.filter(
            panelist__kind=10,
        ).values(
            'panelist__category',
        ).annotate(
            avg=Avg('points'),
        )
        for category in categories:
            category_scores = self.scores.filter(
                panelist__category=category['panelist__category'],
                panelist__kind=10,
            )
            for score in category_scores:
                is_asterisk = abs(score.points - category['avg']) > 5
                if is_asterisk:
                    asterisks.append(category['panelist__category'])
                    continue
        asterisks = list(set(asterisks))
        return asterisks

    def get_dixons(self):
        """
        Check to see if the song produces a spread error (Dixon's Q)

        Returns a list of categories that produced a Dixon's Q.
        """
        # Set flag
        output = []
        # Confidence thresholds
        confidence = {
            '3': 0.941,
            '6': .56,
            '9': .376,
            '12': .437,
            '15': .338,
        }
        # Only use official scores.
        scores = self.scores.filter(
            panelist__kind=10,
        )
        # Get the totals
        aggregates = scores.aggregate(
            cnt=Count('id'),
            max=Max('points'),
            min=Min('points'),
            spread=Max('points') - Min('points'),
        )
        # Check for validity.
        if aggregates['cnt'] < 3:
            return RuntimeError('Panel too small error')
        # Bypass to avoid division by zero
        if not aggregates['spread']:
            return output
        # Order the scores
        ascending = scores.order_by('points')
        descending = scores.order_by('-points')
        # Separate check for single-panel
        if aggregates['cnt'] == '3':
            if abs(ascending[0].points - ascending[1].points) >= 10:
                output.append(ascending[0].panelist.category)
            if abs(descending[0].points - descending[1].points) >= 10:
                output.append(descending[0].panelist.category)
            return output

        # Otherwise, run the checks, both ascending and descending
        critical = confidence[str(aggregates['cnt'])]
        ascending_distance = abs(ascending[0].points - ascending[1].points)
        ascending_q = ascending_distance / aggregates['spread']
        if ascending_q > critical and ascending_distance > 4:
            output.append(ascending[0].panelist.category)
        descending_distance = abs(descending[0].points - descending[1].points)
        descending_q = descending_distance / aggregates['spread']
        if descending_q > critical and descending_distance > 4:
            output.append(descending[0].panelist.category)
        return output


    # Song Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return any([
            self.appearance.round.status == self.appearance.round.STATUS.published,
            self.appearance.round.owners.filter(id__contains=request.user.id),
        ])

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return any([
            'SCJC' in request.user.roles,
            'CA' in request.user.roles,
        ])


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return any([
            all([
                self.appearance.round.owners.filter(id__contains=request.user.id),
                self.appearance.round.status < self.appearance.round.STATUS.verified,
            ]),
        ])
