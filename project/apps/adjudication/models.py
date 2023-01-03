# Standard Library
import datetime
import logging
import uuid
from random import randint
from io import BytesIO
from timezone_field import TimeZoneField
import json
# Third-Party
import pydf
import html
from docx import Document
from django_fsm import FSMIntegerField
from django_fsm import transition
from django_fsm_log.decorators import fsm_log_by
from django_fsm_log.models import StateLog
from dry_rest_permissions.generics import allow_staff_or_superuser
from dry_rest_permissions.generics import authenticated_users
from model_utils import Choices
from model_utils.models import TimeStampedModel
from django.db.models import Subquery, Sum, Max, Avg, StdDev, Count, Q, F, Func
from django.contrib.postgres.fields import ArrayField, JSONField
from django_fsm import RETURN_VALUE
from django.db.models.functions import DenseRank, RowNumber, Rank
from django.core.validators import MaxValueValidator
from django.core.validators import MinValueValidator
from django.db import models
from django.db.models import Min, Max, Avg
from django.db.models.functions import Cast
from django.db.models.constants import LOOKUP_SEP
from django.contrib.postgres.fields.jsonb import KeyTextTransform
from django.core.exceptions import ValidationError
from collections.abc import Iterable

# Django
from django.apps import apps
from django.contrib.contenttypes.fields import GenericRelation
from django.core.files.base import ContentFile
from django.db import models
from django.db.models import Window
from django.template.loader import render_to_string
from django.utils.functional import cached_property
from django.conf import settings
from django.utils.timezone import now
from django.contrib.postgres.fields import DecimalRangeField
from django.contrib.postgres.fields import IntegerRangeField
from phonenumber_field.modelfields import PhoneNumberField
from django.db.models.expressions import RawSQL

from .tasks import build_email
from .tasks import send_publish_email_from_round
from .tasks import send_publish_report_email_from_round
from .tasks import send_psa_email_from_panelist
from .tasks import save_psa_from_panelist
from .tasks import send_complete_email_from_appearance
from .tasks import save_reports_from_round

from .fields import UploadPath
from .fields import LowerEmailField
from .fields import DivisionsField


from .managers import AppearanceManager
from .managers import PanelistManager
from .managers import SongManager
from .managers import ScoreManager

from .helpers import round_up as rnd

log = logging.getLogger(__name__)

from apps.salesforce.decorators import notification_user

class KeyTextTransformFactory:

    def __init__(self, key_name):
        self.key_name = key_name

    def __call__(self, *args, **kwargs):
        return KeyTextTransform(self.key_name, *args, **kwargs)

class JSONF(F):

    def resolve_expression(self, query=None, allow_joins=True, reuse=None, summarize=False, for_save=False):
        rhs = super().resolve_expression(query, allow_joins, reuse, summarize, for_save)

        field_list = self.name.split(LOOKUP_SEP)
        for name in field_list[1:]:
            rhs = KeyTextTransformFactory(name)(rhs)
        return rhs

class Appearance(TimeStampedModel):
    """
    An appearance of a group on stage.

    The Appearance is meant to be a private resource.
    """

    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-30, 'disqualified', 'Disqualified',),
        (-20, 'scratched', 'Scratched',),
        (-10, 'completed', 'Completed',),
        (0, 'new', 'New',),
        (7, 'built', 'Built',),
        (10, 'started', 'Started',),
        (20, 'finished', 'Finished',),
        (25, 'variance', 'Variance',),
        (30, 'verified', 'Verified',),
        (40, 'advanced', 'Advanced',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
        help_text="""The order of appearance for this round.""",
        null=True,
        blank=True,
    )

    draw = models.IntegerField(
        help_text="""The draw for the next round.""",
        null=True,
        blank=True,
    )

    is_private = models.BooleanField(
        help_text="""Copied from entry.""",
        default=False,
    )

    is_single = models.BooleanField(
        help_text="""Single-round group""",
        default=False,
    )

    participants = models.CharField(
        help_text='Director(s) or Members (listed TLBB)',
        max_length=255,
        blank=True,
        default='',
    )

    area = models.CharField(
        help_text='Area representing',
        max_length=255,
        blank=True,
        default='',
    )

    onstage = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    actual_start = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    actual_finish = models.DateTimeField(
        help_text="""
            The actual appearance datetime.""",
        null=True,
        blank=True,
    )

    pos = models.IntegerField(
        help_text='Number of Singers on Stage',
        null=True,
        blank=True,
        verbose_name="Participants on Stage",
    )

    stats = JSONField(
        null=True,
        blank=True,
    )

    round_stats = JSONField(
        null=True,
        blank=True,
    )

    base = models.FloatField(
        help_text="""
            The incoming base score used to determine most-improved winners.""",
        null=True,
        blank=True,
    )

    variance_report = models.FileField(
        upload_to=UploadPath('variance_report'),
        blank=True,
        default='',
    )

    csa_report = models.FileField(
        upload_to=UploadPath('csa_report'),
        blank=True,
        default='',
        verbose_name='CSA Report',
    )

    # Denorm
    group_id = models.UUIDField(
        null=True,
        blank=True,
    )

    # Denorm
    entry_id = models.UUIDField(
        null=True,
        blank=True,
    )

    name = models.CharField(
        help_text="""
            The name of the resource.
        """,
        max_length=255,
        default='',
        blank=True,
    )

    KIND = Choices(
        (32, 'chorus', "Chorus"),
        (41, 'quartet', "Quartet"),
        (46, 'vlq', "VLQ"),
    )

    kind = models.IntegerField(
        help_text="""
            The kind of group.
        """,
        choices=KIND,
        null=True,
        blank=True,
    )

    GENDER = Choices(
        (10, 'male', "Male"),
        (20, 'female', "Female"),
        (30, 'mixed', "Mixed"),
    )

    gender = models.IntegerField(
        help_text="""
            The gender of group.
        """,
        choices=GENDER,
        null=True,
        blank=True,
    )

    DISTRICT = Choices(
        (110, 'bhs', 'BHS'),
        (200, 'car', 'CAR'),
        (205, 'csd', 'CSD'),
        (210, 'dix', 'DIX'),
        (215, 'evg', 'EVG'),
        (220, 'fwd', 'FWD'),
        (225, 'ill', 'ILL'),
        (230, 'jad', 'JAD'),
        (235, 'lol', 'LOL'),
        (240, 'mad', 'MAD'),
        (345, 'ned', 'NED'),
        (350, 'nsc', 'NSC'),
        (355, 'ont', 'ONT'),
        (360, 'pio', 'PIO'),
        (365, 'rmd', 'RMD'),
        (370, 'sld', 'SLD'),
        (375, 'sun', 'SUN'),
        (380, 'swd', 'SWD'),
    )

    district = models.IntegerField(
        choices=DISTRICT,
        null=True,
        blank=True,
    )

    DIVISION = Choices(
        ('EVG', [
            (10, 'evgd1', 'EVG Division I'),
            (20, 'evgd2', 'EVG Division II'),
            (30, 'evgd3', 'EVG Division III'),
            (40, 'evgd4', 'EVG Division IV'),
            (50, 'evgd5', 'EVG Division V'),
        ]),
        ('FWD', [
            (60, 'fwdaz', 'FWD Arizona'),
            (70, 'fwdne', 'FWD Northeast'),
            (80, 'fwdnw', 'FWD Northwest'),
            (90, 'fwdse', 'FWD Southeast'),
            (100, 'fwdsw', 'FWD Southwest'),
        ]),
        ('LOL', [
            (110, 'lol10l', 'LOL 10000 Lakes'),
            (120, 'lolone', 'LOL Division One'),
            (130, 'lolnp', 'LOL Northern Plains'),
            (140, 'lolpkr', 'LOL Packerland'),
            (150, 'lolsw', 'LOL Southwest'),
        ]),
        ('MAD', [
            # (160, 'madatl', 'MAD Atlantic'),
            (170, 'madcen', 'MAD Central'),
            (180, 'madnth', 'MAD Northern'),
            (190, 'madsth', 'MAD Southern'),
            # (200, 'madwst', 'MAD Western'),
        ]),
        ('NED', [
            (210, 'nedgp', 'NED Granite and Pine'),
            (220, 'nedmtn', 'NED Mountain'),
            (230, 'nedpat', 'NED Patriot'),
            (240, 'nedsun', 'NED Sunrise'),
            (250, 'nedyke', 'NED Yankee'),
        ]),
        ('SWD', [
            (260, 'swdne', 'SWD Northeast'),
            (270, 'swdnw', 'SWD Northwest'),
            (280, 'swdse', 'SWD Southeast'),
            (290, 'swdsw', 'SWD Southwest'),
        ]),
    )

    division = models.IntegerField(
        choices=DIVISION,
        null=True,
        blank=True,
    )

    bhs_id = models.IntegerField(
        blank=True,
        null=True,
    )

    code = models.CharField(
        help_text="""
            Short-form code.""",
        max_length=255,
        blank=True,
        null=True,
        default='',
    )

    image_id = models.CharField(
        help_text="""
            Image ID.""",
        max_length=255,
        blank=True,
        default='missing_image',
    )

    charts = ArrayField(
        base_field=models.TextField(
            blank=True,
        ),
        null=True,
        blank=True,
        default=list,
    )

    # Appearance FKs
    owners = models.ManyToManyField(
        settings.AUTH_USER_MODEL,
        blank=True,
        related_name='appearances',
    )

    round = models.ForeignKey(
        'Round',
        related_name='appearances',
        on_delete=models.CASCADE,
    )

    outcomes = models.ManyToManyField(
        'Outcome',
        related_name='appearances',
        blank=True,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='appearances',
    )

    # Appearance Internals
    objects = AppearanceManager()

    def clean(self):
        pass

    class Meta:
        ordering = [
            'num',
        ]
        # unique_together = (
        #     ('round', 'num'),
        # )

    class JSONAPIMeta:
        resource_name = "appearance"

    def __str__(self):
        return str(self.id)
        # return "{0} {1}".format(
        #     self.round,
        #     self.group,
        # )

    # Methods
    def get_owners_emails(self):
        Entry = apps.get_model('registration.entry')
        entry = Entry.objects.filter(
            group_id=self.group_id,
            session_id=self.round.session_id,
        ).first()
        if not len(entry.notification_list):
            raise ValueError("No notification list for {0}".format(entry))
        return ["{0}".format(x.strip()) for x in entry.notification_list.split(',')]

    def get_variance(self):
        Chart = apps.get_model('bhs.chart')
        Group = apps.get_model('bhs.group')
        Score = apps.get_model('adjudication.score')
        Panelist = apps.get_model('adjudication.panelist')

        # Group
        group = Group.objects.get(id=self.group_id)
        # Songs Block
        songs = self.songs.annotate(
            tot_score=Avg(
                'scores__points',
                filter=Q(
                    scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
        ).order_by('num')
        scores = Score.objects.filter(
            panelist__kind=Panelist.KIND.official,
            song__appearance=self,
        ).order_by(
            'category',
            # 'panelist__person__last_name',
            'song__num',
        )
        panelists = self.round.panelists.filter(
            kind=Panelist.KIND.official,
            category__gt=Panelist.CATEGORY.adm,
        ).order_by(
            'category',
            'last_name',
            # 'person__last_name',
        )
        variances = []
        for song in songs:
            if song.chart_id is not None:
                chart = Chart.objects.get(id=song.chart_id)
                song.chart_patched = chart
            variances.extend(song.dixons)
            variances.extend(song.asterisks)
        variances = list(set(variances))
        tot_points = scores.aggregate(sum=Sum('points'))['sum']
        context = {
            'appearance': self,
            'group': group,
            'songs': songs,
            'scores': scores,
            'panelists': panelists,
            'variances': variances,
            'tot_points': tot_points,
        }
        rendered = render_to_string('reports/variance.html', context)
        pdf = pydf.generate_pdf(
            rendered,
            enable_smart_shrinking=False,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
        )
        content = ContentFile(pdf)
        return content

    def save_variance(self):
        content = self.get_variance()
        self.variance_report.save("variance_report", content)

    def mock(self):
        # Mock Appearance
        Chart = apps.get_model('bhs.chart')
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('adjudication.panelist')
        group = Group.objects.get(id=self.group_id)
        if group.kind == group.KIND.chorus:
            # pos = group.members.filter(
            #     status=group.members.model.STATUS.active,
            # ).count()
            self.pos = 60
        # Try to approximate reality
        # prelim = getattr(getattr(self, "entry"), "prelim", None)
        # if not prelim:
        #     prelim = group.appearances.annotate(
        #         avg=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #             )
        #         )
        #     ).latest('round__date').avg or randint(60, 70)

        # ensures score doesn't exceed 100
        if self.base and self.base <= 100:
            prelim = self.base or randint(60, 80)
        else:
            prelim = randint(60, 80)
        songs = self.songs.all()
        for song in songs:
            charts = group.charts.order_by('id')
            x = song.num + self.round.num - 2
            try:
                chart = charts[x]
                song.chart_id = chart.id
                song.title = chart.title
                song.arrangers = chart.arrangers
            except IndexError:
                pass
            song.save()
            scores = song.scores.all()
            for score in scores:
                d = randint(-3, 3)
                score.points = prelim + d
                score.save()
        if self.status == self.STATUS.new:
            raise RuntimeError("Out of state")
        if self.status == self.STATUS.built:
            self.start()
            self.finish()
            self.verify()
            return
        if self.status == self.STATUS.started:
            self.finish()
            self.verify()
            return
        if self.status == self.STATUS.finished:
            self.verify()
            return

    def check_variance(self):
        # Set flag
        variance = False
        # Run checks for all songs and save.
        for song in self.songs.all():
            asterisks = song.get_asterisks()
            if asterisks:
                song.asterisks = asterisks
                variance = True
            dixons = song.get_dixons()
            if dixons:
                song.dixons = dixons
                variance = True
            if variance:
                song.save()
        return variance

    def get_stats(self):
        Round = apps.get_model('adjudication.round')
        rounds = Round.objects.filter(
            session_id=self.round.session_id,
            kind__gte=self.round.kind,
        )
        stats = rounds.aggregate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__group_id=self.group_id,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__group_id=self.group_id,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__group_id=self.group_id,
                ),
            ),
            mus_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__group_id=self.group_id,
                ),
            ),
            sng_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__group_id=self.group_id,
                ),
            ),
            per_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__group_id=self.group_id,
                ),
            ),
            mus_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__group_id=self.group_id,
                ),
            ),
            score_count=Count(
                'appearances__songs__scores',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__group_id=self.group_id,
                ),                
            ),
        )
        stats['tot_score'] = round((stats['tot_points'] / stats['score_count']), 1)
        stats.pop("score_count", None)
        for key, value in stats.items():
            stats[key] = rnd(value, 1)
        return stats

    def get_stats_round(self):
        Round = apps.get_model('adjudication.round')
        rounds = Round.objects.filter(
            id=self.round.id,
        )
        stats = rounds.aggregate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__group_id=self.group_id,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__group_id=self.group_id,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__group_id=self.group_id,
                ),
            ),
            mus_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__group_id=self.group_id,
                ),
            ),
            sng_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__group_id=self.group_id,
                ),
            ),
            per_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__group_id=self.group_id,
                ),
            ),
            mus_score=Avg(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
                    appearances__group_id=self.group_id,
                ),
            ),
            score_count=Count(
                'appearances__songs__scores',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__group_id=self.group_id,
                ),                
            ),
        )
        stats['tot_score'] = round((stats['tot_points'] / stats['score_count']), 1)
        stats.pop("score_count", None)
        for key, value in stats.items():
            stats[key] = rnd(value, 1)
        return stats

    def get_csa(self):
        Chart = apps.get_model('bhs.chart')
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('adjudication.panelist')
        Song = apps.get_model('adjudication.song')
        Score = apps.get_model('adjudication.score')

        # Appearancers Block
        group = Group.objects.get(id=self.group_id)
        stats = Score.objects.select_related(
            'song__appearance__group',
            'song__appearance__round__session',
            'song__appearance__round',
            'panelist',
        ).filter(
            song__appearance__group_id=self.group_id,
            song__appearance__round__session_id=self.round.session_id,
            panelist__kind=Panelist.KIND.official,
        ).aggregate(
            max=Max(
                'song__appearance__round__num',
            ),
            tot_points=Sum(
                'points',
            ),
            mus_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            tot_score=Avg(
                'points',
            ),
            mus_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
        )
        appearances = Appearance.objects.select_related(
            'round',
        ).prefetch_related(
            'songs__scores',
            'songs__scores__panelist',
        ).filter(
            group_id=self.group_id,
            round__session_id=self.round.session_id,
        ).order_by(
            '-round__num',
        ).annotate(
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
            tot_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
        )

        # Monkeypatch
        for key, value in stats.items():
            setattr(group, key, value)
        for appearance in appearances:
            songs = appearance.songs.prefetch_related(
                'scores',
                'scores__panelist',
            ).order_by(
                'num',
            ).annotate(
                tot_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
                tot_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )
            for song in songs:
                try:
                    chart = Chart.objects.get(id=song.chart_id)
                except Chart.DoesNotExist:
                    chart = None
                song.chart_patched = chart
                penalties_map = {
                    30: "†",
                    32: "‡",
                    34: "✠",
                    36: "✶",
                    38: "✢",
                    39: "✦",
                    40: "❉",
                    44: "∏",
                    48: "∇",
                    50: "※",
                }
                items = " ".join([penalties_map[x] for x in song.penalties])
                song.penalties_patched = items
            appearance.songs_patched = songs
        group.appearances_patched = appearances

        # Panelists
        panelists_ids = Panelist.objects.select_related(
        ).filter(
            kind=Panelist.KIND.official,
            round__session_id=self.round.session_id,
            category__gt=10,
        ).order_by('num').distinct('num')
        # Order Panelists by Category
        panelists = Panelist.objects.filter(id__in=panelists_ids).order_by('category', 'num')

        class_map = {
            Panelist.CATEGORY.music: 'warning',
            Panelist.CATEGORY.performance: 'success',
            Panelist.CATEGORY.singing: 'info',
        }

        # Score Block
        default_scoring_row = {}
        initials = []
        for panelist in panelists:
            default_scoring_row[str(panelist.num).zfill(2)] = {'points': '--', 'row_class': class_map[panelist.category]}

            initials.append(("{0}{1}".format(
                panelist.get_category_display()[0].upper(),
                str(panelist.num).zfill(2),
            ), str(panelist.num).zfill(2)))

        # Hackalicious
        category_count = {
            'Music': 0,
            'Performance': 0,
            'Singing': 0,
        }
        for panelist in panelists:
            category_count[panelist.get_category_display()] += 1
        songs = Song.objects.filter(
            appearance__round__session_id=self.round.session_id,
            appearance__group_id=self.group_id,
        ).order_by(
            'appearance__round__kind',
            'num',
        )
        for song in songs:
            try:
                chart = Chart.objects.get(id=song.chart_id)
            except Chart.DoesNotExist:
                chart = None

            song.chart_patched = chart
            scores = song.scores.filter(
                panelist__kind=Panelist.KIND.official,

            ).order_by('panelist__num')
            items = {}
            for score in scores:
                items[str(score.panelist.num).zfill(2)] = {'points': score.points, 'row_class': class_map[score.panelist.category]}
            song.scores_patched = {**default_scoring_row, **items}

        # Category Block
        categories = {
            'Music': [],
            'Performance': [],
            'Singing': [],
        }
        # panelists from above
        for panelist in panelists:
            item = categories[panelist.get_category_display()]
            name = "{0}{1} = {2}".format(
                panelist.get_category_display()[0].upper(),
                str(panelist.num).zfill(2),
                panelist.name
            )
            item.append(name)

        # Penalties Block
        array = Song.objects.filter(
            appearance__round__session_id=self.round.session_id,
            appearance__group_id=self.group_id,
            penalties__len__gt=0,
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            30: "† Scores penalized for Repeating Substantial Portions of a Song (V.A.2)",
            32: "‡ Scores penalized for Instrumental Accompaniment (IX.A.2.a)",
            34: "✠ Scores penalized for Chorus Exceeding 4-Part Texture (IX.A.2.b)",
            36: "✶ Scores penalized for Excessive Melody Not in Inner Part (IX.A.2.c)",
            38: "✢ Scores penalized for Lack of Characteristic Chord Progression (IX.A.2.d)",
            39: "✦ Scores penalized for Excessive Lyrics < 4 parts (IX.A.2.e)",
            40: "❉ Scores penalized for Primarily Patriotic/Religious Intent (IX.A.3)",
            44: "∏ Scores penalized for Not in Good Taste (IX.A.3.b)",
            48: "∇ Scores penalized for Non-members Performing on Stage (XI.A.1)",
            50: "※ Scores penalized for Sound Equipment or Electronic Enhancement (X.B.1-3)",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        context = {
            'appearance': self,
            'round': self.round,
            'group': group,
            'initials': initials,
            'songs': songs,
            'categories': categories,
            'penalties': penalties,
            'category_count': category_count,
        }
        rendered = render_to_string(
            'reports/csa.html',
            context,
        )
        statelog = self.round.statelogs.latest('timestamp')
        footer = 'Published by {0} at {1}'.format(
            statelog.by.name,
            statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
        )
        file = pydf.generate_pdf(
            rendered,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def save_csa(self):
        content = self.get_csa()
        return self.csa_report.save('csa', content)

    def get_complete_email(self):
        Panelist = apps.get_model('adjudication.panelist')
        Score = apps.get_model('adjudication.score')
        Chart = apps.get_model('bhs.chart')
        Group = apps.get_model('bhs.group')

        # Context
        group = Group.objects.get(id=self.group_id)

        stats = Score.objects.select_related(
            'song__appearance__group',
            'song__appearance__round',
            'panelist',
        ).filter(
            song__appearance__group_id=group.id,
            song__appearance__round__session_id=self.round.session_id,
            panelist__kind=Panelist.KIND.official,
        ).aggregate(
            max=Max(
                'song__appearance__round__num',
            ),
            tot_points=Sum(
                'points',
            ),
            mus_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_points=Sum(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            tot_score=Avg(
                'points',
            ),
            mus_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.music,
                )
            ),
            per_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.performance,
                )
            ),
            sng_score=Avg(
                'points',
                filter=Q(
                    panelist__category=Panelist.CATEGORY.singing,
                )
            ),
        )

        appearances = Appearance.objects.select_related(
            'round',
        ).prefetch_related(
            'songs__scores',
            'songs__scores__panelist',
        ).filter(
            group_id=group.id,
            round__session_id=self.round.session_id,
        ).order_by(
            '-round__num',
        ).annotate(
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
            tot_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                ),
            ),
            mus_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.music,
                ),
            ),
            per_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                ),
            ),
            sng_score=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                ),
            ),
        )

        # Monkeypatch
        for key, value in stats.items():
            setattr(group, key, value)
        for a in appearances:
            songs = a.songs.prefetch_related(
                'scores',
                'scores__panelist',
            ).order_by(
                'num',
            ).annotate(
                tot_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
                tot_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_points=Sum(
                    'scores__points',
                    filter=Q(
                        scores__panelist__kind=Panelist.KIND.official,
                        scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )
            for song in songs:
                chart = Chart.objects.get(id=song.chart_id)
                song.chart_patched = chart
                penalties_map = {
                    30: "†",
                    32: "‡",
                    34: "✠",
                    36: "✶",
                    38: "✢",
                    39: "✦",
                    40: "❉",
                    44: "∏",
                    48: "∇",
                    50: "※",
                }
                items = " ".join([penalties_map[x] for x in song.penalties])
                song.penalties_patched = items
            a.songs_patched = songs
        group.appearances_patched = appearances
        context = {'group': group}

        template = 'emails/appearance_complete.txt'
        subject = "[Barberscore] CSA for {0}".format(
            self.name,
        )
        to = self.get_owners_emails()
        cc = self.round.get_owners_emails()

        if self.csa_report:
            pdf = self.csa_report.file
        else:
            pdf = self.get_csa()
        file_name = '{0} CSA.pdf'.format(group.name)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]

        print("CSA EMAIL")
        print("to", to)
        print("cc", cc)

        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email

    def send_complete_email(self):
        if self.status != self.STATUS.completed:
            raise ValueError("Do not send CSAs unless Appearance is Completed")
        if self.round.status != self.round.STATUS.published:
            raise ValueError("Do not send CSAs unless Round is Published")
        email = self.get_complete_email()
        return email.send()

    def has_advanced(self):
        return True if self.draw is not None and self.draw > 0 else False

    # Appearance Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        if self.round.status == self.round.STATUS.published:
            return True
        return request.user in self.round.owners.all()

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'CA',
                'PC',
                'ADM',
            ]
        ))

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        if self.round.status == self.round.STATUS.published:
            return False
        return request.user in self.round.owners.all()

    # Appearance Conditions
    def can_verify(self):
        Group = apps.get_model('bhs.group')
        group = Group.objects.get(id=self.group_id)

        try:
            if group.kind == group.KIND.chorus and not self.pos:
                is_pos = False
            else:
                is_pos = True
        except AttributeError:
            is_pos = False
        return all([
            is_pos,
            not self.songs.filter(scores__points__isnull=True),
        ])

    # Appearance Transitions
    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.new],
        target=STATUS.built,
    )
    def build(self, *args, **kwargs):
        """Sets up the Appearance."""
        Panelist = apps.get_model('adjudication.panelist')
        panelists = self.round.panelists.filter(
            category__gt=Panelist.CATEGORY.adm,
        )
        i = 1
        while i <= 2:  # Number songs constant
            song = self.songs.create(
                num=i
            )
            for panelist in panelists:
                song.scores.create(
                    panelist=panelist,
                )
            i += 1
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.built],
        target=STATUS.started,
    )
    def start(self, *args, **kwargs):
        """Indicates when they start singing."""
        self.actual_start = now()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.started],
        target=STATUS.finished,
    )
    def finish(self, *args, **kwargs):
        """Indicates when they finish singing."""
        self.actual_finish = now()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.finished, STATUS.variance],
        target=RETURN_VALUE(STATUS.variance, STATUS.verified, STATUS.advanced,),
        conditions=[can_verify],
    )
    def verify(self, *args, **kwargs):
        # Checks for variance.  Returns Verified or Variance accordingly.
        if self.status == self.STATUS.finished:
            variance = self.check_variance()
            if variance:
                # Run variance report and save file.
                self.save_variance()
        # Variance is only checked once.
        else:
            variance = False

        self.stats = self.get_stats()
        self.round_stats = self.get_stats_round()

        if variance:
            return self.STATUS.variance
        elif self.has_advanced():
            # Update future Round stats
            future_appearances = Appearance.objects.filter(
                entry_id=self.entry_id,
                round__kind__lt=self.round.kind,
                status__in=[Appearance.STATUS.verified, Appearance.STATUS.advanced]
            )

            for a in future_appearances:
                a.stats = a.get_stats()
                a.save()

            return self.STATUS.advanced
        else:
            return self.STATUS.verified

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.verified, STATUS.advanced],
        target=STATUS.advanced,
    )
    def advance(self, *args, **kwargs):
        # Advances the Group.
        # Delete if CSA report has been created.
        if self.csa_report:
            self.csa_report.delete()
            self.save()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.verified, STATUS.advanced, STATUS.completed],
        target=STATUS.completed,
    )
    def complete(self, *args, **kwargs):
        # Completes the Group.
        # Saves CSA via post-transition signal to avoid race condition
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[
            STATUS.new,
            STATUS.built,
        ],
        target=STATUS.scratched,
    )
    def scratch(self, *args, **kwargs):
        if self.round.num > 1:
            self.stats = self.get_stats()
        # Scratches the group.
        # Remove songs
        songs = self.songs.all()
        songs.delete()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source='*',
        target=STATUS.disqualified,
    )
    def disqualify(self, *args, **kwargs):
        # Disqualify the group.
            # self.songs.delete()
        # Notify the group?
        return


class Outcome(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-10, 'inactive', 'Inactive',),
        (0, 'new', 'New',),
        (10, 'active', 'Active',),
    )

    status = models.IntegerField(
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
        blank=True,
        null=True,
    )

    winner = models.CharField(
        max_length=1024,
        null=True,
        blank=True,
    )

    # Denormalized from BHS Award
    award_id = models.UUIDField(
        null=True,
        blank=True,
    )

    name = models.CharField(
        help_text="""Award Name.""",
        max_length=255,
        null=True,
        blank=True,
    )

    KIND = Choices(
        (32, 'chorus', "Chorus"),
        (41, 'quartet', "Quartet"),
    )

    kind = models.IntegerField(
        choices=KIND,
        null=True,
        blank=True,
    )

    GENDER = Choices(
        (10, 'male', "Male"),
        (20, 'female', "Female"),
        (30, 'mixed', "Mixed"),
    )

    gender = models.IntegerField(
        help_text="""
            The gender to which the award is restricted.  If unselected, this award is open to all combinations.
        """,
        choices=GENDER,
        null=True,
        blank=True,
    )

    LEVEL = Choices(
        (10, 'championship', "Championship"),
        (30, 'qualifier', "Qualifier"),
        (45, 'representative', "Representative"),
        (50, 'deferred', "Deferred"),
        (60, 'manual', "Manual"),
        (70, 'raw', "Improved - Raw"),
        (80, 'standard', "Improved - Standard"),
    )

    level = models.IntegerField(
        choices=LEVEL,
        null=True,
        blank=True,
    )

    SEASON = Choices(
        (1, 'summer', 'Summer',),
        (2, 'midwinter', 'Midwinter',),
        (3, 'fall', 'Fall',),
        (4, 'spring', 'Spring',),
    )

    season = models.IntegerField(
        choices=SEASON,
        null=True,
        blank=True,
    )

    description = models.TextField(
        help_text="""
            The Public description of the award.""",
        max_length=1000,
        null=True,
        blank=True,
    )

    district = models.CharField(
        max_length=255,
        null=True,
        blank=True,
    )

    DIVISION = Choices(
        (10, 'evgd1', 'EVG Division I'),
        (20, 'evgd2', 'EVG Division II'),
        (30, 'evgd3', 'EVG Division III'),
        (40, 'evgd4', 'EVG Division IV'),
        (50, 'evgd5', 'EVG Division V'),
        (60, 'fwdaz', 'FWD Arizona'),
        (70, 'fwdne', 'FWD Northeast'),
        (80, 'fwdnw', 'FWD Northwest'),
        (90, 'fwdse', 'FWD Southeast'),
        (100, 'fwdsw', 'FWD Southwest'),
        (110, 'lol10l', 'LOL 10000 Lakes'),
        (120, 'lolone', 'LOL Division One'),
        (130, 'lolnp', 'LOL Northern Plains'),
        (140, 'lolpkr', 'LOL Packerland'),
        (150, 'lolsw', 'LOL Southwest'),
        # (160, 'madatl', 'MAD Atlantic'),
        (170, 'madcen', 'MAD Central'),
        (180, 'madnth', 'MAD Northern'),
        (190, 'madsth', 'MAD Southern'),
        # (200, 'madwst', 'MAD Western'),
        (210, 'nedgp', 'NED Granite and Pine'),
        (220, 'nedmtn', 'NED Mountain'),
        (230, 'nedpat', 'NED Patriot'),
        (240, 'nedsun', 'NED Sunrise'),
        (250, 'nedyke', 'NED Yankee'),
        (260, 'swdne', 'SWD Northeast'),
        (270, 'swdnw', 'SWD Northwest'),
        (280, 'swdse', 'SWD Southeast'),
        (290, 'swdsw', 'SWD Southwest'),
    )

    division = models.IntegerField(
        choices=DIVISION,
        null=True,
        blank=True,
    )

    AGE = Choices(
        (10, 'seniors', 'Seniors',),
        (20, 'novice', 'Novice',),
        (30, 'youth', 'Youth',),
    )

    age = models.IntegerField(
        choices=AGE,
        null=True,
        blank=True,
    )

    is_novice = models.BooleanField(
        default=False,
        null=True,
        blank=True,
    )

    is_single = models.BooleanField(
        default=False,
        null=True,
        blank=True,
    )

    SIZE = Choices(
        (100, 'p1', 'Plateau 1',),
        (110, 'p2', 'Plateau 2',),
        (120, 'p3', 'Plateau 3',),
        (130, 'p4', 'Plateau 4',),
        (140, 'pa', 'Plateau A',),
        (150, 'paa', 'Plateau AA',),
        (160, 'paaa', 'Plateau AAA',),
        (170, 'paaaa', 'Plateau AAAA',),
        (180, 'pb', 'Plateau B',),
        (190, 'pi', 'Plateau I',),
        (200, 'pii', 'Plateau II',),
        (210, 'piii', 'Plateau III',),
        (220, 'piv', 'Plateau IV',),
        (230, 'small', 'Small',),
    )

    size = models.IntegerField(
        choices=SIZE,
        null=True,
        blank=True,
    )

    size_range = IntegerRangeField(
        null=True,
        blank=True,
    )

    SCOPE = Choices(
        (100, 'p1', 'Plateau 1',),
        (110, 'p2', 'Plateau 2',),
        (120, 'p3', 'Plateau 3',),
        (130, 'p4', 'Plateau 4',),
        (140, 'pa', 'Plateau A',),
        (150, 'paa', 'Plateau AA',),
        (160, 'paaa', 'Plateau AAA',),
        (170, 'paaaa', 'Plateau AAAA',),
        (175, 'paaaaa', 'Plateau AAAAA',),
    )

    scope = models.IntegerField(
        choices=SCOPE,
        null=True,
        blank=True,
    )

    scope_range = DecimalRangeField(
        null=True,
        blank=True,
    )

    tree_sort = models.IntegerField(
        # unique=True,
        # editable=False,
        null=True,
        blank=True,
    )

    printed = models.BooleanField(
        help_text="""
            Show this outcome on the OSS.""",
        default=True,
    )

    print_on_finals_oss = models.BooleanField(
        help_text="""
            Show this outcome on the Finals OSS.""",
        default=False,
        verbose_name='Include on Finals OSS',
    )

    # FKs
    round = models.ForeignKey(
        'Round',
        related_name='outcomes',
        on_delete=models.CASCADE,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='outcomes',
    )

    # Methods
    def get_winner(self):
        Award = apps.get_model('bhs.award')
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('adjudication.panelist')
        award = Award.objects.get(id=self.award_id)
        if self.round.kind != self.round.KIND.finals and not award.is_single:
            return "(Result determined in Finals)"
        if award.level == award.LEVEL.deferred:
            return "(Result determined post-contest)"
        if award.level in [award.LEVEL.manual, award.LEVEL.raw, award.LEVEL.standard]:
            return "MUST ENTER WINNER MANUALLY"
        # if award.level == award.LEVEL.raw:
        #     group_ids = Group.objects.filter(
        #         appearances__contenders__outcome=self,
        #     ).values_list('id', flat=True)
        #     group = Group.objects.filter(
        #         id__in=group_ids,
        #     ).annotate(
        #         avg=Avg(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         base=Avg(
        #             'appearances__base',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #             ),
        #         ),
        #         diff=F('avg') - F('base'),
        #         sng=Sum(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #                 appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing
        #             ),
        #         ),
        #         per=Sum(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #                 appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance
        #             ),
        #         ),
        #     ).filter(
        #         diff__gt=0,
        #     ).order_by(
        #        '-diff',
        #     ).first()
        #     try:
        #         winner = group.name
        #     except AttributeError:
        #         winner = "(No Award Winner)"
        #     return winner
        # if award.level == award.LEVEL.standard:
        #     group_ids = Group.objects.filter(
        #         appearances__contenders__outcome=self,
        #     ).values_list('id', flat=True)
        #     groups = Group.objects.filter(
        #         id__in=group_ids,
        #     ).annotate(
        #         score=Avg(
        #             'appearances__songs__scores__points',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #                 appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         base=Avg(
        #             'appearances__base',
        #             filter=Q(
        #                 appearances__round__session=self.round.session,
        #             ),
        #         ),
        #     )
        #     totals = groups.aggregate(
        #         Avg('score'),
        #         Avg('base'),
        #     )
        #     winner = groups.annotate(
        #         score_ratio=F('score') / totals['score__avg'],
        #         base_ratio=F('base') / totals['base__avg'],
        #         diff=F('score_ratio') - F('base_ratio'),
        #     ).order_by(
        #         '-diff',
        #     ).first().name
        if award.level == award.LEVEL.qualifier:
            threshold = award.threshold
            winners = self.appearances.filter(
                stats__isnull = False,
                stats__tot_score__gte=threshold,
            ).order_by(
                'name',
            ).values_list('name', flat=True)

            # group_ids = self.contenders.filter(
            #     status__gt=0,
            # ).values_list(
            #     'appearance__group_id',
            #     flat=True,
            # )
            # qualifiers = Group.objects.filter(
            #     id__in=group_ids,
            # ).annotate(
            #     avg=Avg(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__round__session=self.round.session,
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #         ),
            #     ),
            # ).filter(
            #     avg__gte=threshold,
            # ).order_by(
            #     'name',
            # ).values_list('name', flat=True)
            # group_ids = self.contenders.filter(
            #     appearance__stats__tot_score__gte=threshold,
            # ).values_list(
            #     'appearance__group_id',
            #     flat=True,
            # )
            qualifiers = list(winners)
            if qualifiers:
                return ", ".join(qualifiers)
            return "(No Qualifiers)"
        if award.level in [award.LEVEL.championship, award.LEVEL.representative]:
            winner = self.appearances.filter(stats__isnull = False).order_by(
                'stats__tot_points',
                'stats__sng_points',
                'stats__per_points',
            ).last()
            # winner = Group.objects.get(id=group_id)
            # winner = Group.objects.filter(
            #     appearances__contenders__outcome=self,
            # ).annotate(
            #     tot=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #         ),
            #     ),
            #     sng=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing
            #         ),
            #     ),
            #     per=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance
            #         ),
            #     ),
            # ).earliest(
            #     '-tot',
            #     '-sng',
            #     '-per',
            # )
            if winner:
                return winner.name
            return "(No Recipient)"
        raise RuntimeError("Level mismatch")

    # Internals
    # class Meta:
    #     unique_together = (
    #         ('round', 'award',),
    #         ('round', 'num',),
    #     )

    class JSONAPIMeta:
        resource_name = "outcome"

    def __str__(self):
        return self.name or str(self.id)

    def clean(self):
        pass

    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        if self.round.status == self.round.STATUS.published:
            return True
        return request.user in self.round.owners.all()

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'PC',
                'CA',
                'ADM',
            ]
        ))


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        if self.round.status == self.round.STATUS.published:
            return False
        return request.user in self.round.owners.all()


class Panelist(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (-10, 'inactive', 'Inactive',),
        (-5, 'released', 'Released',),
        (0, 'new', 'New',),
        (10, 'active', 'Active',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.active,
    )

    num = models.IntegerField(
        blank=True,
        null=True,
    )

    KIND = Choices(
        (10, 'official', 'Official'),
        (20, 'practice', 'Practice'),
        (30, 'observer', 'Observer'),
    )

    kind = models.IntegerField(
        choices=KIND,
    )

    CATEGORY_SHORT_NAMES = {
        9: "PC",
        10: "ADM",
        30: "MUS",
        40: "PER",
        50: "SNG",
    }

    CATEGORY = Choices(
        (5, 'drcj', 'DRCJ'),
        (9, 'pc', 'PC'),
        (10, 'adm', 'ADM'),
        (30, 'music', 'Music'),
        (40, 'performance', 'Performance'),
        (50, 'singing', 'Singing'),
    )

    category = models.IntegerField(
        choices=CATEGORY,
        null=True,
        blank=True,
    )

    psa_report = models.FileField(
        upload_to=UploadPath('psa_report'),
        blank=True,
        default='',
    )

    person_id = models.UUIDField(
        null=True,
        blank=True,
    )

    name = models.CharField(
        help_text="""
            The prefix of the person.""",
        max_length=255,
        blank=True,
        default='',
    )

    first_name = models.CharField(
        help_text="""
            The first name of the person.""",
        max_length=255,
        blank=True,
        default='',
    )

    last_name = models.CharField(
        help_text="""
            The last name of the person.""",
        max_length=255,
        blank=True,
        default='',
    )

    DISTRICT = Choices(
        (110, 'bhs', 'BHS'),
        (200, 'car', 'CAR'),
        (205, 'csd', 'CSD'),
        (210, 'dix', 'DIX'),
        (215, 'evg', 'EVG'),
        (220, 'fwd', 'FWD'),
        (225, 'ill', 'ILL'),
        (230, 'jad', 'JAD'),
        (235, 'lol', 'LOL'),
        (240, 'mad', 'MAD'),
        (345, 'ned', 'NED'),
        (350, 'nsc', 'NSC'),
        (355, 'ont', 'ONT'),
        (360, 'pio', 'PIO'),
        (365, 'rmd', 'RMD'),
        (370, 'sld', 'SLD'),
        (375, 'sun', 'SUN'),
        (380, 'swd', 'SWD'),
    )

    district = models.IntegerField(
        choices=DISTRICT,
        null=True,
        blank=True,
    )

    area = models.CharField(
        help_text="""
            District""",
        max_length=10,
        blank=True,
        default='',
    )

    email = LowerEmailField(
        help_text="""
            The contact email of the resource.""",
        blank=True,
        null=True,
    )

    cell_phone = PhoneNumberField(
        help_text="""
            The cell phone number of the resource.  Include country code.""",
        blank=True,
        null=True,
    )

    airports = ArrayField(
        base_field=models.CharField(
            blank=True,
            max_length=3,
        ),
        null=True,
        blank=True,
        default=list,
    )

    image = models.ImageField(
        upload_to=UploadPath('image'),
        null=True,
        blank=True,
    )

    bhs_id = models.IntegerField(
        null=True,
        blank=True,
    )

    # FKs
    round = models.ForeignKey(
        'Round',
        related_name='panelists',
        on_delete=models.CASCADE,
    )

    owners = models.ManyToManyField(
        settings.AUTH_USER_MODEL,
        related_name='panelists',
        blank=True,
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='panelists',
    )

    # Internals
    objects = PanelistManager()

    # class Meta:
        # unique_together = (
        #     ('round', 'num',),
        #     ('round', 'person', 'category',),
        # )

    class JSONAPIMeta:
        resource_name = "panelist"

    def __str__(self):
        return str(self.id)
        # return "{0} {1}".format(
        #     self.round,
        #     self.person,
        # )

    def clean(self):
        if self.kind > self.KIND.practice:
            raise ValidationError(
                {'category': 'Panel may only contain Official and Practice'}
            )
        if self.num and self.num > 50 and self.kind == self.KIND.official:
            raise ValidationError(
                {'num': 'Official Num must be less than 50'}
            )
        if self.num and self.num <= 50 and self.kind == self.KIND.practice:
            raise ValidationError(
                {'num': 'Practice Num must be greater than 50'}
            )
        if self.num and self.num and self.category == self.CATEGORY.pc:
            raise ValidationError(
                {'num': 'PCs must not have a num.'}
            )
        if self.num and self.num and self.category == self.CATEGORY.adm:
            raise ValidationError(
                {'num': 'ADMs must not have a num.'}
            )

    # Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return True

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'PC',
                'CA',
                'ADM',
            ]
        ))

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'PC',
                'CA',
                'ADM',
            ]
        ))
        # if self.round.status >= self.round.STATUS.started:
        #     return request.user in self.round.owners.all()


    def get_psa(self):
        Group = apps.get_model('bhs.group')
        Chart = apps.get_model('bhs.chart')
        Appearance = apps.get_model('adjudication.appearance')
        Score = apps.get_model('adjudication.score')
        # Score block
        group_ids = self.round.appearances.exclude(
            # Don't include advancers or MTs on PSA
            draw__gt=0,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications
            status=Appearance.STATUS.disqualified,
        ).values_list('group_id', flat=True)
        groups = Group.objects.filter(
            id__in=group_ids,
        ).prefetch_related(
            'appearances__songs__scores',
            'appearances__songs__scores__panelist',
        ).annotate(
            tot_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__round__session_id=self.round.session_id,
                ),
            ),
            per_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    appearances__round__session_id=self.round.session_id,
                ),
            ),
            sng_points=Sum(
                'appearances__songs__scores__points',
                filter=Q(
                    appearances__songs__scores__panelist__kind=Panelist.KIND.official,
                    appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    appearances__round__session_id=self.round.session_id,
                ),
            ),
        ).order_by(
            '-tot_points',
            '-sng_points',
            '-per_points',
        )
        # Monkeypatching
        class_map = {
            Panelist.CATEGORY.music: 'badge badge-warning mono-font',
            Panelist.CATEGORY.performance: 'badge badge-success mono-font',
            Panelist.CATEGORY.singing: 'badge badge-info mono-font',
        }
        for group in groups:
            appearances = group.appearances.filter(
                round__session_id=self.round.session_id,
            ).order_by('round__kind')
            for appearance in appearances:
                songs = appearance.songs.order_by(
                    'num',
                ).prefetch_related(
                    'scores',
                    'scores__panelist',
                ).annotate(
                    avg=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    dev=StdDev(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    Music=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    Performance=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    Singing=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart_patched = chart
                    scores2 = song.scores.select_related(
                        'panelist',
                    ).filter(
                        panelist__kind=Panelist.KIND.official,
                    ).order_by('points')
                    out = []
                    for score in scores2:
                        if score.points == 0:
                            score.points = "00"
                        span_class = class_map[score.panelist.category]
                        if score.panelist == self:
                            span_class = "{0} black-font".format(span_class)
                        out.append((score.points, span_class))
                    song.scores_patched = out
                    panelist_score = song.scores.get(
                        panelist=self,
                    )
                    category = self.get_category_display()
                    diff = panelist_score.points - getattr(song, category)
                    song.diff_patched = diff
                    pp = song.scores.get(panelist=self).points
                    song.pp = pp
                appearance.songs_patched = songs
            group.appearances_patched = appearances
        panelist = self

        context = {
            'panelist': panelist,
            'groups': groups,
        }
        rendered = render_to_string(
            'reports/psa.html',
            context,
        )
        statelog = self.round.statelogs.latest('timestamp')
        footer = 'Published by {0} at {1}'.format(
            statelog.by.name,
            statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
        )
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def save_psa(self):
        content = self.get_psa()
        return self.psa_report.save('psa', content)

    def get_psa_email(self):
        context = {'panelist': self}

        template = 'emails/panelist_released.txt'
        subject = "[Barberscore] PSA for {0}".format(
            self.name,
        )
        to = ["{0} <{1}>".format(self.name, self.email)]
        cc = self.round.get_owners_emails()

        if self.psa_report:
            pdf = self.psa_report.file
        else:
            pdf = self.get_psa()
        file_name = '{0} PSA.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]

        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email

    def send_psa_email(self):
        email = self.get_psa_email()
        return email.send()

    # Transitions
    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.new, STATUS.active, STATUS.released],
        target=STATUS.released
    )
    def release(self, *args, **kwargs):
        # Saves PSA through post-transition signal to avoid race condition
        return


class Round(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
        (10, 'built', 'Built',),
        (20, 'started', 'Started',),
        (25, 'completed', 'Completed',),
        (27, 'finalized', 'Finalized',),
        (30, 'published', 'Published',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    KIND = Choices(
        (1, 'finals', 'Finals'),
        (2, 'semis', 'Semi-Finals'),
        (3, 'quarters', 'Quarter-Finals'),
    )

    kind = models.IntegerField(
        choices=KIND,
    )

    num = models.IntegerField(
        default=1,
    )

    spots = models.IntegerField(
        default=10,
    )

    date = models.DateField(
        blank=True,
        null=True,
    )

    footnotes = models.TextField(
        help_text="""
            Freeform text field; will print on OSS.""",
        blank=True,
    )

    oss_report = models.FileField(
        upload_to=UploadPath('oss_report'),
        blank=True,
        default='',
    )

    sa_report = models.FileField(
        upload_to=UploadPath('sa_report'),
        blank=True,
        default='',
    )

    legacy_oss = models.FileField(
        upload_to=UploadPath('legacy_oss'),
        blank=True,
        default='',
    )

    # Build artifact
    registration = JSONField(
        blank=True,
        null=True,
    )

    # Convention Denorm
    convention_id = models.UUIDField(
        blank=True,
        null=True,
    )

    name = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    DISTRICT_FULL_NAMES = {
        110: "Barbershop Harmony Society",
        200: "Cardinal",
        205: "Central States",
        210: "Dixie",
        215: "Evergreen",
        220: "Far Western",
        225: "Illinois",
        230: "Johnny Appleseed",
        235: "Land 'O Lakes",
        240: "Mid-Atlantic",
        345: "Northeastern",
        350: "Carolinas",
        355: "Ontario",
        360: "Pioneer",
        365: "Rocky Mountain",
        370: "Seneca Land",
        375: "Sunshine",
        380: "Southwestern",        
    }

    DISTRICT = Choices(
        (110, 'bhs', 'BHS'),
        (200, 'car', 'CAR'),
        (205, 'csd', 'CSD'),
        (210, 'dix', 'DIX'),
        (215, 'evg', 'EVG'),
        (220, 'fwd', 'FWD'),
        (225, 'ill', 'ILL'),
        (230, 'jad', 'JAD'),
        (235, 'lol', 'LOL'),
        (240, 'mad', 'MAD'),
        (345, 'ned', 'NED'),
        (350, 'nsc', 'NSC'),
        (355, 'ont', 'ONT'),
        (360, 'pio', 'PIO'),
        (365, 'rmd', 'RMD'),
        (370, 'sld', 'SLD'),
        (375, 'sun', 'SUN'),
        (380, 'swd', 'SWD'),
    )

    district = models.IntegerField(
        choices=DISTRICT,
        blank=True,
        null=True,
    )

    SEASON = Choices(
        # (1, 'summer', 'Summer',),
        # (2, 'midwinter', 'Midwinter',),
        (3, 'fall', 'Fall',),
        (4, 'spring', 'Spring',),
    )

    season = models.IntegerField(
        choices=SEASON,
        blank=True,
        null=True,
    )

    PANEL = Choices(
        (1, 'single', "Single"),
        (2, 'double', "Double"),
        (3, 'triple', "Triple"),
        (4, 'quadruple', "Quadruple"),
        (5, 'quintiple', "Quintiple"),
    )

    panel = models.IntegerField(
        choices=PANEL,
        blank=True,
        null=True,
    )

    YEAR_CHOICES = []
    for r in reversed(range(1939, (datetime.datetime.now().year + 11))):
        YEAR_CHOICES.append((r, r))

    year = models.IntegerField(
        choices=YEAR_CHOICES,
        blank=True,
        null=True,
    )

    open_date = models.DateField(
        blank=True,
        null=True,
    )

    close_date = models.DateField(
        blank=True,
        null=True,
    )

    start_date = models.DateField(
        blank=True,
        null=True,
    )

    end_date = models.DateField(
        blank=True,
        null=True,
    )

    venue_name = models.CharField(
        help_text="""
            The venue name (when available).""",
        max_length=255,
        blank=True,
        default='(TBD)',
    )

    location = models.CharField(
        help_text="""
            The general location in the form "City, State".""",
        max_length=255,
        blank=True,
        default='(TBD)',
    )

    timezone = TimeZoneField(
        help_text="""
            The local timezone of the convention.""",
        blank=True,
        null=True,
    )

    DIVISION = Choices(
        ('EVG', [
            (10, 'evgd1', 'EVG Division I'),
            (20, 'evgd2', 'EVG Division II'),
            (30, 'evgd3', 'EVG Division III'),
            (40, 'evgd4', 'EVG Division IV'),
            (50, 'evgd5', 'EVG Division V'),
        ]),
        ('FWD', [
            (60, 'fwdaz', 'FWD Arizona'),
            (70, 'fwdne', 'FWD Northeast'),
            (80, 'fwdnw', 'FWD Northwest'),
            (90, 'fwdse', 'FWD Southeast'),
            (100, 'fwdsw', 'FWD Southwest'),
        ]),
        ('LOL', [
            (110, 'lol10l', 'LOL 10000 Lakes'),
            (120, 'lolone', 'LOL Division One'),
            (130, 'lolnp', 'LOL Northern Plains'),
            (140, 'lolpkr', 'LOL Packerland'),
            (150, 'lolsw', 'LOL Southwest'),
        ]),
        ('MAD', [
            # (160, 'madatl', 'MAD Atlantic'),
            (170, 'madcen', 'MAD Central'),
            (180, 'madnth', 'MAD Northern'),
            (190, 'madsth', 'MAD Southern'),
            # (200, 'madwst', 'MAD Western'),
        ]),
        ('NED', [
            (210, 'nedgp', 'NED Granite and Pine'),
            (220, 'nedmtn', 'NED Mountain'),
            (230, 'nedpat', 'NED Patriot'),
            (240, 'nedsun', 'NED Sunrise'),
            (250, 'nedyke', 'NED Yankee'),
        ]),
        ('SWD', [
            (260, 'swdne', 'SWD Northeast'),
            (270, 'swdnw', 'SWD Northwest'),
            (280, 'swdse', 'SWD Southeast'),
            (290, 'swdsw', 'SWD Southwest'),
        ]),
    )

    divisions = DivisionsField(
        help_text="""Only select divisions if required.  If it is a district-wide convention do not select any.""",
        base_field=models.IntegerField(
            choices=DIVISION,
        ),
        blank=True,
        default=list,
    )

    image_id = models.CharField(
        max_length=255,
        blank=True,
        default='missing_image',
    )

    # Session Denorm
    session_id = models.UUIDField(
        blank=True,
        null=True,
    )

    session_nomen = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    SESSION_KIND = Choices(
        (32, 'chorus', "Chorus"),
        (41, 'quartet', "Quartet"),
        (42, 'mixed', "Mixed"),
        (43, 'senior', "Senior"),
        (44, 'youth', "Youth"),
        (45, 'unknown', "Unknown"),
        (46, 'vlq', "VLQ"),
    )

    session_kind = models.IntegerField(
        help_text="""
            The kind of session.  Generally this will be either quartet or chorus.
        """,
        choices=SESSION_KIND,
        blank=True,
        null=True,
    )

    # FKs
    owners = models.ManyToManyField(
        settings.AUTH_USER_MODEL,
        related_name='rounds',
    )

    # Relations
    statelogs = GenericRelation(
        StateLog,
        related_query_name='rounds',
    )

    # Properties
    @cached_property
    def nomen(self):
        return "{0} {1}".format(
            self.session_nomen,
            self.get_kind_display(),
        )

    # Internals
    class Meta:
        constraints = [
            models.UniqueConstraint(
                name='unique_round',
                fields=[
                    'session_id',
                    'kind',
                ]
            )
        ]
        get_latest_by = [
            'num',
        ]

    class JSONAPIMeta:
        resource_name = "round"

    def __str__(self):
        return self.nomen

    # Methods
    def get_owners_emails(self):
        Session = apps.get_model('registration.session')
        session = Session.objects.get(id=self.session_id)
        if not len(session.notification_list):
            raise ValueError("No notification list for {0}".format(session))
        return ["{0}".format(x.strip()) for x in session.notification_list.split(',')]

    def get_oss(self, published_by=None, zoom=1):
        Group = apps.get_model('bhs.group')
        Chart = apps.get_model('bhs.chart')
        Panelist = apps.get_model('adjudication.panelist')
        Appearance = apps.get_model('adjudication.appearance')
        Song = apps.get_model('adjudication.song')

        excluded_appearance_statuses = [
            Appearance.STATUS.disqualified, # Don't include disqualifications.
        ]
        if self.num == 1:
            # Don't include scratches
            excluded_appearance_statuses.append(Appearance.STATUS.scratched)

        # Score Block - get completeds
        publics = self.appearances.filter(
            is_private=False,
        ).exclude(
            # Don't include advancers on OSS
            draw__gt=0,
        ).exclude(
            status__in=excluded_appearance_statuses,
        ).exclude(
            # Don't include mic testers on OSS
            num__lte=0,
        ).all().order_by(
            RawSQL("(stats->>%s)::Numeric", ("tot_points",)).desc(nulls_last=True),
            RawSQL("(stats->>%s)::Numeric", ("sng_points",)).desc(nulls_last=True),
            RawSQL("(stats->>%s)::Numeric", ("per_points",)).desc(nulls_last=True),
        )

        # groups = Group.objects.filter(
        #     id__in=group_ids,
        # ).prefetch_related(
        #     'appearances',
        #     'appearances__songs__scores',
        #     'appearances__songs__scores__panelist',
        #     'appearances__round__session',
        # ).annotate(
        #     max_round=Max(
        #         'appearances__round__num',
        #         filter=Q(
        #             appearances__num__gt=0,
        #             appearances__round__session_id=self.session_id,
        #         ),
        #     ),
        #     tot_points=Sum(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     sng_points=Sum(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     per_points=Sum(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     mus_points=Sum(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     tot_score=Avg(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     mus_score=Avg(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.music,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     per_score=Avg(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     sng_score=Avg(
        #         'appearances__songs__scores__points',
        #         filter=Q(
        #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
        #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #             appearances__round__session_id=self.session_id,
        #             appearances__round__num__lte=self.num,
        #         ),
        #     ),
        #     tot_rank=Window(
        #         expression=RowNumber(),
        #         order_by=(
        #             F('tot_points').desc(),
        #             F('sng_points').desc(),
        #             F('per_points').desc(),
        #         )
        #     ),
        # ).order_by(
        #     '-tot_points',
        #     '-sng_points',
        #     '-per_points',
        # )

        songs_with_panelities = []

        if self.kind == self.KIND.finals:
            previous_rounds = Round.objects.filter(
                session_id=self.session_id,
            ).exclude(
                id=self.id,
            ).order_by('kind')

        # Monkeypatching
        i = self.spots
        for public in publics:
            i += 1
            public.tot_score_avg = 0 
            public.tot_rank = i
            tot_points = 0
            appearances = Appearance.objects.filter(
                group_id=public.group_id,
                round__session_id=self.session_id,
                round__num__lte=self.num,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).order_by(
                'round__kind',
            ).annotate(
                tot_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                tot_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
                mus_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.music,
                    ),
                ),
                per_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.performance,
                    ),
                ),
                sng_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                        songs__scores__panelist__category=Panelist.CATEGORY.singing,
                    ),
                ),
            )

            # print('count', appearances.count())

            scored_rounds = 0
            for appearance in appearances:
                if appearance.status != Appearance.STATUS.scratched:
                    scored_rounds += 1
                    songs = appearance.songs.prefetch_related(
                        'scores',
                        'scores__panelist',
                    ).order_by(
                        'num',
                    ).annotate(
                        tot_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                            ),
                        ),
                        mus_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.music,
                            ),
                        ),
                        per_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.performance,
                            ),
                        ),
                        sng_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.singing,
                            ),
                        ),
                    )
                    for song in songs:
                        if song.chart_id:
                            song.chart_patched = "{0} [{1}]".format(
                                song.title,
                                song.arrangers,
                            )
                        else:
                            song.chart_patched = ""

                        penalties_map = {
                            30: "†",
                            32: "‡",
                            34: "✠",
                            36: "✶",
                            38: "✢",
                            39: "✦",
                            40: "❉",
                            44: "∏",
                            48: "∇",
                            50: "※",
                        }
                        items = " ".join([penalties_map[x] for x in song.penalties])
                        if song.penalties and song.penalties not in songs_with_panelities:
                            songs_with_panelities.append(song.penalties)
                        song.penalties_patched = items
                    appearance.songs_patched = songs
                    public.tot_score_avg += appearance.tot_score
                    if public.stats is None:
                        public.stats = appearance.stats

            public.tot_score_avg = public.tot_score_avg / scored_rounds
            public.appearances_patched = appearances

            contesting = public.outcomes.order_by(
                'num',
            ).values_list(
                'num',
                flat=True
            )

            if self.kind == self.KIND.finals:
                for previous_round in previous_rounds:
                    previous_appearances = Appearance.objects.filter(
                            entry_id=public.entry_id,
                            is_private=False,
                        ).exclude(
                            status__in=excluded_appearance_statuses,
                        ).exclude(
                            # Don't include mic testers on OSS
                            num__lte=0,
                        ).exclude(
                            id=public.id
                        )

                    for previous_appearance in previous_appearances:

                        previously_contesting = previous_appearance.outcomes.filter(
                            print_on_finals_oss=True,
                        ).order_by(
                            'num',
                        ).values_list(
                            'num',
                            flat=True
                        )

                        merged_contesting = contesting | previously_contesting
                        contesting = merged_contesting.distinct()

            public.contesting_patched = ", ".join([str(x) for x in contesting])
            public.pos_patched = public.pos
            public.participants_patched = public.participants
            group = Group.objects.get(id=public.group_id)
            public.district = group.district
            if public.area == 'BHS':
                public.display_area = self.DISTRICT[group.district]
            else:
                public.display_area = public.area
            public.name = group.name

        # Penalties Block
        # array = Song.objects.select_related(
        # ).filter(
        #     appearance__round__session_id=self.session_id, # Using any session appearance
        #     penalties__len__gt=0, # Where there are penalties
        #     appearance__is_private=False, # If a public appearance

        #     # appearance__in=publics,  # Only completeds
        # ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            30: "† Scores penalized for Repeating Substantial Portions of a Song (V.A.2)",
            32: "‡ Scores penalized for Instrumental Accompaniment (IX.A.2.a)",
            34: "✠ Scores penalized for Chorus Exceeding 4-Part Texture (IX.A.2.b)",
            36: "✶ Scores penalized for Excessive Melody Not in Inner Part (IX.A.2.c)",
            38: "✢ Scores penalized for Lack of Characteristic Chord Progression (IX.A.2.d)",
            39: "✦ Scores penalized for Excessive Lyrics < 4 parts (IX.A.2.e)",
            40: "❉ Scores penalized for Primarily Patriotic/Religious Intent (IX.A.3)",
            44: "∏ Scores penalized for Not in Good Taste (IX.A.3.b)",
            48: "∇ Scores penalized for Non-members Performing on Stage (XI.A.1)",
            50: "※ Scores penalized for Sound Equipment or Electronic Enhancement (X.B.1-3)",
        }
        penalties = sorted(list(set(penalties_map[x] for l in songs_with_panelities for x in l)))

        # Missing flag
        # use group_ids - these are the completeds
        is_missing = bool(Song.objects.filter(
            appearance__round__session_id=self.session_id,
            chart_id__isnull=True,
            appearance__in=publics,
        ))
        # Eval Only Block
        privates = self.appearances.prefetch_related(
        ).filter(
            is_private=True,
        ).exclude(
            status__in=[
                Appearance.STATUS.disqualified,
                Appearance.STATUS.scratched,
            ],
        ).order_by(
        ).values_list('name', flat=True)
        privates = list(privates)

        # Disqualification Block
        disqualifications = self.appearances.prefetch_related(
        ).filter(
            status=Appearance.STATUS.disqualified,
        ).order_by(
        ).values_list('name', flat=True)
        disqualifications = list(disqualifications)

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancer_group_ids = self.appearances.filter(
                draw__gt=0,
            ).select_related(
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'group_id',
            )
            advancers = []
            for draw, group_id in advancer_group_ids:
                name = Group.objects.get(id=group_id).name
                advancers.append((draw, name))
            try:
                mt_id = self.appearances.get(
                    draw=0,
                ).group_id
                mt = Group.objects.get(id=mt_id).name
                advancers.append(('MT', mt))
            except self.appearances.model.DoesNotExist:
                pass
        else:
            advancers = None

        # Panelist Block
        panelists_raw = self.panelists.select_related(
        ).filter(
            kind=Panelist.KIND.official,
            category__gte=Panelist.CATEGORY.pc,
        ).order_by(
            'num',
        )
        categories_map = {
            9: 'PC',
            10: 'ADM',
            30: 'MUS',
            40: 'PER',
            50: 'SNG',
        }
        panelists = []
        for key, value in categories_map.items():
            sections = panelists_raw.filter(
                category=key,
            ).select_related(
            ).order_by(
                'num',
            )
            persons = []
            for x in sections:
                try:
                    persons.append(
                        "{0}".format(
                            x.name,
                            # x.get_district_display(),
                        )
                    )
                except AttributeError:
                    persons.append("(Unknown)")
            if persons: 
                names = ", ".join(persons)
                panelists.append((value, names))

        # Outcome Block
        items = self.outcomes.select_related(
            # 'award',
        ).filter(
            printed=True,
        ).order_by(
            'tree_sort',
            'num',
        ).values_list(
            'num',
            'name',
            'winner',
        )
        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} {1}".format(item[0], item[1]),
                    item[2],
                )
            )

        if self.kind == self.KIND.finals:
            for pr in previous_rounds:
                items = pr.outcomes.select_related(
                    # 'award',
                ).filter(
                    print_on_finals_oss=True,
                ).order_by(
                    'tree_sort',
                    'num',
                ).values_list(
                    'num',
                    'name',
                    'winner',
                )
                for item in items:
                    outcomes.append(
                        (
                            "{0} {1}".format(item[0], item[1]),
                            item[2],
                        )
                    )

        context = {
            'round': self,
            'publics': publics,
            'penalties': penalties,
            'privates': privates,
            'disqualifications': disqualifications,
            'advancers': advancers,
            'panelists': panelists,
            'outcomes': outcomes,
            'is_missing': is_missing,
        }
        rendered = render_to_string('reports/oss.html', context)

        if self.get_district_display() == 'BHS':
            Session = apps.get_model('registration.session')
            session = Session.objects.get(id=self.session_id)

            if session.convention.name == 'International Youth Convention':
                page_size = 'Legal'
            elif session.kind == session.KIND.quartet and self.kind == self.KIND.semis:
                page_size = 'Legal'
            else:
                page_size = 'Letter'
        else:
            if self.kind == self.KIND.finals:
                if publics.count() >= 10:
                    page_size = 'Legal'
                else:
                    page_size = 'Letter'
            elif self.kind == self.KIND.semis:
                if publics.count() >= 12:
                    page_size = 'Legal'
                else:
                    page_size = 'Letter'
            else:
                page_size = 'Legal'
        try:
            if published_by is not None:
                footer = 'Published by {0} at {1}'.format(
                    published_by,
                    datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S %Z"),
                )
            else:
                statelog = self.statelogs.latest('timestamp')
                footer = 'Published by {0} at {1}'.format(
                    statelog.by.name,
                    statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
                )
        except StateLog.DoesNotExist:
            footer = '(Unknown)'
        file = pydf.generate_pdf(
            rendered,
            page_size=page_size,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
            zoom=zoom,
        )
        content = ContentFile(file)
        return content

    def save_oss(self):
        oss = self.get_oss()
        return self.oss_report.save('oss', oss)

    # Placeholder for single-page recursion tool.
    # from builtins import round as rnd
    # from PyPDF2 import PdfFileReader

    # r = Round.objects.get(id='0c02d601-357b-4315-9b73-063f563aa410')

    # f = r.get_oss()
    # obj = PdfFileReader(f)
    # num_pages = obj.getNumPages()

    # zoom = 1
    # while num_pages > 1:
    #     zoom = rnd((zoom - .05),2)
    #     print(zoom)
    #     f = r.get_oss(zoom=zoom)
    #     obj = PdfFileReader(f)
    #     num_pages = obj.getNumPages()
    # print(type(f))

    # r.oss.save('oss', f)

    def get_sa(self, published_by=None):
        Group = apps.get_model('bhs.group')
        Chart = apps.get_model('bhs.chart')
        Person = apps.get_model('bhs.person')
        Appearance = apps.get_model('adjudication.appearance')
        Panelist = apps.get_model('adjudication.panelist')
        Song = apps.get_model('adjudication.song')

        # Score Block
        # group_ids = self.appearances.exclude(
        #     # DOn't include scratches
        #     status=Appearance.STATUS.scratched,
        # ).exclude(
        #     # Don't include disqualifications.
        #     status=Appearance.STATUS.disqualified,
        # ).exclude(
        #     # Don't include advancers on SA
        #     draw__gt=0,
        # ).values_list('group_id', flat=True)
        # for group_id in group_ids:
        #     group_appearances = Appearance.objects.prefetch_related(
        #         'songs__scores',
        #         'songs__scores__panelist',
        #     ).filter(
        #         group_id=group_id,
        #         round__session_id=self.session_id,
        #     ).annotate(
        #         tot_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         mus_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.music,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         per_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         sng_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         tot_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         mus_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.music,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         per_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         sng_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #                 round__session_id=self.session_id,
        #                 round__num__lte=self.num,
        #             ),
        #         ),
        #         tot_rank=Window(
        #             expression=RowNumber(),
        #             order_by=(
        #                 F('tot_points').desc(),
        #                 F('sng_points').desc(),
        #                 F('per_points').desc(),
        #             )
        #         ),
        #         mus_rank=Window(
        #             expression=Rank(),
        #             order_by=F('mus_points').desc(),
        #         ),
        #         per_rank=Window(
        #             expression=Rank(),
        #             order_by=F('per_points').desc(),
        #         ),
        #         sng_rank=Window(
        #             expression=Rank(),
        #             order_by=F('sng_points').desc(),
        #         ),
        #     ).order_by(
        #         '-tot_points',
        #         '-sng_points',
        #         '-per_points',
        #     )

        mus_persons_qs = Panelist.objects.filter(
            category=Panelist.CATEGORY.music,
            round__session_id=self.session_id,
        ).order_by(
            'num',
        ).distinct('num')

        mus_persons = []
        for p in mus_persons_qs:
            practice = bool(p.kind == Panelist.KIND.practice)
            num = "{0}".format('{:02d}'.format(p.num))
            mus_persons.append((p.name, practice, num))

        per_persons_qs = Panelist.objects.filter(
            category=Panelist.CATEGORY.performance,
            round__session_id=self.session_id,
        ).order_by(
            'num',
        ).distinct('num')
        per_persons = []
        for p in per_persons_qs:
            practice = bool(p.kind == Panelist.KIND.practice)
            num = "{0}".format('{:02d}'.format(p.num))
            per_persons.append((p.name, practice, num))

        sng_persons_qs = Panelist.objects.filter(
            category=Panelist.CATEGORY.singing,
            round__session_id=self.session_id,
        ).order_by(
            'num',
        ).distinct('num')
        sng_persons = []
        for p in sng_persons_qs:
            practice = bool(p.kind == Panelist.KIND.practice)
            num = "{0}".format('{:02d}'.format(p.num))
            sng_persons.append((p.name, practice, num))
        # per_persons_qs = persons.filter(
        #     panelists__category=Panelist.CATEGORY.performance,
        #     panelists__round__session_id=self.session_id,
        # ).order_by(
        #     'panelists__num',
        # ).distinct()
        # per_persons = []
        # for p in per_persons_qs:
        #     practice = bool(p.panelists.get(round=self).kind == Panelist.KIND.practice)
        #     per_persons.append((p.name, practice, p.initials))

        # sng_persons_qs = persons.filter(
        #     panelists__category=Panelist.CATEGORY.singing,
        #     panelists__round__session_id=self.session_id,
        # ).order_by(
        #     'panelists__num',
        # ).distinct()
        # sng_persons = []
        # for p in sng_persons_qs:
        #     practice = bool(p.panelists.get(round=self).kind == Panelist.KIND.practice)
        #     sng_persons.append((p.name, practice, p.initials))

        # Monkeypatching
        # for all_appearance in all_appearances:
        #     # Populate the group block
        #     appearances = group.appearances.filter(
        #         round__session=self.session,
        #         round__num__lte=self.num,
        #     ).prefetch_related(
        #         'songs__scores',
        #         'songs__scores__panelist',
        #     ).annotate(
        #         tot_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         sng_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #             ),
        #         ),
        #         per_points=Sum(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #             ),
        #         ),
        #         tot_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #             ),
        #         ),
        #         mus_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.music,
        #             ),
        #         ),
        #         per_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.performance,
        #             ),
        #         ),
        #         sng_score=Avg(
        #             'songs__scores__points',
        #             filter=Q(
        #                 songs__scores__panelist__kind=Panelist.KIND.official,
        #                 songs__scores__panelist__category=Panelist.CATEGORY.singing,
        #             ),
        #         ),
        #     ).order_by(
        #         'round__kind',
        #     )
        #     for appearance in appearances:
        #         songs = appearance.songs.prefetch_related(
        #             'scores',
        #             'scores__panelist',
        #         ).order_by(
        #             'num',
        #         ).annotate(
        #             tot_score=Avg(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                 ),
        #             ),
        #             mus_score=Avg(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.music,
        #                 ),
        #             ),
        #             per_score=Avg(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.performance,
        #                 ),
        #             ),
        #             sng_score=Avg(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.singing,
        #                 ),
        #             ),
        #             tot_dev=StdDev(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                 ),
        #             ),
        #             mus_dev=StdDev(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.music,
        #                 ),
        #             ),
        #             per_dev=StdDev(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.performance,
        #                 ),
        #             ),
        #             sng_dev=StdDev(
        #                 'scores__points',
        #                 filter=Q(
        #                     scores__panelist__kind=Panelist.KIND.official,
        #                     scores__panelist__category=Panelist.CATEGORY.singing,
        #                 ),
        #             ),
        #         )
        #         for song in songs:
        #             penalties_map = {
        #                 30: "†",
        #                 32: "‡",
        #                 34: "✠",
        #                 36: "✶",
        #                 38: "✢",
        #                 39: "✦",
        #                 40: "❉",
        #                 50: "※",
        #             }
        #             items = " ".join([penalties_map[x] for x in song.penalties])
        #             song.penalties_patched = items


        #         for song in songs:
        #             chart = Chart.objects.get(id=song.chart_id)
        #             song.chart_patched = chart

        #             mus_scores = []
        #             for person in mus_persons_qs:
        #                 raw_music_scores = song.scores.filter(
        #                     panelist__person=person,
        #                 ).order_by(
        #                     'panelist__num',
        #                 )
        #                 if raw_music_scores:
        #                     for m in raw_music_scores:
        #                         diff = abs(m.points - m.song.mus_score) > 5
        #                         practice = bool(m.panelist.kind == Panelist.KIND.practice)
        #                         mus_scores.append((m.points, diff, practice))
        #                 else:
        #                     mus_scores.append((None, False, False))
        #             song.mus_scores = mus_scores

        #             per_scores = []
        #             for person in per_persons_qs:
        #                 raw_performance_scores = song.scores.filter(
        #                     panelist__person=person,
        #                 ).order_by(
        #                     'panelist__num',
        #                 )
        #                 if raw_performance_scores:
        #                     for m in raw_performance_scores:
        #                         diff = abs(m.points - m.song.per_score) > 5
        #                         practice = bool(m.panelist.kind == Panelist.KIND.practice)
        #                         per_scores.append((m.points, diff, practice))
        #                 else:
        #                     per_scores.append((None, False, False))
        #             song.per_scores = per_scores

        #             sng_scores = []
        #             for person in sng_persons_qs:
        #                 raw_singing_scores = song.scores.filter(
        #                     panelist__person=person,
        #                 ).order_by(
        #                     'panelist__num',
        #                 )
        #                 if raw_singing_scores:
        #                     for m in raw_singing_scores:
        #                         diff = abs(m.points - m.song.sng_score) > 5
        #                         practice = bool(m.panelist.kind == Panelist.KIND.practice)
        #                         sng_scores.append((m.points, diff, practice))
        #                 else:
        #                     sng_scores.append((None, False, False))
        #             song.sng_scores = sng_scores


        #         appearance.songs_patched = songs
        #     group.appearances_patched = appearances

        excluded_appearance_statuses = [
            Appearance.STATUS.disqualified, # Don't include disqualifications.
        ]
        if self.num == 1:
            # Don't include scratches
            excluded_appearance_statuses.append(Appearance.STATUS.scratched)

        groups = self.appearances.filter(
            # is_private=False,
            num__gt=0,
        ).exclude(
            # Don't include advancers on SA
            draw__gt=0,
        ).exclude(
            status__in=excluded_appearance_statuses
        ).exclude(
            # Don't include mic testers on OSS
            # num__lte=0,
        # ).annotate(
        #     # tot_round_rank=Window(
        #     #     expression=RowNumber(),
        #     #     partition_by=[F('round_id')],
        #     #     order_by=(
        #     #         F('round_stats__tot_points').desc(),
        #     #         F('round_stats__sng_points').desc(),
        #     #         F('round_stats__per_points').desc(),
        #     #     )
        #     # ),
        #     mus_round_rank=Window(
        #         expression=Rank(),
        #         partition_by=[F('round_id')],
        #         order_by=F('round_stats__mus_points').desc(),
        #     ),
        #     per_round_rank=Window(
        #         expression=Rank(),
        #         partition_by=[F('round_id')],
        #         order_by=F('round_stats__per_points').desc(),
        #     ),
        #     sng_round_rank=Window(
        #         expression=Rank(),
        #         partition_by=[F('round_id')],
        #         order_by=F('round_stats__sng_points').desc(),
        #     ),
        ).all().order_by(
            RawSQL("(stats->>%s)::Numeric", ("tot_points",)).desc(nulls_last=True),
            RawSQL("(stats->>%s)::Numeric", ("sng_points",)).desc(nulls_last=True),
            RawSQL("(stats->>%s)::Numeric", ("per_points",)).desc(nulls_last=True),
        )

        round_rankings = Appearance.objects.filter(
            round__session_id=self.session_id,
            round__num__lte=self.num,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).annotate(
            mus_round_rank=Window(
                expression=Rank(),
                partition_by=[F('round_id')],
                order_by=Cast(JSONF('round_stats__mus_points'), models.FloatField()).desc(),
            ),
            per_round_rank=Window(
                expression=Rank(),
                partition_by=[F('round_id')],
                order_by=Cast(JSONF('round_stats__per_points'), models.FloatField()).desc(),
            ),
            sng_round_rank=Window(
                expression=Rank(),
                partition_by=[F('round_id')],
                order_by=Cast(JSONF('round_stats__sng_points'), models.FloatField()).desc(),
            ),
        )

        rankings = {}
        for appearance in round_rankings:
            rankings[appearance.id] = appearance

        for group in groups:
            group.tot_score_avg = 0
            appearances = Appearance.objects.filter(
                group_id=group.group_id,
                round__session_id=self.session_id,
                round__num__lte=self.num,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).order_by(
                'round__kind',
            ).annotate(
                tot_score=Avg(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
            )
            # ).annotate(
            #     tot_rank=Window(
            #         expression=RowNumber(),
            #         partition_by=[F('group_id')],
            #         order_by=(
            #             F('round_stats__tot_points').desc(),
            #             F('round_stats__sng_points').desc(),
            #             F('round_stats__per_points').desc(),
            #         )
            #     ),
            #     mus_round_rank=Window(
            #         expression=Rank(),
            #         partition_by=[F('group_id')],
            #         order_by=F('round_stats__mus_points').desc(),
            #     ),
            #     per_round_rank=Window(
            #         expression=Rank(),
            #         partition_by=[F('group_id')],
            #         order_by=F('round_stats__per_points').desc(),
            #     ),
            #     sng_round_rank=Window(
            #         expression=Rank(),
            #         partition_by=[F('group_id')],
            #         order_by=F('round_stats__sng_points').desc(),
            #     ),
            # )

            # ).annotate(
            #     tot_points=Sum(
            #         'songs__scores__points',
            #         filter=Q(
            #             songs__scores__panelist__kind=Panelist.KIND.official,
            #         ),
            #     ),
            #     tot_score=Avg(
            #         'songs__scores__points',
            #         filter=Q(
            #             songs__scores__panelist__kind=Panelist.KIND.official,
            #         ),
            #     ),
            #     mus_score=Avg(
            #         'songs__scores__points',
            #         filter=Q(
            #             songs__scores__panelist__kind=Panelist.KIND.official,
            #             songs__scores__panelist__category=Panelist.CATEGORY.music,
            #         ),
            #     ),
            #     per_score=Avg(
            #         'songs__scores__points',
            #         filter=Q(
            #             songs__scores__panelist__kind=Panelist.KIND.official,
            #             songs__scores__panelist__category=Panelist.CATEGORY.performance,
            #         ),
            #     ),
            #     sng_score=Avg(
            #         'songs__scores__points',
            #         filter=Q(
            #             songs__scores__panelist__kind=Panelist.KIND.official,
            #             songs__scores__panelist__category=Panelist.CATEGORY.singing,
            #         ),
            #     ),
            # )
            scored_rounds = 0
            for appearance in appearances:
                if appearance.status != Appearance.STATUS.scratched:
                    scored_rounds += 1
                    songs = appearance.songs.prefetch_related(
                        'scores',
                        'scores__panelist',
                    ).order_by(
                        'num',
                    ).annotate(
                        tot_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                            ),
                        ),
                        mus_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.music,
                            ),
                        ),
                        per_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.performance,
                            ),
                        ),
                        sng_score=Avg(
                            'scores__points',
                            filter=Q(
                                scores__panelist__kind=Panelist.KIND.official,
                                scores__panelist__category=Panelist.CATEGORY.singing,
                            ),
                        ),
                    )

                    for song in songs:
                        if song.chart_id:
                            song.chart_patched = "{0} [{1}]".format(
                                song.title,
                                song.arrangers,
                            )
                        else:
                            song.chart_patched = ""
                        mus_scores = []
                        for panelist in mus_persons_qs:
                            raw_music_scores = song.scores.filter(
                                panelist__person_id=panelist.person_id,
                                panelist__kind=panelist.kind
                            ).order_by(
                                'panelist__num',
                            )

                            if raw_music_scores:
                                for m in raw_music_scores:
                                    diff = abs(m.points - m.song.mus_score) > 5
                                    practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                    mus_scores.append((m.points, diff, practice))
                            else:
                                mus_scores.append((None, False, False))
                        song.mus_scores = mus_scores

                        per_scores = []
                        for panelist in per_persons_qs:
                            raw_performance_scores = song.scores.filter(
                                panelist__person_id=panelist.person_id,
                                panelist__kind=panelist.kind
                            ).order_by(
                                'panelist__num',
                            )

                            if raw_performance_scores:
                                for m in raw_performance_scores:
                                    diff = abs(m.points - m.song.per_score) > 5
                                    practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                    per_scores.append((m.points, diff, practice))
                            else:
                                per_scores.append((None, False, False))
                        song.per_scores = per_scores

                        sng_scores = []
                        for panelist in sng_persons_qs:
                            raw_singing_scores = song.scores.filter(
                                panelist__person_id=panelist.person_id,
                                panelist__kind=panelist.kind
                            ).order_by(
                                'panelist__num',
                            )
                            if raw_singing_scores:
                                for m in raw_singing_scores:
                                    diff = abs(m.points - m.song.sng_score) > 5
                                    practice = bool(m.panelist.kind == Panelist.KIND.practice)
                                    sng_scores.append((m.points, diff, practice))
                            else:
                                sng_scores.append((None, False, False))
                        song.sng_scores = sng_scores
                        penalties_map = {
                            30: "†",
                            32: "‡",
                            34: "✠",
                            36: "✶",
                            38: "✢",
                            39: "✦",
                            40: "❉",
                            44: "∏",
                            48: "∇",
                            50: "※",
                        }
                        items = " ".join([penalties_map[x] for x in song.penalties])
                        song.penalties_patched = items
                    appearance.songs_patched = songs
                    appearance.rankings = rankings[appearance.id]
                    group.tot_score_avg += appearance.tot_score
                    if group.stats is None:
                        group.stats = appearance.stats
                        
            group.tot_score_avg = group.tot_score_avg / scored_rounds
            group.appearances_patched = appearances

        # Penalties Block
        array = Song.objects.select_related(
            'appearance__round',
            'appearance__group',
        ).filter(
            appearance__round__session_id=self.session_id, # Using any session appearance
            penalties__len__gt=0, # Where there are penalties
            appearance__is_private=False, # If a public appearance
            # appearance__in=groups,  # Only completeds
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            30: "† Scores penalized for Repeating Substantial Portions of a Song (V.A.2)",
            32: "‡ Scores penalized for Instrumental Accompaniment (IX.A.2.a)",
            34: "✠ Scores penalized for Chorus Exceeding 4-Part Texture (IX.A.2.b)",
            36: "✶ Scores penalized for Excessive Melody Not in Inner Part (IX.A.2.c)",
            38: "✢ Scores penalized for Lack of Characteristic Chord Progression (IX.A.2.d)",
            39: "✦ Scores penalized for Excessive Lyrics < 4 parts (IX.A.2.e)",
            40: "❉ Scores penalized for Primarily Patriotic/Religious Intent (IX.A.3)",
            44: "∏ Scores penalized for Not in Good Taste (IX.A.3.b)",
            48: "∇ Scores penalized for Non-members Performing on Stage (XI.A.1)",
            50: "※ Scores penalized for Sound Equipment or Electronic Enhancement (X.B.1-3)",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        # Build stats
        # stats = Song.objects.select_related(
        #     'appearance__round',
        # ).prefetch_related(
        #     'scores__panelist__kind',
        # ).filter(
        #     appearance__round__session_id=self.session_id,
        #     appearance__round__num__lte=self.num,
        # ).annotate(
        #     tot_dev=StdDev(
        #         'scores__points',
        #         filter=Q(
        #             scores__panelist__kind=Panelist.KIND.official,
        #         ),
        #     ),
        #     mus_dev=StdDev(
        #         'scores__points',
        #         filter=Q(
        #             scores__panelist__kind=Panelist.KIND.official,
        #             scores__panelist__category=Panelist.CATEGORY.music,
        #         ),
        #     ),
        #     per_dev=StdDev(
        #         'scores__points',
        #         filter=Q(
        #             scores__panelist__kind=Panelist.KIND.official,
        #             scores__panelist__category=Panelist.CATEGORY.performance,
        #         ),
        #     ),
        #     sng_dev=StdDev(
        #         'scores__points',
        #         filter=Q(
        #             scores__panelist__kind=Panelist.KIND.official,
        #             scores__panelist__category=Panelist.CATEGORY.singing,
        #         ),
        #     ),
        # ).aggregate(
        #     tot=Avg('tot_dev'),
        #     mus=Avg('mus_dev'),
        #     per=Avg('per_dev'),
        #     sng=Avg('sng_dev'),
        # )

        # Category Stats
        stats = self.get_session_stats()

        context = {
            'round': self,
            'groups': groups,
            'mus_persons': mus_persons,
            'per_persons': per_persons,
            'sng_persons': sng_persons,
            'stats': stats,
            'penalties': penalties,
        }
        rendered = render_to_string('reports/sa.html', context)
        # return rendered

        if published_by is not None:
            footer = 'Published by {0} at {1}'.format(
                published_by,
                datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S %Z"),
            )
        else:
            statelog = self.statelogs.latest('timestamp')
            footer = 'Published by {0} at {1}'.format(
                statelog.by.name,
                statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
            )
 
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Landscape',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Encode Sans',
            footer_font_size=6,
        )
        content = ContentFile(file)
        return content

    def get_session_stats(self):
        stats = {}
        # If current round == 1 and more than one round exists calculate Grand Total Figures.
        if self.kind > 1:
            round_name = Round.KIND[self.kind]

            # Retrieve Grand Total Stats for a specific round
            stats['Total'] = {}
            stats['Total'][round_name] = self.get_round_stats(self.kind, None)

            for category in Panelist.CATEGORY:
                if category[0] >= 30:
                    stats[category[1]] = {}
        
                    # Retrieve category stats for a specific round
                    stats[category[1]][round_name] = self.get_round_stats(self.kind, category[0])        
        else:
            # How many rounds are there???
            rounds = Round.objects.filter(
                session_id=self.session_id,
            ).order_by('kind')

            stats['Total'] = {}
            if rounds.count() > 1:
                stats['Total']['Grand'] = self.get_round_stats(None, None)

                for round in rounds:
                    round_name = Round.KIND[round.kind]
                    stats['Total'][round_name] = self.get_round_stats(round.kind, None)
            else:
                stats['Total']['Finals'] = self.get_round_stats(self.kind, None)                

            for category in Panelist.CATEGORY:
                if category[0] >= 30:
                    stats[category[1]] = {}
                    
                    if rounds.count() > 1:
                        stats[category[1]]['Grand'] = self.get_round_stats(None, category[0])

                    for round in rounds:
                        round_name = Round.KIND[round.kind]
                    
                        # Retrieve category stats for a specific round
                        stats[category[1]][round_name] = self.get_round_stats(round.kind, category[0])
        return stats

    def get_round_stats(self, round_num, category):
        filter_round = ""
        if round_num is not None:
            filter_round = " AND round.kind = {0}".format(round_num)

        filter_category = ""
        if category is not None:
            filter_category = " AND panelist.category = {0}".format(category)

        query = '''
            SELECT 
                scores.id,
                ROUND(AVG(scores.resid) OVER (),2) AS diff,
                ROUND(stddev_samp(scores.resid) OVER (),2) AS stddev
            FROM (
                SELECT round.id, song.id as song_id, appearance.name, song.title, panelist.category, panelist.last_name, score.points, round.id as round_id, round.kind,

                avg(score.points) over(partition by song.id,panelist.category) as AVG_SCORE,

                ABS(avg(score.points) over(partition by song.id,panelist.category) - score.points) as resid

                FROM public.adjudication_score as score

                LEFT JOIN adjudication_song as song ON song.id = score.song_id
                LEFT JOIN adjudication_appearance as appearance ON song.appearance_id = appearance.id
                LEFT JOIN adjudication_panelist as panelist ON panelist.id = score.panelist_id
                LEFT JOIN adjudication_round as round ON round.id = appearance.round_id
                LEFT JOIN registration_session as session on session.id = round.session_id

                WHERE session.id = '%s' -- Session ID
                AND panelist.kind = 10 -- Official scores only
                %s -- Round Number
                %s -- Category
                    
                ORDER BY appearance.num, score.song_id, panelist.num

            ) as scores

            LIMIT 1
        ''' % (self.session_id, filter_round, filter_category)

        return Score.objects.raw(query)[0]

    def save_sa(self):
        sa = self.get_sa()
        return self.sa_report.save('sa', sa)

    def save_reports(self):
        oss = self.get_oss()
        self.oss_report.save('oss', oss, save=False)
        sa = self.get_sa()
        self.sa_report.save('sa', sa, save=False)
        return self.save()

    def get_legacy_oss(self):
        Chart = apps.get_model('bhs.chart')
        Panelist = apps.get_model('adjudication.panelist')
        Appearance = apps.get_model('adjudication.appearance')
        Song = apps.get_model('adjudication.song')
        Score = apps.get_model('adjudication.score')

        # Get the Groups
        group_ids = self.appearances.filter(
            is_private=False,
        ).exclude(
            # Don't include advancers on OSS
            draw__gt=0,
        ).exclude(
            # Don't include scratches
            status=Appearance.STATUS.scratched,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            # Don't include mic testers on OSS
            num__lte=0,
        ).values_list('group_id', flat=True)

        session = self.session

        rounds = session.rounds.filter(
            num__lte=self.num
        ).order_by('kind')

        # Monkeypatching
        for round in rounds:
            appearances = round.appearances.filter(
                group_id__in=group_ids,
            ).prefetch_related(
                'songs__scores',
                'songs__scores__panelist',
            ).annotate(
                sub_points=Sum(
                    'songs__scores__points',
                    filter=Q(
                        songs__scores__panelist__kind=Panelist.KIND.official,
                    ),
                ),
            )
            for appearance in appearances:
                if appearance.round.num == 1:
                    appearance.prev_points = None
                    appearance.sub_points = None
                songs = appearance.songs.prefetch_related(
                    'scores',
                    'scores__panelist',
                ).order_by(
                    'num',
                ).annotate(
                    tot_score=Avg(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                        ),
                    ),
                    mus_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.music,
                        ),
                    ),
                    per_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.performance,
                        ),
                    ),
                    sng_points=Sum(
                        'scores__points',
                        filter=Q(
                            scores__panelist__kind=Panelist.KIND.official,
                            scores__panelist__category=Panelist.CATEGORY.singing,
                        ),
                    ),
                )
                for song in songs:
                    chart = Chart.objects.get(id=song.chart_id)
                    song.chart_patched = chart
                    penalties_map = {
                        30: "†",
                        32: "‡",
                        34: "✠",
                        36: "✶",
                        38: "✢",
                        39: "✦",
                        40: "❉",
                        44: "∏",
                        48: "∇",
                        50: "※",
                    }
                    items = " ".join([penalties_map[x] for x in song.penalties])
                    song.penalties_patched = items
                raise RuntimeError('contender')
                contesting = appearance.contenders.order_by(
                    'outcome__num',
                ).values_list('outcome__num', flat=True)
                appearance.contesting = ", ".join([str(x) for x in contesting])
                tot_score = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num__lte=round.num,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(tot_score=Avg('points'))['tot_score']
                appearance.tot_score = tot_score
                tot_points = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num__lte=round.num,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(tot_points=Sum('points'))['tot_points']
                appearance.tot_points = tot_points if tot_points else 0
                prev_points = Score.objects.filter(
                    song__appearance__round__session=self.session,
                    song__appearance__round__num=round.num - 1,
                    song__appearance__group=appearance.group,
                    panelist__kind=Panelist.KIND.official,
                ).aggregate(prev_points=Sum('points'))['prev_points']
                appearance.prev_points = prev_points
                appearance.songs_patched = songs
            appearances = sorted(appearances, reverse=True, key=lambda t: t.tot_points)
            round.appearances_patched = appearances

        # Penalties Block
        array = Song.objects.select_related(
            'appearance__round',
            'appearance__group',
        ).filter(
            appearance__round__session=self.session, # Using any session appearance
            penalties__len__gt=0, # Where there are penalties
            appearance__is_private=False, # If a public appearance
            appearance__group_id__in=group_ids,  # Only completeds
        ).distinct().values_list('penalties', flat=True)
        penalties_map = {
            30: "† Scores penalized for Repeating Substantial Portions of a Song (V.A.2)",
            32: "‡ Scores penalized for Instrumental Accompaniment (IX.A.2.a)",
            34: "✠ Scores penalized for Chorus Exceeding 4-Part Texture (IX.A.2.b)",
            36: "✶ Scores penalized for Excessive Melody Not in Inner Part (IX.A.2.c)",
            38: "✢ Scores penalized for Lack of Characteristic Chord Progression (IX.A.2.d)",
            39: "✦ Scores penalized for Excessive Lyrics < 4 parts (IX.A.2.e)",
            40: "❉ Scores penalized for Primarily Patriotic/Religious Intent (IX.A.3)",
            44: "∏ Scores penalized for Not in Good Taste (IX.A.3.b)",
            48: "∇ Scores penalized for Non-members Performing on Stage (XI.A.1)",
            50: "※ Scores penalized for Sound Equipment or Electronic Enhancement (X.B.1-3)",
        }
        penalties = sorted(list(set(penalties_map[x] for l in array for x in l)))

        # Missing flag
        # use group_ids - these are the completeds
        is_missing = bool(Song.objects.filter(
            appearance__round__session=self.session,
            chart_id__isnull=True,
            appearance__group_id__in=group_ids,
        ))

        # Eval Only Block
        privates = self.appearances.prefetch_related(
            'group',
        ).filter(
            is_private=True,
        ).order_by(
            'group__name',
        ).values_list('group__name', flat=True)
        privates = list(privates)

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancers = self.appearances.filter(
                draw__gt=0,
            ).select_related(
                'group',
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'group__name',
            )
            advancers = list(advancers)
            try:
                mt = self.appearances.get(
                    draw=0,
                ).group.name
                advancers.append(('MT', mt))
            except self.appearances.model.DoesNotExist:
                pass
        else:
            advancers = None

        # Panelist Block
        panelists_raw = self.panelists.select_related(
        ).filter(
            kind=Panelist.KIND.official,
            category__gte=Panelist.CATEGORY.pc,
        ).order_by(
            'num',
        )
        categories_map = {
            9: 'PC',
            10: 'ADM',
            30: 'Music',
            40: 'Performance',
            50: 'Singing',
        }
        panelists = {}
        for key, value in categories_map.items():
            sections = panelists_raw.filter(
                category=key,
            ).select_related(
            ).order_by(
                'num',
            )
            persons = []
            for x in sections:
                persons.append("{0} - {1}".format(x.name, x.area))
            panelists[value] = persons


        # Outcome Block
        items = self.outcomes.select_related(
            # 'award',
        ).order_by(
            'num',
        ).values_list(
            'num',
            # 'award__name',
            'name',
        )
        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} is the winner of the {1}".format(item[2], item[1])
                )
            )
        contests = []
        for item in items:
            contests.append(
                "{0}={1}".format(item[0], item[1])
            )
        contest_legend = ", ".join(contests)

        context = {
            'round': self,
            'contest_legend': contest_legend,
            'rounds': rounds,
            'penalties': penalties,
            'privates': privates,
            'advancers': advancers,
            'panelists': panelists,
            'outcomes': outcomes,
            'is_missing': is_missing,
        }
        rendered = render_to_string('reports/legacy_oss.html', context)

        # if self.session.rounds.count() == 1:
        #     if groups.count() < 16:
        #         page_size = 'Letter'
        #     else:
        #         page_size = 'Legal'
        # else:
        #     if groups.count() > 8 and self.kind == self.KIND.finals:
        #         page_size = 'Legal'
        #     else:
        #         page_size = 'Letter'
        page_size = 'Letter'
        try:
            statelog = self.statelogs.latest('timestamp')
            footer = 'Published by {0} at {1}'.format(
                statelog.by.name,
                statelog.timestamp.strftime("%Y-%m-%d %H:%M:%S %Z"),
            )
        except StateLog.DoesNotExist:
            footer = ""
        file = pydf.generate_pdf(
            rendered,
            page_size=page_size,
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
            footer_right=footer,
            footer_font_name='Courier',
            footer_font_size=6,
            # zoom=zoom,
        )
        content = ContentFile(file)
        return content

    def save_legacy_oss(self):
        content = self.get_legacy_oss()
        self.legacy_oss.save('legacy_oss', content)

    def get_participants_emails(self):
        Entry = apps.get_model('registration.entry')
        appearances = self.appearances.distinct()
        for appearance in appearances:
            entry = Entry.objects.filter(
                group_id=appearance.group_id,
                session_id=self.session_id
            ).first()
            seen = set()
            if len(entry.notification_list):
                result = [
                    "{0}".format(email,)
                    for email in entry.notification_list.split(',')
                    if not (
                        "{0}".format(email,) in seen or seen.add(
                            "{0}".format(email,)
                        )
                    )
                ]
        return result

    def get_titles(self):
        Chart = apps.get_model('bhs.chart')
        Group = apps.get_model('bhs.group')
        Song = apps.get_model('adjudication.song')
        appearances = self.appearances.filter(
            draw__gt=0,
        ).order_by(
            'draw',
        )
        for appearance in appearances:
            songs = appearance.songs.order_by('num')
            titles = []
            for song in songs:
                try:
                    chart = Chart.objects.get(id=song.chart_id)
                except Chart.DoesNotExist:
                    chart = None
                try:
                    title = chart.title
                except AttributeError:
                    title = "Unknown (Not in Repertory)"
                row = "{0} Song {1}: {2}".format(
                    song.appearance.round.get_kind_display(),
                    song.num,
                    title,
                )
                titles.append(row)
            appearance.titles_patched = titles

        context = {
            'appearances': appearances,
            'round': self,
        }
        rendered = render_to_string('reports/titles.html', context)
        file = pydf.generate_pdf(
            rendered,
            page_size='Letter',
            orientation='Portrait',
            margin_top='5mm',
            margin_bottom='5mm',
        )
        content = ContentFile(file)
        return content

    def get_announcements(self):
        Panelist = apps.get_model('adjudication.panelist')
        Award = apps.get_model('bhs.award')
        Group = apps.get_model('bhs.group')
        Appearance = apps.get_model('adjudication.appearance')
        appearances = self.appearances.filter(
            draw__gt=0,
        ).order_by(
            'draw',
        )
        group_ids = appearances.values_list('group_id', flat=True)
        mt = self.appearances.filter(
            draw=0,
        ).first()
        outcomes = self.outcomes.filter(
            printed=True,
        ).order_by(
            '-tree_sort',
            '-num',
        )
        if self.kind == self.KIND.finals:
            winners = self.appearances.exclude(num=0).exclude(
                status__in=[
                    Appearance.STATUS.built,
                    Appearance.STATUS.disqualified,
                    Appearance.STATUS.scratched,
                ],
            ).exclude(
                is_private=True,
            ).order_by(
                'stats__tot_points',
                'stats__sng_points',
                'stats__per_points',
            )


            # groups = Group.objects.filter(
            #     appearances__round__session_id=self.session_id,
            # ).exclude(
            #     status__in=[
            #         Appearance.STATUS.disqualified,
            #         Appearance.STATUS.scratched,
            #     ],
            # ).prefetch_related(
            #     'appearances__songs__scores',
            #     'appearances__songs__scores__panelist',
            #     'appearances__round__session',
            # ).annotate(
            #     tot_points=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__round__session_id=self.session_id,
            #             appearances__round__num__lte=self.num,
            #         ),
            #     ),
            #     sng_points=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.singing,
            #             appearances__round__session_id=self.session_id,
            #             appearances__round__num__lte=self.num,
            #         ),
            #     ),
            #     per_points=Sum(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__songs__scores__panelist__category=Panelist.CATEGORY.performance,
            #             appearances__round__session_id=self.session_id,
            #             appearances__round__num__lte=self.num,
            #         ),
            #     ),
            #     raw_score=Avg(
            #         'appearances__songs__scores__points',
            #         filter=Q(
            #             appearances__songs__scores__panelist__kind=Panelist.KIND.official,
            #             appearances__round__session_id=self.session_id,
            #             appearances__round__num__lte=self.num,
            #         ),
            #     ),
            #     tot_score=Func(
            #         F('raw_score'),
            #         function='ROUND',
            #         template='%(function)s(%(expressions)s, 1)'
            #     ),
            # ).order_by(
            #     'tot_points',
            #     'sng_points',
            #     'per_points',
            # )
            winners = list(winners)[-5:]
        else:
            winners = None
        pos = self.appearances.aggregate(sum=Sum('pos'))['sum']
        # context = {
        #     'round': self,
        #     'appearances': appearances,
        #     'mt': mt,
        #     'outcomes': outcomes,
        #     'groups': groups,
        #     'pos': pos,
        # }

        # rendered = render_to_string('reports/announcements.html', context)
        # file = pydf.generate_pdf(
        #     rendered,
        #     page_size='Letter',
        #     orientation='Portrait',
        #     margin_top='5mm',
        #     margin_bottom='5mm',
        # )
        # content = ContentFile(file)

        document = Document()

        # document.add_heading(
        #     'Announcements {0}'.format(self)
        # )
        document.add_heading('Evaluations')
        document.add_paragraph(
            "Remember to state Start time(s) for each round, Rotation time and general location. "
        )
        document.add_heading('Administration')
        document.add_paragraph(
            "CSA: sent to group admins"
        )
        if pos:
            document.add_paragraph(
                "Total participants on stage: {0}".format(pos)
            )
        document.add_heading('Awards')
        for outcome in outcomes:
            award = Award.objects.get(id=outcome.award_id)

            if outcome is None or outcome.winner is None or len(outcome.winner) == 0:
                if outcome.level == Outcome.LEVEL.manual or outcome.level == Outcome.LEVEL.raw:
                    document.add_paragraph("{0}: {1}".format(award.name, "No winner was entered"))
                else:
                    document.add_paragraph("{0}: {1}".format(award.name, "This award does not have a winner"))
            else:
                document.add_paragraph("{0}: {1}".format(award.name, html.unescape(outcome.winner)))

        if appearances:
            document.add_heading('Draw')
            for appearance in appearances:
                group = Group.objects.get(id=appearance.group_id)
                document.add_paragraph(
                    "{0}: {1}".format(appearance.draw, group.name),
                    # style='List Bullet',
                )
            if mt is not None:
                gp = Group.objects.get(id=mt.group_id)
                document.add_paragraph(
                    "MT: {0}".format(html.unescape(gp.name)),
                    # style='List Bullet',
                )            
        if winners:
            document.add_heading('Results')
            for winner in winners:
                group = Group.objects.get(id=winner.group_id)
                document.add_paragraph(
                    "With a score of {0}, a {1} average: {2}".format(
                        winner.stats['tot_points'],
                        winner.stats['tot_score'],
                        html.unescape(group.name),
                    ),
                    style='List Bullet',
                )

        buff = BytesIO()
        document.save(buff)
        content = ContentFile(buff.getvalue())
        return content

    def get_labels(self, request, data):
        Appearance = apps.get_model('adjudication.appearance')
        Panelist = apps.get_model('adjudication.panelist')
        Round = apps.get_model('adjudication.round')
        Entry = apps.get_model('registration.entry')


        # Default Filters...
        appearance_filters = {}
        panelist_filters = {}
        panelist_filters.update({'round_id': self.id})
        appearance_filters.update({'round_id': self.id})

        # All Judges
        # All Contestants

        # Specific Judge
        if len(data) and isinstance(data['judges'], Iterable) and len(data['judges']):
            judge_ids = []
            for p in data['judges']:
                judge_ids.append(p['personId'])
            panelist_filters.update({'person_id__in': judge_ids})

        # Specific Contestant
        if len(data) and isinstance(data['groups'], Iterable) and len(data['groups']):
            appearance_ids = []
            for a in data['groups']:
                appearance_ids.append(a['groupId'])
            appearance_filters.update({'group_id__in': appearance_ids})

        ### ORDER
            # Category
                # Judge Last Name
                    # Round
                        # Order of Appearance

        # ### JUDGES ###
        panelists = Panelist.objects.select_related('round').filter(
                round__id=self.id,
                category__gt=Panelist.CATEGORY.adm,
                **panelist_filters
            ).order_by(
                'category',
                'last_name',
                'round__session_kind',
                'round__kind',
            )

        ## Set via POST var
        if len(data) and 'isReversed' in data:
            reverse_order = data['isReversed']
        else:
            reverse_order = False

        # Reverse Order
        appearance_order = 'num'
        if reverse_order:
            appearance_order = "-{0}".format(appearance_order)

        ### APPEARANCES ###
        appearances = Appearance.objects.select_related('round').filter(
                round__id=self.id,
                **appearance_filters
            ).order_by(
                'round__session_kind',
                'round__kind',
                appearance_order
            )

        # Create new document
        document = "{\\rtf1\\ansi\\ansicpg1252\\cocoartf2513\n"
        document += "\\cocoatextscaling0\\cocoaplatform0{\\fonttbl\\f0\\fswiss\\fcharset0 ArialMT;\\f1\\fswiss\\fcharset0 Arial-BoldMT;}\n"
        document += "{\\colortbl;\\red255\\green255\\blue255;\\red0\\green0\\blue0;}\n"
        document += "{\\*\\expandedcolortbl;;\\cssrgb\\c0\\c0\\c0;}\n"
        document += "{\\info\n"
        document += "{\\title LEFTLEFTLEFT}\n"
        document += "{\\author Barberscore}}\\margl720\\margr720\\margb720\\margt720\\vieww22680\\viewh14600\\viewkind0\n"
        document += "\\deftab720\n"
        document += "\\pard\\tqc\\tx5400\\tqr\\tx10800\\pardeftab720\\ri0\\partightenfactor0\n"
        document += "\n"
        document += "\\f0\\fs24 \\cf2 \\\n"

        directors = {}
        completed = []
        for judge in panelists:
            if judge.person_id not in completed:
                for appearance in appearances:

                    document += "\\page {0}. {1}\t\\\n".format(
                        (appearance.num if appearance.num != 0 else 'MT'),
                        html.unescape(appearance.name)
                        )

                    if appearance.kind == Appearance.KIND.chorus and appearance.group_id not in directors:
                        entry = Entry.objects.filter(
                            group_id=appearance.group_id,
                            session_id=appearance.round.session_id,
                        ).first()
                        directors[appearance.group_id] = entry.participants
                        document += "        {0}\t\t{1}\n".format(directors[appearance.group_id], judge.last_name)
                    else:
                        document += "         \t\t{0}\n".format(judge.last_name)

                    document += "\\f1\\b   \n"
                    document += "\\fs44 #{0}\n".format(judge.num)
                    document += "\\f0\\b0\\fs24 \\\n"
                    document += "\\pard\\tqc\\tx5400\\tqr\\tx10800\\pardeftab720\\ri0\\qc\\partightenfactor0\n"
                    document += "\\cf2 {0} District - {1} {2}, {3}\\\n".format(
                            self.DISTRICT_FULL_NAMES[self.district],
                            Appearance.KIND[appearance.round.session_kind],
                            Round.KIND[appearance.round.kind],
                            appearance.round.date.strftime("%-m/%-d/%Y"),                            
                    )
                    document += "\\pard\\tqc\\tx5400\\tqr\\tx10800\\pardeftab720\\ri0\\partightenfactor0\n"
                    document += "\\cf2 \t\t\\\n"

                completed.append(judge.person_id)

        # End of File...
        document += "\\\n"
        document += "}"

        content = ContentFile(document)
        return content

    def base_filename(self):
        return '{0}{1}'.format(
            self.get_district_display(),
            self.start_date.strftime("%Y%m%d")
        )

    def get_judge_emails(self):
        postfix = ''
        if (settings.EMAIL_ADMINS_ONLY):
            postfix = '.invalid'
        Panelist = apps.get_model('adjudication.panelist')
        judges = self.panelists.filter(
            # status=Panelist.STATUS.active,
            category__gt=Panelist.CATEGORY.adm,
            # person__email__isnull=False,
        ).order_by(
            'kind',
            'category',
            'last_name',
            'first_name',
        )
        seen = set()
        result = [
            "{0} ({1} {2}) <{3}{4}>".format(judge.name, judge.get_kind_display(), judge.get_category_display(), judge.email, postfix,)
            for judge in judges
            if not (
                "{0} ({1} {2}) <{3}{4}>".format(judge.name, judge.get_kind_display(), judge.get_category_display(), judge.email, postfix,) in seen or seen.add(
                    "{0} ({1} {2}) <{3}{4}>".format(judge.name, judge.get_kind_display(), judge.get_category_display(), judge.email, postfix,)
                )
            )
        ]
        return result

    def mock(self):
        """Mock Round"""
        Appearance = apps.get_model('adjudication.appearance')
        if self.status != self.STATUS.started:
            raise RuntimeError("Round not Started")
        appearances = self.appearances.exclude(
            status=Appearance.STATUS.verified,
        ).exclude(
            # Don't include disqualifications.
            status=Appearance.STATUS.disqualified,
        ).exclude(
            # Don't include scratches.
            status=Appearance.STATUS.scratched,
        )
        for appearance in appearances:
            appearance.mock()
            appearance.save()
        return

    def get_publish_email(self):
        Appearance = apps.get_model('adjudication.appearance')
        Group = apps.get_model('bhs.group')
        Panelist = apps.get_model('adjudication.panelist')
        completes = self.appearances.filter(
            is_private=False,
            status=Appearance.STATUS.completed,
        ).order_by(
            '-stats__tot_points',
            '-stats__sng_points',
            '-stats__per_points',
        )

        # Draw Block
        if self.kind != self.KIND.finals:
            # Get advancers
            advancers = self.appearances.filter(
                status=Appearance.STATUS.advanced,
            ).order_by(
                'draw',
            ).values_list(
                'draw',
                'name',
            )
            advancers = list(advancers)
            try:
                mt = self.appearances.get(
                    draw=0,
                ).name
                advancers.append(('MT', mt))
            except Appearance.DoesNotExist:
                pass
        else:
            advancers = None

        # Outcome Block
        items = self.outcomes.select_related(
            # 'award',
        ).order_by(
            'num',
        ).values_list(
            'num',
            # 'award__name',
            'name',
            'winner'
        )

        outcomes = []
        for item in items:
            outcomes.append(
                (
                    "{0} {1}".format(item[0], item[1]),
                    item[2],
                )
            )

        context = {
            'round': self,
            'advancers': advancers,
            'completes': completes,
            'outcomes': outcomes,
        }
        template = 'emails/round_publish.txt'
        subject = "[Barberscore] {0} Results and OSS".format(
            self,
        )
        to = self.get_owners_emails()
        cc = self.get_judge_emails()
        bcc = self.get_participants_emails()

        if self.oss_report:
            pdf = self.oss_report.file
        else:
            pdf = self.get_oss()
        file_name = '{0} OSS.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]
        context['bcc'] = [x.partition(" <")[0] for x in bcc]

        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            bcc=bcc,
            attachments=attachments,
        )
        return email

    def send_publish_email(self):
        email = self.get_publish_email()
        if self.status != self.STATUS.published:
            raise RuntimeError("Round not published")
        return email.send()

    def get_publish_report_email(self):
        template = 'emails/round_publish_report.txt'
        context = {
            'round': self,
        }
        subject = "[Barberscore] {0} Reports and SA".format(
            self.nomen,
        )
        to = self.get_owners_emails()
        cc = self.get_judge_emails()
        attachments = []
        if self.sa_report:
            pdf = self.sa_report.file
        else:
            pdf = self.get_sa()
        file_name = '{0} SA.pdf'.format(self)
        attachments = [(
            file_name,
            pdf,
            'application/pdf',
        )]
        email = build_email(
            template=template,
            context=context,
            subject=subject,
            to=to,
            cc=cc,
            attachments=attachments,
        )
        return email

    def send_publish_report_email(self):
        email = self.get_publish_report_email()
        return email.send()        

    # Round Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        return True

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'CA',
                'PC',
                'ADM',
            ]
        ))

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        if self.status == self.STATUS.published:
            return False
        return bool(any([
            request.user in self.owners.all(),
        ]))

    # Round Conditions
    def can_build(self):
        return True

    def can_complete(self):
        Appearance = apps.get_model('adjudication.appearance')

        # Confirm verified performances exist
        valid = all([
            self.appearances.filter(
                num__gt=0,
                status__in=[
                    Appearance.STATUS.verified,
                    Appearance.STATUS.disqualified,
                    Appearance.STATUS.scratched,
                ]
            ),
        ])

        # Confirm all performances have been verified
        valid = not all([
            self.appearances.filter(
                num__gt=0,
                status__in=[
                    Appearance.STATUS.completed,
                    Appearance.STATUS.new,
                    Appearance.STATUS.built,
                    Appearance.STATUS.started,
                    Appearance.STATUS.finished,
                    Appearance.STATUS.variance,
                ]
            ),
        ])

        return valid

    def can_finalize(self):
        return all([
            not self.outcomes.filter(
                name='MUST ENTER WINNER MANUALLY',
            ),
        ])

    def can_publish(self):
        return True

    def determine_winners(self):
        # Run outcomes
        outcomes = self.outcomes.all()
        for outcome in outcomes:
            if outcome.level is not Outcome.LEVEL.raw:
                outcome.winner = outcome.get_winner()
                outcome.save()
        return     

    # Round Transitions
    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source='*',
        target=STATUS.new,
    )
    def reset(self, *args, **kwargs):
        self.oss_report.delete()
        self.sa_report.delete()
        panelists = self.panelists.all()
        appearances = self.appearances.all()
        outcomes = self.outcomes.all()
        panelists.delete()
        outcomes.delete()
        appearances.delete()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=STATUS.new,
        target=STATUS.built,
        conditions=[can_build],
    )
    def build(self, *args, **kwargs):
        """Build the Round"""
        # Reset for indempodence
        self.reset()

        # Fetch models
        Session = apps.get_model('registration.session')
        Assignment = apps.get_model('registration.assignment')
        Entry = apps.get_model('registration.entry')
        Round = apps.get_model('adjudication.round')
        Chart = apps.get_model('bhs.chart')

        # Get objects for build
        session = Session.objects.get(id=self.session_id)

        # Check for initial round
        if self.num == 1:
            # If first Round
            prior_round = None

        else:
            # If not first round
            prior_round = Round.objects.get(
                session_id=session.id,
                num=self.num-1,
            )

        # Build panel from assignments
        cas = session.assignments.filter(
            category__in=[
                Assignment.CATEGORY.pc,
                Assignment.CATEGORY.adm,
            ]
        ).order_by(
            'kind',
            'category',
            'last_name',
            'first_name',
        )
        judges = session.assignments.filter(
            category__in=[
                Assignment.CATEGORY.singing,
                Assignment.CATEGORY.performance,
                Assignment.CATEGORY.music,
            ],
            kind=Assignment.KIND.official,
        ).order_by(
            'kind',
            'category',
            'last_name',
            'first_name',
        )
        practices = session.assignments.filter(
            category__in=[
                Assignment.CATEGORY.singing,
                Assignment.CATEGORY.performance,
                Assignment.CATEGORY.music,
            ],
            kind=Assignment.KIND.practice,
        ).order_by(
            'kind',
            'category',
            'last_name',
            'first_name',
        )

        for ca in cas:
            area = ca.get_district_display() if ca.district else ''
            self.panelists.create(
                kind=ca.kind,
                category=ca.category,
                person_id=ca.person_id,
                name=ca.name,
                first_name=ca.first_name,
                last_name=ca.last_name,
                area=area,
                email=ca.email,
                cell_phone=ca.cell_phone,
                bhs_id=ca.bhs_id,
            )

        i = 0
        for judge in judges:
            i += 1
            area = judge.get_district_display() if judge.district else ''
            self.panelists.create(
                num=i,
                kind=judge.kind,
                category=judge.category,
                person_id=judge.person_id,
                name=judge.name,
                first_name=judge.first_name,
                last_name=judge.last_name,
                area=area,
                email=judge.email,
                cell_phone=judge.cell_phone,
                bhs_id=judge.bhs_id,
            )
        i = 50
        for practice in practices:
            i += 1
            area = practice.get_district_display() if practice.district else ''
            self.panelists.create(
                num=i,
                kind=practice.kind,
                category=practice.category,
                person_id=practice.person_id,
                name=practice.name,
                first_name=practice.first_name,
                last_name=practice.last_name,
                area=area,
                email=practice.email,
                cell_phone=practice.cell_phone,
                bhs_id=practice.bhs_id,
            )

        # Build Outcomes (Awards)
        if not prior_round:
            # Put all contests with entries in the round for first round
            contests = session.contests.filter(
                entries__isnull=False,
            ).order_by(
                'tree_sort',
            ).distinct()
        else:
            # for subsequent rounds, create from prior outcomes
            # that are not single-round awards
            contests = prior_round.outcomes.filter(
                is_single=False,
            )
        i = 0
        for contest in contests:
            # Only number contests on initial; otherwise use existing.
            i += 1
            if prior_round:
                num = contest.num
            else:
                num = i
            self.outcomes.create(
                num=num,
                award_id=contest.award_id,
                name=contest.name,
                kind=contest.kind,
                gender=contest.gender,
                level=contest.level,
                season=contest.season,
                district=contest.district,
                division=contest.division,
                age=contest.age,
                is_novice=contest.is_novice,
                is_single=contest.is_single,
                tree_sort=contest.tree_sort,
            )

        # Create Appearances
        if not prior_round:
            # If first round, build appearances from entries
            entries = session.entries.filter(
                status=Entry.STATUS.approved,
            ).order_by('draw')
            for entry in entries:
                # Force draw = 0 for MTs
                if entry.is_mt:
                    entry.draw = 0
                is_single = not bool(entry.contests.filter(is_single=False))
                charts_raw = Chart.objects.filter(
                    groups__id=entry.group_id,
                ).values(
                    'id',
                    'title',
                    'arrangers',
                ).order_by('title')
                for c in charts_raw:
                    c['pk'] = str(c.pop('id'))
                charts = [json.dumps(x) for x in charts_raw]

                if entry.kind == entry.KIND.quartet:
                    # area = entry.get_district_display()
                    # entry object here is from the registration_entry table.
                    # need to determine the correct value for quartets...
                    area = entry.area
                else:
                    area = entry.chapters

                appearance = self.appearances.create(
                    num=entry.draw,
                    is_private=entry.is_private,
                    is_single=is_single,
                    participants=entry.participants,
                    area=area,
                    group_id=entry.group_id,
                    entry_id=entry.id,
                    name=entry.name,
                    kind=entry.kind,
                    gender=entry.gender,
                    district=entry.district,
                    division=entry.division,
                    bhs_id=entry.bhs_id,
                    code=entry.code,
                    base=entry.base,
                    image_id=entry.image_id,
                    charts=charts,
                )
                appearance.owners.set(entry.owners.all())
                award_ids = list(entry.contests.values_list('award_id', flat=True))
                outcomes = self.outcomes.filter(award_id__in=award_ids)
                appearance.outcomes.set(outcomes)
        else:
            # If subsequent round, build appearances based on
            # ADVANCERS from prior round.
            prior_appearances = prior_round.appearances.filter(
                status=Appearance.STATUS.advanced,
            )
            for prior_appearance in prior_appearances:
                charts_raw = Chart.objects.filter(
                    groups__id=prior_appearance.group_id,
                ).values(
                    'id',
                    'title',
                    'arrangers',
                ).order_by('title')
                for c in charts_raw:
                    c['pk'] = str(c.pop('id'))
                charts = [json.dumps(x) for x in charts_raw]
                
                # Create and start group
                appearance = self.appearances.create(
                    num=prior_appearance.draw,
                    is_private=prior_appearance.is_private,
                    is_single=prior_appearance.is_single,
                    participants=prior_appearance.participants,
                    area=prior_appearance.area,
                    group_id=prior_appearance.group_id,
                    entry_id=prior_appearance.entry_id,
                    name=prior_appearance.name,
                    kind=prior_appearance.kind,
                    gender=prior_appearance.gender,
                    district=prior_appearance.district,
                    division=prior_appearance.division,
                    bhs_id=prior_appearance.bhs_id,
                    code=prior_appearance.code,
                    base=prior_appearance.base,
                    image_id=prior_appearance.image_id,
                    charts=prior_appearance.charts,
                )
                appearance.owners.set(prior_appearance.owners.all())
                award_ids = list(prior_appearance.outcomes.values_list('award_id', flat=True))
                outcomes = self.outcomes.filter(
                    award_id__in=award_ids
                )
                appearance.outcomes.set(outcomes)
            # Create MT
            mts = prior_round.appearances.filter(
                draw__lte=0,
            )
            for mt in mts:
                appearance = self.appearances.create(
                    num=mt.draw,
                    is_private=mt.is_private,
                    is_single=mt.is_single,
                    participants=mt.participants,
                    area=mt.area,
                    group_id=mt.group_id,
                    entry_id=mt.entry_id,
                    name=mt.name,
                    kind=mt.kind,
                    gender=mt.gender,
                    district=mt.district,
                    division=mt.division,
                    bhs_id=mt.bhs_id,
                    code=mt.code,
                    base=mt.base,
                    image_id=mt.image_id,
                    charts=mt.charts,
                )
                appearance.owners.set(mt.owners.all())
                award_ids = list(mt.outcomes.values_list('award_id', flat=True))
                outcomes = self.outcomes.filter(
                    award_id__in=award_ids
                )
                appearance.outcomes.set(outcomes)
            return


    @notification_user
    @fsm_log_by
    @transition(field=status, source=[STATUS.built], target=STATUS.started)
    def start(self, *args, **kwargs):
        # Build the appearances
        appearances = self.appearances.all()
        for appearance in appearances:
            appearance.build()
            appearance.save()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.started, STATUS.completed,],
        target=STATUS.completed,
        conditions=[can_complete],)
    def complete(self, *args, **kwargs):
        Appearance = apps.get_model('adjudication.appearance')
        Panelist = apps.get_model('adjudication.panelist')

        self.determine_winners()

        # If there is no next round simply return
        if self.kind == self.KIND.finals:
            return

        # Otherwise, figure out the Draw.
        # First, get spots available
        spots = self.spots

        # Get all multi appearances and annotate average.
        multis = self.appearances.filter(
            status=Appearance.STATUS.verified,
            is_single=False,
            num__gt=0
        ).annotate(
            avg=Avg(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                )
            ),
            tot_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                )
            ),
            sng_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.singing,
                )
            ),
            per_points=Sum(
                'songs__scores__points',
                filter=Q(
                    songs__scores__panelist__kind=Panelist.KIND.official,
                    songs__scores__panelist__category=Panelist.CATEGORY.performance,
                )
            ),
        )
        # If spots are constricted, find those who advance
        mt = None
        if spots:
            # All those above 75.0 advance automatically, regardless of spots available
            if self.get_district_display() == 'BHS':
                ordered = multis.order_by(
                    '-tot_points',
                    '-sng_points',
                    '-per_points',
                )
                advancer__ids = [x.id for x in ordered[:spots]]
                try:
                    mt = ordered[spots:spots+1][0]                
                except IndexError:
                    # Index is not present
                    # ...usually when the number of spots exceeds the number of competitors.
                    pass
            else:
                automatics = multis.filter(
                    avg__gte=75.0,
                )
                # generate list of the advancers, as appearance IDs
                advancer__ids = [x.id for x in automatics]
                # Generate remaining multi appearances
                remains = multis.exclude(
                    id__in=advancer__ids,
                ).order_by(
                    '-tot_points',
                    '-sng_points',
                    '-per_points',
                )
                # Figure out remaining spots
                diff = spots - automatics.count()
                # If there are additional remaining spots, add them up to available
                if diff > 0:
                    adds = remains[:diff]
                    for a in adds:
                        advancer__ids.append(a.id)
                    try:
                        mt = remains[diff:diff+1][0]
                    except IndexError:
                        mt = None
                else:
                    # If MT available add, otherwise none
                    try:
                        mt = remains.first()
                    except AttributeError:
                        mt = None
        # Otherwise, advance all
        else:
            advancer__ids = [a.id for a in multis]

        # Reset draw
        self.appearances.update(draw=None)

        # Randomize the advancers and set the initial draw
        appearances = self.appearances.filter(
            id__in=advancer__ids,
        ).order_by('?')
        i = 1
        for appearance in appearances:
            appearance.draw = i
            appearance.save()
            i += 1
        # create Mic Tester at draw 0
        if mt:
            mt.draw = 0
            mt.save()
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.completed],
        target=STATUS.finalized,
        conditions=[can_finalize],)
    def finalize(self, *args, **kwargs):
        Appearance = apps.get_model('adjudication.appearance')
        Panelist = apps.get_model('adjudication.panelist')
        completed_appearances = self.appearances.filter(
            status=Appearance.STATUS.verified,
        ).exclude(
            draw__gt=0,
        )
        for appearance in completed_appearances:
            appearance.complete()
            appearance.save()
        advancing_appearances = self.appearances.filter(
            status=Appearance.STATUS.verified,
            draw__gt=0,
        )
        for appearance in advancing_appearances:
            appearance.advance()
            appearance.save()
        panelists = self.panelists.filter(
            category__gt=Panelist.CATEGORY.adm,
        )
        for panelist in panelists:
            panelist.release()
            panelist.save()
        # Saves reports through transition signal to avoid race condition
        return

    @notification_user
    @fsm_log_by
    @transition(
        field=status,
        source=[STATUS.finalized],
        target=STATUS.published,
        conditions=[can_publish],)
    def publish(self, *args, **kwargs):
        """Publishes the results and notifies all parties"""
        Appearance = apps.get_model('adjudication.appearance')
        Panelist = apps.get_model('adjudication.panelist')
        # Send the OSS
        send_publish_email_from_round.delay(self)
        # Send the CSAs
        completed_appearances = self.appearances.filter(
            status=Appearance.STATUS.completed,
        )
        for appearance in completed_appearances:
            send_complete_email_from_appearance.delay(appearance)
        # Send the SAs
        send_publish_report_email_from_round.delay(self)
        # # Send the PSAs
        # panelists = self.panelists.filter(
        #     category__gt=Panelist.CATEGORY.adm,
        #     status=Panelist.STATUS.released,
        # )
        # for panelist in panelists:
        #     send_psa_email_from_panelist.delay(panelist)
        return


class Score(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
        (10, 'verified', 'Verified',),
        # (20, 'entered', 'Entered',),
        (25, 'cleared', 'Cleared',),
        (30, 'flagged', 'Flagged',),
        (35, 'revised', 'Revised',),
        (40, 'confirmed', 'Confirmed',),
        # (50, 'final', 'Final',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    # CATEGORY = Choices(
    #     (30, 'music', 'Music'),
    #     (40, 'performance', 'Performance'),
    #     (50, 'singing', 'Singing'),
    # )

    # category = models.IntegerField(
    #     choices=CATEGORY,
    # )

    # KIND = Choices(
    #     (10, 'official', 'Official'),
    #     (20, 'practice', 'Practice'),
    #     (30, 'composite', 'Composite'),
    # )

    # kind = models.IntegerField(
    #     choices=KIND,
    # )

    # num = models.IntegerField(
    # )

    points = models.IntegerField(
        help_text="""
            The number of points (0-100)""",
        null=True,
        blank=True,
        validators=[
            MaxValueValidator(
                100,
                message='Points must be between 0 - 100',
            ),
            MinValueValidator(
                0,
                message='Points must be between 0 - 100',
            ),
        ]
    )

    # FKs
    song = models.ForeignKey(
        'Song',
        related_name='scores',
        on_delete=models.CASCADE,
    )

    panelist = models.ForeignKey(
        'Panelist',
        related_name='scores',
        on_delete=models.CASCADE,
    )

    objects = ScoreManager()

    # class Meta:
    #     unique_together = (
    #         ('song', 'panelist',),
    #     )

    class JSONAPIMeta:
        resource_name = "score"

    def __str__(self):
        return str(self.id)


    # Score Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        # Assigned owners can always see
        if request.user in self.song.appearance.round.owners.all():
            return True
        return False

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'CA',
                'PC',
                'ADM',
            ]
        ))

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        if self.song.appearance.round.status == self.song.appearance.round.STATUS.published:
            return False
        return request.user in self.song.appearance.round.owners.all()


class Song(TimeStampedModel):
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
    )

    STATUS = Choices(
        (0, 'new', 'New',),
    )

    status = FSMIntegerField(
        help_text="""DO NOT CHANGE MANUALLY unless correcting a mistake.  Use the buttons to change state.""",
        choices=STATUS,
        default=STATUS.new,
    )

    num = models.IntegerField(
    )

    asterisks = ArrayField(
        base_field=models.IntegerField(
        ),
        default=list,
        blank=True,
    )

    dixons = ArrayField(
        base_field=models.IntegerField(
        ),
        default=list,
        blank=True,
    )

    PENALTY = Choices(
        (30, 'repetition', 'Repeating Substantial Portions of a Song (V.A.2)'),
        (32, 'accompaniment', 'Instrumental Accompaniment (IX.A.2.a)'),
        (34, 'texture', 'Chorus Exceeding 4-Part Texture (IX.A.2.b)'),
        (36, 'excessive', 'Excessive Melody Not in Inner Part (IX.A.2.c)'),
        (38, 'progression', 'Lack of Characteristic Chord Progression (IX.A.2.d)'),
        (39, 'lyrics', 'Excessive Lyrics < 4 parts (IX.A.2.e)'),
        (40, 'patreg', 'Primarily Patriotic/Religious Intent (IX.A.3)'),
        (44, 'notingoodtaste', 'Not in Good Taste (IX.A.3.b)'),
        (48, 'nonmembersperforming', 'Non-members Performing on Stage (XI.A.1)'),
        (50, 'enhancement', 'Sound Equipment or Electronic Enhancement (X.B.1-3)'),
    )

    penalties = ArrayField(
        base_field=models.IntegerField(
            choices=PENALTY,
        ),
        default=list,
        blank=True,
    )

    stats = JSONField(
        null=True,
        blank=True,
    )

    # Chart Denorm
    chart_id = models.UUIDField(
        null=True,
        blank=True,
    )

    name = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    title = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    arrangers = models.CharField(
        max_length=255,
        blank=True,
        default='',
    )

    # FKs
    appearance = models.ForeignKey(
        'Appearance',
        related_name='songs',
        on_delete=models.CASCADE,
    )

    # Internals
    objects = SongManager()

    class Meta:
        # unique_together = (
        #     ('appearance', 'num',),
        # )
        get_latest_by = ['num']

    class JSONAPIMeta:
        resource_name = "song"

    def __str__(self):
        return str(self.id)

    # Methods
    def get_asterisks(self):
        """
        Check to see if the song produces a category variance (asterisk)

        Returns a list of categories that produced an asterisk.
        """
        # Set Flag
        asterisks = []
        # Get Averages by category
        categories = self.scores.filter(
            panelist__kind=10,
        ).values(
            'panelist__category',
        ).annotate(
            avg=Avg('points'),
        )
        for category in categories:
            category_scores = self.scores.filter(
                panelist__category=category['panelist__category'],
                panelist__kind=10,
            )
            for score in category_scores:
                is_asterisk = abs(score.points - category['avg']) > 5
                if is_asterisk:
                    asterisks.append(category['panelist__category'])
                    continue
        asterisks = list(set(asterisks))
        return asterisks

    def get_dixons(self):
        """
        Check to see if the song produces a spread error (Dixon's Q)

        Returns a list of categories that produced a Dixon's Q.
        """
        # Set flag
        output = []
        # Confidence thresholds
        confidence = {
            '3': 0.941,
            '6': .56,
            '9': .376,
            '12': .437,
            '15': .338,
        }
        # Only use official scores.
        scores = self.scores.filter(
            panelist__kind=10,
        )
        # Get the totals
        aggregates = scores.aggregate(
            cnt=Count('id'),
            max=Max('points'),
            min=Min('points'),
            spread=Max('points') - Min('points'),
        )
        # Check for validity.
        if aggregates['cnt'] < 3:
            return RuntimeError('Panel too small error')
        # Bypass to avoid division by zero
        if not aggregates['spread']:
            return output
        # Order the scores
        ascending = scores.order_by('points')
        descending = scores.order_by('-points')
        # Separate check for single-panel
        if aggregates['cnt'] == '3':
            if abs(ascending[0].points - ascending[1].points) >= 10:
                output.append(ascending[0].panelist.category)
            if abs(descending[0].points - descending[1].points) >= 10:
                output.append(descending[0].panelist.category)
            return output

        # Otherwise, run the checks, both ascending and descending
        if str(aggregates['cnt']) in confidence:
            critical = confidence[str(aggregates['cnt'])]
            ascending_distance = abs(ascending[0].points - ascending[1].points)
            ascending_q = ascending_distance / aggregates['spread']
            if ascending_q > critical and ascending_distance > 4:
                output.append(ascending[0].panelist.category)
            descending_distance = abs(descending[0].points - descending[1].points)
            descending_q = descending_distance / aggregates['spread']
            if descending_q > critical and descending_distance > 4:
                output.append(descending[0].panelist.category)
        return output


    # Song Permissions
    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_read_permission(request):
        return True

    @allow_staff_or_superuser
    @authenticated_users
    def has_object_read_permission(self, request):
        if self.appearance.round.status == self.appearance.round.STATUS.published:
            return True
        return request.user in self.appearance.round.owners.all()

    @staticmethod
    @allow_staff_or_superuser
    @authenticated_users
    def has_write_permission(request):
        return bool(request.user.roles.filter(
            name__in=[
                'SCJC',
                'CA',
                'PC',
                'ADM',
            ]
        ))


    @allow_staff_or_superuser
    @authenticated_users
    def has_object_write_permission(self, request):
        if self.appearance.round.status == self.appearance.round.STATUS.published:
            return False
        return request.user in self.appearance.round.owners.all()

